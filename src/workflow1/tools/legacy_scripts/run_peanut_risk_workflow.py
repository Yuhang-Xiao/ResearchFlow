from __future__ import annotations

import math
import re
import sys
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd


ROOT = Path(__file__).resolve().parents[1]
RAW_FILE = ROOT / "data/01_raw/PEANUT2023-20241.xlsx"
DOC_DIR = ROOT / "references/notes"
SUMMARY_FILE = ROOT / "references/processed_summaries/peanut_research_plan_summary.md"


def ensure_dirs() -> None:
    for rel in [
        "data/03_primary",
        "data/04_feature",
        "reports",
        "reports/tables",
        "reports/figures",
        "references/processed_summaries",
        "project_state",
    ]:
        (ROOT / rel).mkdir(parents=True, exist_ok=True)


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def to_csv(df: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(path, index=False, encoding="utf-8-sig")


def normalize_text(x) -> str:
    if pd.isna(x):
        return ""
    s = str(x).strip()
    if s.lower() in {"nan", "none", "null"}:
        return ""
    return re.sub(r"\s+", " ", s)


def read_docx_full() -> tuple[Path, list[str], list[list[list[str]]]]:
    docs = list(DOC_DIR.glob("*.docx"))
    target = None
    for p in docs:
        if "物流与供应链管理前言-研究计划-肖宇航" in p.name:
            target = p
            break
    if target is None:
        raise FileNotFoundError("未找到研究计划 Word 文档：references/notes/物流与供应链管理前言-研究计划-肖宇航.docx")
    try:
        from docx import Document

        doc = Document(target)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)
        return target, paragraphs, tables
    except Exception as exc:  # pragma: no cover - reported to user
        raise RuntimeError(f"研究计划 Word 文档无法读取：{type(exc).__name__}: {exc}") from exc


def extract_doc_summary(doc_path: Path, paragraphs: list[str], tables: list[list[list[str]]]) -> str:
    def has_any(text: str, words: list[str]) -> bool:
        return any(w.lower() in text.lower() for w in words)

    sections = {
        "背景": ["AFB1", "AFB₁", "黄曲霉", "风险", "监管", "MOE"],
        "计数面板": ["计数面板", "省份", "月份", "环节", "抽检总批次", "不合格"],
        "清洗表": ["清洗表", "浓度", "单位", "对数正态", "分布参数"],
        "MOE_EDI": ["MOE", "EDI", "暴露", "体重", "消费量", "BMDL"],
        "Beta_Binomial": ["Beta", "Binomial", "贝叶斯", "后验", "先验", "遗忘因子"],
        "POMDP": ["POMDP", "belief", "信念", "隐状态", "观测", "belief-MDP"],
        "DQN": ["DQN", "Deep Q", "强化学习", "动作掩码", "预算", "产能"],
    }
    snippets = {k: [] for k in sections}
    for p in paragraphs:
        for key, words in sections.items():
            if has_any(p, words) and len(snippets[key]) < 5:
                snippets[key].append(p)

    table_md = []
    for table in tables:
        if not table:
            continue
        table_md.append("| " + " | ".join(table[0]) + " |")
        table_md.append("| " + " | ".join(["---"] * len(table[0])) + " |")
        for row in table[1:]:
            table_md.append("| " + " | ".join(row) + " |")

    summary = f"""# 花生/AFB1 风险监管研究计划摘要

## 参考文档

- 文件：`{doc_path.relative_to(ROOT).as_posix()}`
- 读取方式：使用 `python-docx` 提取全文段落与表格；共读取 {len(paragraphs)} 个非空段落、{len(tables)} 个表格。

## 研究主题

研究主题为“基于受限 POMDP 与深度强化学习的食品供应链风险暴露动态监管策略研究”。场景聚焦花生原料供应链中黄曲霉毒素 B1（AFB1）等霉菌毒素污染风险，在预算与抽检产能约束下，构建动态抽检资源配置与风险暴露降低策略。

## 花生/AFB1 风险监管背景

{chr(10).join('- ' + s for s in snippets['背景'][:4])}

## 计数面板的数据结构要求

- 计数面板是后续 Beta-Binomial 信念更新、受限 POMDP / belief-MDP 与 DQN 动态决策的核心数据结构。
- 基本索引应优先为 `省份—月份—供应链环节`；月份不可用时退化为 `省份—年份—供应链环节`。
- 每个索引单元至少需要抽检总批次数、合格/不合格批次数、AFB1 相关记录数、AFB1 相关不合格批次数、浓度可用记录数和数据完整性标记。
{chr(10).join('- ' + s for s in snippets['计数面板'][:3])}

## 清洗表的数据结构要求

- 清洗表服务于浓度分布、暴露评估、风险度量与模型派生特征，不替代计数面板。
- 对 AFB1 相关记录提取浓度数值，进行单位一致化、异常值检查、限量匹配和超标倍数计算。
- 在省份—时间—环节尺度上估计浓度均值、分位数、最大值及 lognormal 参数，用于后续 Monte Carlo 与 MOE/EDI。
{chr(10).join('- ' + s for s in snippets['清洗表'][:3])}

## MOE/EDI 风险度量要求

- 需要 AFB1 浓度分布参数、消费量、平均体重、人口/权重和毒理学基准剂量下限（如 BMDL）等外部参数。
- EDI 用于估计每日摄入暴露，MOE 用于刻画暴露水平与毒理学效应点之间距离；风险惩罚可由 MOE 阈值缺口或风险等级映射得到。
{chr(10).join('- ' + s for s in snippets['MOE_EDI'][:3])}

## Beta-Binomial 信念更新要求

- 观测信号来自每期抽检批次数与不合格/检出批次数。
- 对每个省份—环节维护 Beta 分布超参数，使用二项抽样观测进行共轭更新。
- 跨期传播应引入遗忘因子，使历史信息对当前风险保持适度衰减。
{chr(10).join('- ' + s for s in snippets['Beta_Binomial'][:3])}

## POMDP / belief-MDP 建模要求

- 真实污染率或非合规概率视为不可直接观测的隐状态；抽检结果是带随机性的观测信号。
- 通过信念分布把部分可观测问题转化为 belief-MDP，状态包含 Beta 超参数、后验均值、后验方差、风险特征、预算与产能等。
{chr(10).join('- ' + s for s in snippets['POMDP'][:3])}

## DQN 动态监管优化要求

- DQN 不直接使用真实隐状态，而是在信念—风险特征空间学习政府抽检策略。
- 动作为跨省×环节的整数抽检批次数，需要离散档位、预算硬约束、产能上限和 action masking。
- Reward/损失应同时考虑 MOE 驱动的暴露风险、抽检成本、处置/召回损失和信息价值。
{chr(10).join('- ' + s for s in snippets['DQN'][:4])}

## 当前原始数据优先构建变量

- 时间变量：`年份`、`月份`、`年月`。
- 空间变量：`生产省份_清洗`、`通报单位_清洗`、`抽样地区_候选`。
- 供应链变量：`抽检环节_清洗`、`供应链环节`。
- 产品变量：`原始产品分类`、`原始产品名称`、`产品一级类`、`产品二级类`、`产品细分类`。
- 标签变量：`是否合格`、`是否不合格`、`是否检出`、`是否超标`、`是否AFB1相关`、`AFB1识别依据`。
- 浓度变量：`原始检测数值`、`初检浓度值`、`复检浓度值`、`最终采用浓度值`、`浓度单位`、`法规限量_数值`、`法规限量_单位`、`超标倍数`、`浓度清洗状态`。
- 面板变量：按省份、年月/年份、供应链环节汇总抽检总数、不合格数、AFB1 相关数和浓度可用数。

## 研究计划表格摘录

{chr(10).join(table_md[:30]) if table_md else '无表格。'}
"""
    write_text(SUMMARY_FILE, summary)
    return summary


def classify_role(col: str) -> str:
    c = col.lower()
    rules = [
        ("时间字段", ["日期", "时间", "年月", "生产日期", "通报时间"]),
        ("地区字段", ["省", "地区", "地址", "通报单位"]),
        ("产品字段", ["产品", "食品", "分类", "名称"]),
        ("抽检环节字段", ["抽检环节", "环节"]),
        ("判定结果字段", ["判定结果", "结果"]),
        ("不合格项目字段", ["不合格项目", "不合格规范列"]),
        ("检测数值字段", ["检测数值", "检验项目"]),
        ("法规限制字段", ["法规限制", "限量"]),
        ("AFB1相关字段", ["黄曲霉", "afb", "b1", "b₁"]),
    ]
    for role, keys in rules:
        if any(k.lower() in c for k in keys):
            return role
    return "其他字段"


def schema_inventory(df: pd.DataFrame, sheet_name: str, raw_meta: dict) -> pd.DataFrame:
    rows = []
    n = len(df)
    for col in df.columns:
        s = df[col]
        nonnull = s.dropna()
        sample_values = [normalize_text(x) for x in nonnull.head(5).tolist()]
        rows.append(
            {
                "文件名": RAW_FILE.name,
                "sheet": sheet_name,
                "总行数": n,
                "总列数": len(df.columns),
                "字段名": col,
                "推断类型": str(s.dtype),
                "缺失数": int(s.isna().sum()),
                "缺失率": round(float(s.isna().mean()), 6) if n else 0,
                "唯一值数量": int(s.nunique(dropna=True)),
                "候选角色": classify_role(str(col)),
                "样例值": "；".join(sample_values)[:500],
            }
        )
    inv = pd.DataFrame(rows)
    inv["是否候选主键"] = inv.apply(
        lambda r: "是" if r["唯一值数量"] == raw_meta["rows"] and r["缺失数"] == 0 else "否",
        axis=1,
    )
    return inv


PROVINCES = [
    "北京", "天津", "河北", "山西", "内蒙古", "辽宁", "吉林", "黑龙江", "上海", "江苏", "浙江", "安徽",
    "福建", "江西", "山东", "河南", "湖北", "湖南", "广东", "广西", "海南", "重庆", "四川", "贵州",
    "云南", "西藏", "陕西", "甘肃", "青海", "宁夏", "新疆", "台湾", "香港", "澳门",
]


def clean_province(x) -> str:
    s = normalize_text(x)
    if not s:
        return ""
    for p in PROVINCES:
        if p in s:
            return p
    s = re.sub(r"(省|市|自治区|壮族自治区|回族自治区|维吾尔自治区|特别行政区)$", "", s)
    return s[:20]


def clean_stage(x) -> tuple[str, str]:
    s = normalize_text(x)
    if not s:
        return "", "其他"
    if any(k in s for k in ["种植", "基地", "源头", "初加工", "收购", "仓储", "原料"]):
        stage = "源头/初加工"
    elif any(k in s for k in ["生产", "加工", "生产加工"]):
        stage = "生产"
    elif any(k in s for k in ["流通", "销售", "经营", "市场", "超市", "网抽", "网络", "电商"]):
        stage = "流通"
    elif any(k in s for k in ["餐饮", "食堂", "饭店", "小吃"]):
        stage = "餐饮"
    else:
        stage = "其他"
    return s, stage


def classify_product(cat, name) -> tuple[str, str, str]:
    c = normalize_text(cat)
    n = normalize_text(name)
    text = c + " " + n
    lvl1 = "花生及其制品" if "花生" in text or "仁" in n else (c or "未识别")
    if any(k in text for k in ["花生油", "植物油"]):
        lvl2, lvl3 = "花生油/油脂", "花生油"
    elif any(k in text for k in ["炒货", "坚果", "花生仁", "生干", "烘炒", "油炸", "熟制"]):
        lvl2 = "花生仁/炒货坚果"
        if "熟" in text or "炒" in text or "油炸" in text:
            lvl3 = "熟制/炒货花生"
        elif "生" in text:
            lvl3 = "生干花生仁"
        else:
            lvl3 = "花生仁及坚果制品"
    elif any(k in text for k in ["花生酱", "酱"]):
        lvl2, lvl3 = "花生调味/酱制品", "花生酱"
    elif any(k in text for k in ["糕点", "糖", "饼", "食品"]):
        lvl2, lvl3 = "含花生复合食品", "含花生复合食品"
    else:
        lvl2, lvl3 = (c or "其他花生相关食品"), (n[:30] or "未识别")
    return lvl1, lvl2, lvl3


AFB_PATTERN = re.compile(r"(黄\s*曲\s*霉|AFB\s*1|AFB1|B₁|B1|黄曲霉毒素\s*B)", re.IGNORECASE)
OTHER_TOXIN_PATTERN = re.compile(r"(玉米赤霉烯酮|赭曲霉|脱氧雪腐镰刀菌烯醇|展青霉素|伏马菌素)")


def afb1_detect(row: pd.Series) -> tuple[bool, str, bool]:
    cols = ["不合格项目", "不合格规范列", "检验项目", "检测数值", "备注", "标准"]
    hits = []
    suspicious = False
    for col in cols:
        if col not in row:
            continue
        val = normalize_text(row.get(col))
        if not val:
            continue
        if AFB_PATTERN.search(val):
            hits.append(f"{col}含AFB1关键词")
            if OTHER_TOXIN_PATTERN.search(val) and "黄曲霉" not in val:
                suspicious = True
    if hits:
        return True, "；".join(hits), suspicious
    text = " ".join(normalize_text(row.get(c)) for c in cols if c in row)
    if "生物毒素" in text and ("B" in text or "黄曲霉" in text):
        return False, "生物毒素/B类可疑但未明确AFB1", True
    return False, "", False


UNIT_FACTORS = {
    "μg/kg": 1.0,
    "µg/kg": 1.0,
    "ug/kg": 1.0,
    "UG/KG": 1.0,
    "mg/kg": 1000.0,
    "MG/KG": 1000.0,
}


def normalize_unit(u: str) -> str:
    if not u:
        return ""
    u = u.replace("μ", "μ").replace("µ", "μ").strip()
    u = re.sub(r"\s+", "", u)
    if u.lower() == "ug/kg":
        return "μg/kg"
    if u.lower() == "mg/kg":
        return "mg/kg"
    if u.lower() == "cfu/g":
        return "CFU/g"
    return u


def convert_value(value: float | None, unit: str) -> tuple[float | None, str]:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return None, normalize_unit(unit)
    u = normalize_unit(unit)
    if u == "mg/kg":
        return float(value) * 1000.0, "μg/kg"
    if u in {"μg/kg", "ug/kg", "µg/kg"}:
        return float(value), "μg/kg"
    return float(value), u


def extract_num_unit(text: str) -> tuple[float | None, str, str]:
    s = normalize_text(text)
    if not s:
        return None, "", "空值"
    if re.search(r"未检出|不得检出|阴性", s):
        return 0.0, "", "未检出/不得检出"
    pat = re.compile(
        r"([<>≤≥=]*\s*-?\d+(?:\.\d+)?)\s*(μg/kg|µg/kg|ug/kg|mg/kg|g/100g|mg/g|CFU/g)?",
        re.IGNORECASE,
    )
    m = pat.search(s)
    if not m:
        return None, "", "无法提取数值"
    val = float(re.sub(r"[<>≤≥=\s]", "", m.group(1)))
    unit = normalize_unit(m.group(2) or "")
    conv, out_unit = convert_value(val, unit)
    return conv, out_unit, "已解析"


def parse_concentration(text: str) -> dict:
    s = normalize_text(text)
    result = {
        "初检浓度值": None,
        "复检浓度值": None,
        "最终采用浓度值": None,
        "浓度单位": "",
        "浓度清洗状态": "",
    }
    if not s:
        result["浓度清洗状态"] = "检测数值为空"
        return result
    if re.search(r"未检出|不得检出|阴性", s):
        result.update({"初检浓度值": 0.0, "最终采用浓度值": 0.0, "浓度清洗状态": "未检出/不得检出"})
        return result
    unit_pat = r"(μg/kg|µg/kg|ug/kg|mg/kg|g/100g|mg/g|CFU/g)"
    first = re.search(r"初检(?:结果)?[:：]?\s*[^0-9\-<>≤≥=]*([<>≤≥=]*\s*-?\d+(?:\.\d+)?)\s*" + unit_pat + "?", s, re.I)
    second = re.search(r"复检(?:结果)?[:：]?\s*[^0-9\-<>≤≥=]*([<>≤≥=]*\s*-?\d+(?:\.\d+)?)\s*" + unit_pat + "?", s, re.I)

    def parse_match(m):
        if not m:
            return None, ""
        val = float(re.sub(r"[<>≤≥=\s]", "", m.group(1)))
        unit = normalize_unit(m.group(2) or "")
        return convert_value(val, unit)

    fv, fu = parse_match(first)
    sv, su = parse_match(second)
    if fv is not None:
        result["初检浓度值"], result["浓度单位"] = fv, fu
    if sv is not None:
        result["复检浓度值"], result["浓度单位"] = sv, su or result["浓度单位"]
    if sv is not None:
        result["最终采用浓度值"] = sv
        result["浓度清洗状态"] = "初检复检均解析，最终采用复检值"
    elif fv is not None:
        result["最终采用浓度值"] = fv
        result["浓度清洗状态"] = "解析初检值，最终采用初检值"
    else:
        val, unit, status = extract_num_unit(s)
        result["初检浓度值"] = val
        result["最终采用浓度值"] = val
        result["浓度单位"] = unit
        result["浓度清洗状态"] = status
    return result


def parse_limit(text: str) -> tuple[float | None, str, str]:
    val, unit, status = extract_num_unit(text)
    return val, unit, status


def make_risk_summary(df: pd.DataFrame, group_cols: list[str]) -> pd.DataFrame:
    g = df.groupby(group_cols, dropna=False)
    out = g.agg(
        抽检总批次数=("序号", "count"),
        合格批次数=("是否合格", "sum"),
        不合格批次数=("是否不合格", "sum"),
        AFB1相关记录数=("是否AFB1相关", "sum"),
        AFB1相关不合格批次数=("AFB1相关不合格", "sum"),
        浓度可用记录数=("最终采用浓度值", lambda s: int(s.notna().sum())),
    ).reset_index()
    out["不合格率"] = out["不合格批次数"] / out["抽检总批次数"].replace(0, np.nan)
    out["AFB1相关不合格率"] = out["AFB1相关不合格批次数"] / out["AFB1相关记录数"].replace(0, np.nan)
    return out


def safe_markdown_table(df: pd.DataFrame, max_rows: int = 20) -> str:
    if df.empty:
        return "无记录。"
    small = df.head(max_rows).copy()
    small = small.astype(object).where(pd.notna(small), "")
    cols = [str(c) for c in small.columns]
    rows = []
    rows.append("| " + " | ".join(cols) + " |")
    rows.append("| " + " | ".join(["---"] * len(cols)) + " |")
    for _, row in small.iterrows():
        vals = []
        for c in small.columns:
            val = str(row[c]).replace("\n", " ").replace("|", "/")
            vals.append(val[:300])
        rows.append("| " + " | ".join(vals) + " |")
    return "\n".join(rows)


def plot_outputs(df: pd.DataFrame, count_panel: pd.DataFrame, conc: pd.DataFrame, quality: pd.DataFrame) -> list[str]:
    fig_paths = []

    def esc(x) -> str:
        return (
            str(x)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )

    def save_svg(name: str, title: str, body: str, width: int = 900, height: int = 520):
        path = ROOT / f"reports/figures/{name}.svg"
        svg = f"""<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">
<rect width="100%" height="100%" fill="#ffffff"/>
<text x="{width/2}" y="34" text-anchor="middle" font-family="Microsoft YaHei, SimHei, Arial" font-size="22" font-weight="700" fill="#222">{esc(title)}</text>
{body}
</svg>
"""
        write_text(path, svg)
        fig_paths.append(path.relative_to(ROOT).as_posix())

    def bar_chart(name: str, title: str, labels, values, color="#4C78A8", horizontal=False, y_label=""):
        labels = [str(x) if str(x) else "缺失" for x in labels]
        values = [float(v) if pd.notna(v) else 0.0 for v in values]
        width, height = 900, 520
        left, right, top, bottom = 150 if horizontal else 70, 40, 70, 90
        plot_w, plot_h = width - left - right, height - top - bottom
        maxv = max(values) if values else 1.0
        maxv = maxv if maxv > 0 else 1.0
        parts = [f'<line x1="{left}" y1="{top+plot_h}" x2="{left+plot_w}" y2="{top+plot_h}" stroke="#333"/>']
        if horizontal:
            n = max(len(labels), 1)
            step = plot_h / n
            bar_h = min(24, step * 0.65)
            for i, (lab, val) in enumerate(zip(labels, values)):
                y = top + i * step + (step - bar_h) / 2
                w = val / maxv * plot_w
                parts.append(f'<text x="{left-8}" y="{y+bar_h*0.7}" text-anchor="end" font-family="Microsoft YaHei, SimHei, Arial" font-size="13" fill="#333">{esc(lab[:16])}</text>')
                parts.append(f'<rect x="{left}" y="{y}" width="{w}" height="{bar_h}" fill="{color}"/>')
                parts.append(f'<text x="{left+w+5}" y="{y+bar_h*0.7}" font-family="Arial" font-size="12" fill="#333">{val:.2f}</text>')
        else:
            n = max(len(labels), 1)
            step = plot_w / n
            bar_w = min(44, step * 0.65)
            for i, (lab, val) in enumerate(zip(labels, values)):
                x = left + i * step + (step - bar_w) / 2
                h = val / maxv * plot_h
                y = top + plot_h - h
                parts.append(f'<rect x="{x}" y="{y}" width="{bar_w}" height="{h}" fill="{color}"/>')
                parts.append(f'<text transform="translate({x+bar_w/2},{top+plot_h+18}) rotate(35)" text-anchor="start" font-family="Microsoft YaHei, SimHei, Arial" font-size="12" fill="#333">{esc(lab[:12])}</text>')
                parts.append(f'<text x="{x+bar_w/2}" y="{y-5}" text-anchor="middle" font-family="Arial" font-size="11" fill="#333">{val:.0f}</text>')
        if y_label:
            parts.append(f'<text x="20" y="{top+plot_h/2}" transform="rotate(-90 20,{top+plot_h/2})" text-anchor="middle" font-family="Microsoft YaHei, SimHei, Arial" font-size="13" fill="#333">{esc(y_label)}</text>')
        save_svg(name, title, "\n".join(parts), width, height)

    def line_chart(name: str, title: str, labels, values, color="#F58518"):
        labels = [str(x) for x in labels]
        values = [float(v) if pd.notna(v) else 0.0 for v in values]
        width, height = 900, 520
        left, right, top, bottom = 70, 40, 70, 90
        plot_w, plot_h = width - left - right, height - top - bottom
        maxv = max(values) if values else 1.0
        maxv = maxv if maxv > 0 else 1.0
        n = max(len(values), 1)
        pts = []
        for i, val in enumerate(values):
            x = left + (plot_w * i / max(n - 1, 1))
            y = top + plot_h - val / maxv * plot_h
            pts.append((x, y, val))
        parts = [
            f'<line x1="{left}" y1="{top+plot_h}" x2="{left+plot_w}" y2="{top+plot_h}" stroke="#333"/>',
            f'<line x1="{left}" y1="{top}" x2="{left}" y2="{top+plot_h}" stroke="#333"/>',
            '<polyline points="' + " ".join(f"{x},{y}" for x, y, _ in pts) + f'" fill="none" stroke="{color}" stroke-width="3"/>',
        ]
        for i, (x, y, val) in enumerate(pts):
            parts.append(f'<circle cx="{x}" cy="{y}" r="5" fill="{color}"/>')
            parts.append(f'<text x="{x}" y="{y-10}" text-anchor="middle" font-family="Arial" font-size="12" fill="#333">{val:.2f}</text>')
            parts.append(f'<text x="{x}" y="{top+plot_h+24}" text-anchor="middle" font-family="Arial" font-size="12" fill="#333">{esc(labels[i])}</text>')
        save_svg(name, title, "\n".join(parts), width, height)

    def hist_chart(name: str, title: str, values, color="#9D755D"):
        vals = [float(v) for v in values if pd.notna(v) and float(v) >= 0]
        if not vals:
            save_svg(name, title, '<text x="450" y="260" text-anchor="middle" font-family="Microsoft YaHei, SimHei, Arial" font-size="18" fill="#666">无可用浓度数据</text>')
            return
        counts, edges = np.histogram(vals, bins=min(30, max(5, int(math.sqrt(len(vals))))))
        labels = [f"{edges[i]:.1f}" for i in range(len(counts))]
        bar_chart(name, title, labels, counts, color=color, horizontal=False, y_label="记录数")

    by_year = make_risk_summary(df, ["年份"]).sort_values("年份")
    bar_chart("peanut_year_sample_count", "年份—抽检批次", by_year["年份"].astype(str), by_year["抽检总批次数"], "#4C78A8", y_label="抽检总批次数")
    line_chart("peanut_year_noncompliance_rate", "年份—不合格率（%）", by_year["年份"].astype(str), by_year["不合格率"] * 100)

    region = make_risk_summary(df, ["生产省份_清洗"]).sort_values(["不合格率", "抽检总批次数"], ascending=False).head(20)
    bar_chart("peanut_region_risk_ranking", "省份风险排序（Top 20，不合格率%）", region["生产省份_清洗"].astype(str), region["不合格率"] * 100, "#E45756", horizontal=True)

    cat = make_risk_summary(df, ["产品二级类"]).sort_values(["不合格率", "抽检总批次数"], ascending=False).head(20)
    bar_chart("peanut_category_risk_ranking", "产品类别风险排序（Top 20，不合格率%）", cat["产品二级类"].astype(str), cat["不合格率"] * 100, "#72B7B2", horizontal=True)

    stage = make_risk_summary(df, ["供应链环节"]).sort_values("抽检总批次数", ascending=False)
    bar_chart("peanut_stage_risk_comparison", "供应链环节风险对比（不合格率%）", stage["供应链环节"].astype(str), stage["不合格率"] * 100, "#54A24B", y_label="不合格率（%）")

    afb = df[df["是否AFB1相关"]].copy()
    afb_stage = afb.groupby("供应链环节").size().sort_values(ascending=False)
    bar_chart("peanut_afb1_record_distribution", "AFB1 相关记录分布", afb_stage.index.astype(str), afb_stage.values, "#B279A2", y_label="记录数")

    usable = conc["最终采用浓度值"].dropna() if "最终采用浓度值" in conc else pd.Series(dtype=float)
    hist_chart("peanut_afb1_concentration_distribution", "AFB1 浓度分布", usable, "#9D755D")

    q = quality.copy()
    bar_chart("peanut_data_quality_issues", "数据质量问题", q["问题类型"], q["记录数"], "#FF9DA6", horizontal=True)

    return fig_paths


def update_project_state(outputs: list[str], conclusions: dict) -> None:
    today = datetime.now().strftime("%Y-%m-%d")
    handoff = f"""# Conversation Handoff

## 本轮任务目标

启动花生风险监管自动科研流程；仅使用 `data/01_raw/PEANUT2023-20241.xlsx` 与 `references/notes/物流与供应链管理前言-研究计划-肖宇航.docx`，不使用 `PEANUTwithProb0627.xlsx`。

## 使用的数据文件

- 原始数据：`data/01_raw/PEANUT2023-20241.xlsx`
- 原始数据保持未修改。

## 已读取的研究计划文档

- `references/notes/物流与供应链管理前言-研究计划-肖宇航.docx`
- 摘要：`references/processed_summaries/peanut_research_plan_summary.md`

## 已生成的核心数据文件

{chr(10).join('- `' + o + '`' for o in outputs if o.startswith('data/'))}

## 已生成的核心报告

{chr(10).join('- `' + o + '`' for o in outputs if o.startswith('reports/') or o.startswith('references/processed'))}

## 当前主要结论

- 原始数据可读取，sheet 结构正常，关键字段可识别。
- 已构建清洗分析主表、AFB1 浓度清洗表、省份—时间—供应链环节计数面板、EDA 汇总表与基础图表。
- AFB1 识别采用关键词精确匹配与可疑记录人工复核标记，不把所有生物毒素默认视为 AFB1。
- 浓度字段优先从 `检测数值` 提取，含初检/复检时最终采用复检值；法规限制从 `法规限制` 提取并标准化。

## 当前数据问题

- 仍存在字段缺失、浓度/限量无法解析、地区或环节为空、AFB1 可疑但不确定等记录，详见 `reports/tables/peanut_cleaning_issue_log.csv`。
- MOE/EDI、POMDP/DQN 仍缺少消费量、人口、体重、BMDL、抽检成本、预算、产能、召回损失和供应链响应参数。

## 后续建模条件

- MOE/EDI：{conclusions['moe']}
- Beta-Binomial：{conclusions['beta']}
- POMDP / belief-MDP：{conclusions['pomdp']}
- DQN：{conclusions['dqn']}

## 下一步建议

优先补充外部参数，并在当前计数面板上实现 Beta-Binomial 信念更新原型；随后把 AFB1 浓度分布与消费量/体重/BMDL 对齐，形成 MOE/EDI 风险特征。

## 新对话继续 Prompt

请继续花生 AFB1 风险监管流程。先读取 `project_state/conversation_handoff.md`、`project_state/next_step.md`、`reports/peanut_full_workflow_summary.md`、`data/04_feature/peanut_count_panel.csv` 和 `data/04_feature/peanut_concentration_clean_table.csv`。在不修改 `data/01_raw` 的前提下，优先补充/读取外部参数，构建 Beta-Binomial 信念更新原型，并判断 MOE/EDI 参数缺口。
"""
    write_text(ROOT / "project_state/conversation_handoff.md", handoff)
    write_text(
        ROOT / "project_state/current_focus.md",
        "# Current Focus\n\n当前焦点：基于 `PEANUT2023-20241.xlsx` 的花生/AFB1 风险监管数据基础已完成首轮构建。下一阶段应围绕 Beta-Binomial 信念更新、MOE/EDI 外部参数补充、POMDP/belief-MDP 环境变量设定展开。\n",
    )
    write_text(
        ROOT / "project_state/next_step.md",
        "# Next Step\n\n优先使用 `data/04_feature/peanut_count_panel.csv` 构建 Beta-Binomial 信念更新原型，并补充 MOE/EDI 所需消费量、人口、体重、BMDL、预算、成本和产能参数。\n",
    )
    with (ROOT / "project_state/changelog.md").open("a", encoding="utf-8") as f:
        f.write(f"\n## {today}\n\n- 完成 PEANUT2023-20241 花生/AFB1 风险监管首轮自动科研流程，生成 schema、清洗主表、计数面板、浓度清洗表、EDA、图表、可行性报告和交接文件。\n")
    with (ROOT / "project_state/decision_log.md").open("a", encoding="utf-8") as f:
        f.write(f"\n## {today}\n\n### 限定本轮数据源并构建花生/AFB1 数据基础\n\nRationale: 用户明确要求本轮只使用 `PEANUT2023-20241.xlsx` 和指定研究计划，不使用 `PEANUTwithProb0627.xlsx`。\n\nImpact: 所有派生数据与报告均基于指定原始数据；原始数据未修改。后续建模应从本轮计数面板和浓度清洗表继续。\n\n### AFB1 浓度最终值采用复检优先规则\n\nRationale: 用户要求对“初检结果/复检结果”分别提取，并构建最终采用浓度值。\n\nImpact: 当复检浓度可解析时使用复检值，否则使用初检或单一检测值；无法解析记录进入 issue log。\n")
    with (ROOT / "project_state/lessons_learned.md").open("a", encoding="utf-8") as f:
        f.write(f"\n## {today}\n\n- 花生 AFB1 任务中，`检测数值` 可能同时包含初检与复检，应复检优先并保留原始文本。\n- AFB1 识别不能用“生物毒素”泛化，需要 `黄曲霉`、`AFB1`、`B1/B₁` 等上下文关键词与人工复核标记。\n- 后续 POMDP/DQN 前必须先有省份—时间—环节计数面板，以及成本、预算、产能、消费量、人口、体重、BMDL 等外部参数。\n")
    roadmap = ROOT / "project_state/roadmap.yaml"
    if roadmap.exists():
        text = roadmap.read_text(encoding="utf-8", errors="replace")
        text = text.replace("status: pending\n    goals:\n      - Implement actual cleaning only after a dataset and research goal are provided.", "status: completed\n    goals:\n      - Implement actual cleaning only after a dataset and research goal are provided.")
        text = text.replace("status: pending\n    goals:\n      - Export figures, tables, and report drafts.", "status: completed\n    goals:\n      - Export figures, tables, and report drafts.")
        text += f"\n# Updated {today}: peanut_data_foundation_completed\n"
        write_text(roadmap, text)


def main() -> int:
    ensure_dirs()
    if not RAW_FILE.exists():
        raise FileNotFoundError("PEANUT2023-20241.xlsx 不存在，已停止。")
    try:
        with zipfile.ZipFile(RAW_FILE) as zf:
            if zf.testzip() is not None:
                raise RuntimeError("Excel zip 结构异常")
    except zipfile.BadZipFile as exc:
        raise RuntimeError("Excel 文件格式异常或损坏") from exc

    doc_summary_note = ""
    try:
        doc_path, paragraphs, tables = read_docx_full()
        extract_doc_summary(doc_path, paragraphs, tables)
    except Exception as exc:
        doc_summary_note = f"研究计划读取问题：{type(exc).__name__}: {exc}"
        write_text(SUMMARY_FILE, f"# 花生研究计划摘要\n\n{doc_summary_note}\n")

    try:
        xl = pd.ExcelFile(RAW_FILE)
        sheet_names = xl.sheet_names
        if not sheet_names:
            raise RuntimeError("Excel 没有可读取 sheet。")
        if "Sheet1" in sheet_names:
            sheet = "Sheet1"
        else:
            sheet = sheet_names[0]
        df_raw = pd.read_excel(RAW_FILE, sheet_name=sheet)
    except Exception as exc:
        raise RuntimeError(f"Excel 读取错误：{type(exc).__name__}: {exc}") from exc
    if df_raw.empty or df_raw.shape[1] == 0:
        raise RuntimeError("sheet 结构异常：数据为空或没有列。")

    raw_meta = {"rows": len(df_raw), "cols": len(df_raw.columns), "sheet": sheet, "sheets": sheet_names}
    inv = schema_inventory(df_raw, sheet, raw_meta)
    to_csv(inv, ROOT / "reports/tables/schema_inventory_PEANUT2023-20241.csv")

    roles = inv.groupby("候选角色")["字段名"].apply(lambda x: "、".join(map(str, x))).reset_index()
    schema_md = f"""# PEANUT2023-20241 Schema Inventory

## 文件格式识别

- 文件：`data/01_raw/PEANUT2023-20241.xlsx`
- 格式：Excel `.xlsx`
- sheet names：{', '.join(sheet_names)}
- 本轮读取 sheet：`{sheet}`
- 行数：{raw_meta['rows']}
- 列数：{raw_meta['cols']}

## 字段角色识别

{safe_markdown_table(roles, 30)}

## 字段清单与质量摘要

{safe_markdown_table(inv[['字段名','推断类型','缺失率','唯一值数量','候选角色','是否候选主键','样例值']], 60)}

## 关键字段判断

- 时间字段：`生产日期`、`通报时间`
- 地区字段：`生产省份`、`生产企业地址`、`抽样单位地址`、`通报单位`
- 产品字段：`产品分类`、`产品名称`
- 抽检环节字段：`抽检环节`、`是否属于网抽`
- 判定结果字段：`判定结果`
- 不合格项目字段：`不合格项目分类`、`不合格项目`、`不合格规范列`
- 检测数值字段：`检测数值`
- 法规限制字段：`法规限制`
- AFB1 相关字段：主要从 `不合格项目`、`不合格规范列`、`检验项目`、`检测数值`、`备注`、`标准` 中识别。
"""
    write_text(ROOT / "reports/schema_inventory_PEANUT2023-20241.md", schema_md)

    required = ["通报时间", "生产省份", "产品名称", "判定结果", "检测数值", "法规限制", "抽检环节"]
    missing_required = [c for c in required if c not in df_raw.columns]
    if missing_required:
        raise RuntimeError("关键字段完全无法识别，缺少：" + "、".join(missing_required))

    empty_cols = [c for c in df_raw.columns if df_raw[c].isna().all()]
    unnamed_cols = [c for c in df_raw.columns if str(c).startswith("Unnamed")]
    drop_cols = sorted(set(empty_cols + unnamed_cols), key=str)
    df = df_raw.drop(columns=drop_cols, errors="ignore").copy()
    df["原始行号"] = np.arange(2, len(df) + 2)
    df["是否重复记录"] = df.duplicated(subset=[c for c in df.columns if c != "原始行号"], keep="first")
    duplicate_count = int(df["是否重复记录"].sum())
    df = df[~df["是否重复记录"]].copy()

    time_col = "通报时间"
    dates = pd.to_datetime(df[time_col], errors="coerce")
    fallback_dates = pd.to_datetime(df.get("生产日期"), errors="coerce")
    dates = dates.fillna(fallback_dates)
    df["年份"] = dates.dt.year.astype("Int64")
    df["月份"] = dates.dt.month.astype("Int64")
    df["年月"] = dates.dt.strftime("%Y-%m")

    df["生产省份_清洗"] = df["生产省份"].map(clean_province)
    df["通报单位_清洗"] = df["通报单位"].map(clean_province) if "通报单位" in df else ""
    addr_col = "抽样单位地址" if "抽样单位地址" in df else "生产企业地址"
    df["抽样地区_候选"] = df[addr_col].map(clean_province) if addr_col in df else ""

    df["原始产品分类"] = df["产品分类"].map(normalize_text) if "产品分类" in df else ""
    df["原始产品名称"] = df["产品名称"].map(normalize_text) if "产品名称" in df else ""
    product_classes = df.apply(lambda r: classify_product(r.get("产品分类"), r.get("产品名称")), axis=1)
    df["产品一级类"] = [x[0] for x in product_classes]
    df["产品二级类"] = [x[1] for x in product_classes]
    df["产品细分类"] = [x[2] for x in product_classes]

    stages = df["抽检环节"].map(clean_stage)
    df["抽检环节_清洗"] = [x[0] for x in stages]
    df["供应链环节"] = [x[1] for x in stages]
    if "是否属于网抽" in df:
        mask_web = df["是否属于网抽"].map(normalize_text).str.contains("是|网", regex=True, na=False)
        df.loc[mask_web, "供应链环节"] = "流通"

    for src, dest in [
        ("生产企业名称", "生产企业_清洗"),
        ("生产企业地址", "生产企业地址_清洗"),
        ("抽样单位名称", "抽样单位_清洗"),
        ("检验机构", "检验机构_清洗"),
    ]:
        df[dest] = df[src].map(normalize_text) if src in df else ""

    result = df["判定结果"].map(normalize_text)
    df["是否不合格"] = result.str.contains("不合格|问题|阳性|超标", regex=True, na=False)
    df["是否合格"] = result.str.contains("合格", regex=True, na=False) & ~df["是否不合格"]
    df["不合格项目_标准化"] = df["不合格项目"].map(normalize_text) if "不合格项目" in df else ""
    afb_info = df.apply(afb1_detect, axis=1)
    df["是否AFB1相关"] = [x[0] for x in afb_info]
    df["AFB1识别依据"] = [x[1] for x in afb_info]
    afb_suspicious = pd.Series([x[2] for x in afb_info], index=df.index)
    df["污染物名称_标准化"] = np.where(df["是否AFB1相关"], "黄曲霉毒素B1（AFB1）", df["不合格项目_标准化"])
    df["污染物类别_标准化"] = np.where(df["是否AFB1相关"], "真菌毒素/生物毒素", np.where(df["不合格项目_标准化"].ne(""), "其他不合格项目", "未识别"))

    conc_rows = df["检测数值"].map(parse_concentration)
    for col in ["初检浓度值", "复检浓度值", "最终采用浓度值", "浓度单位", "浓度清洗状态"]:
        df[col] = [r[col] for r in conc_rows]
    df["原始检测数值"] = df["检测数值"].map(normalize_text)
    lim_rows = df["法规限制"].map(parse_limit)
    df["法规限量_数值"] = [x[0] for x in lim_rows]
    df["法规限量_单位"] = [x[1] for x in lim_rows]
    df["法规限量解析状态"] = [x[2] for x in lim_rows]
    df["是否检出"] = df["最终采用浓度值"].fillna(np.nan).gt(0)
    df["是否超标"] = False
    both = df["最终采用浓度值"].notna() & df["法规限量_数值"].notna()
    df.loc[both, "是否超标"] = df.loc[both, "最终采用浓度值"] > df.loc[both, "法规限量_数值"]
    df.loc[~both & df["是否不合格"] & df["是否AFB1相关"], "是否超标"] = True
    df["超标倍数"] = np.where(both & (df["法规限量_数值"] > 0), df["最终采用浓度值"] / df["法规限量_数值"], np.nan)
    df["AFB1相关不合格"] = df["是否AFB1相关"] & df["是否不合格"]

    review_reasons = []
    for _, r in df.iterrows():
        reasons = []
        if not r["生产省份_清洗"]:
            reasons.append("生产省份缺失或无法规范")
        if pd.isna(r["年份"]):
            reasons.append("时间无法解析")
        if r["供应链环节"] == "其他":
            reasons.append("抽检环节无法明确归类")
        if afb_suspicious.loc[r.name]:
            reasons.append("AFB1识别可疑需复核")
        if r["是否AFB1相关"] and pd.isna(r["最终采用浓度值"]):
            reasons.append("AFB1相关但浓度无法解析")
        if r["是否AFB1相关"] and pd.isna(r["法规限量_数值"]):
            reasons.append("AFB1相关但法规限量无法解析")
        review_reasons.append("；".join(reasons))
    df["复核原因"] = review_reasons
    df["是否建议人工复核"] = df["复核原因"].ne("")

    primary_cols = list(df.columns)
    df.to_excel(ROOT / "data/03_primary/peanut_cleaned_analysis_ready.xlsx", index=False)
    to_csv(df, ROOT / "data/03_primary/peanut_cleaned_analysis_ready.csv")

    var_dict = pd.DataFrame(
        [
            {"字段名": c, "字段类型": str(df[c].dtype), "含义": classify_role(c), "缺失率": round(float(df[c].isna().mean()), 6)}
            for c in primary_cols
        ]
    )
    to_csv(var_dict, ROOT / "reports/tables/peanut_variable_dictionary.csv")
    label_dict = pd.DataFrame(
        [
            {"标签字段": "是否合格", "取值": "True/False", "含义": "判定结果为合格且非不合格"},
            {"标签字段": "是否不合格", "取值": "True/False", "含义": "判定结果含不合格/问题/阳性/超标"},
            {"标签字段": "是否检出", "取值": "True/False", "含义": "最终采用浓度值大于0"},
            {"标签字段": "是否超标", "取值": "True/False", "含义": "浓度超过法规限量，或AFB1相关且判定不合格但限量无法解析"},
            {"标签字段": "是否AFB1相关", "取值": "True/False", "含义": "关键字段包含黄曲霉/AFB1/B1/B₁等合理变体"},
            {"标签字段": "AFB1相关不合格", "取值": "True/False", "含义": "同时为AFB1相关与不合格"},
        ]
    )
    to_csv(label_dict, ROOT / "reports/tables/peanut_label_dictionary.csv")

    issue_records = []
    issue_defs = [
        ("删除完全空列或Unnamed列", duplicate_count * 0, len(drop_cols), "列级问题", "已从分析主表隔离"),
        ("重复记录", duplicate_count, duplicate_count, "行级问题", "已在主表去重，保留首条"),
        ("时间无法解析", int(df["年份"].isna().sum()), int(df["年份"].isna().sum()), "字段问题", "建议回查通报时间/生产日期"),
        ("生产省份缺失或无法规范", int(df["生产省份_清洗"].eq("").sum()), int(df["生产省份_清洗"].eq("").sum()), "字段问题", "建议人工补全省份"),
        ("抽检环节无法明确归类", int(df["供应链环节"].eq("其他").sum()), int(df["供应链环节"].eq("其他").sum()), "字段问题", "后续可细化环节词典"),
        ("AFB1相关但浓度无法解析", int((df["是否AFB1相关"] & df["最终采用浓度值"].isna()).sum()), int((df["是否AFB1相关"] & df["最终采用浓度值"].isna()).sum()), "浓度问题", "保留原始值并复核"),
        ("AFB1相关但法规限量无法解析", int((df["是否AFB1相关"] & df["法规限量_数值"].isna()).sum()), int((df["是否AFB1相关"] & df["法规限量_数值"].isna()).sum()), "限量问题", "保留原始值并复核"),
        ("建议人工复核", int(df["是否建议人工复核"].sum()), int(df["是否建议人工复核"].sum()), "复核标记", "详见复核原因"),
    ]
    for name, n, rows, typ, action in issue_defs:
        issue_records.append({"问题类型": name, "记录数": rows, "问题层级": typ, "处理方式": action})
    issue_log = pd.DataFrame(issue_records)
    to_csv(issue_log, ROOT / "reports/tables/peanut_cleaning_issue_log.csv")

    grouping_time = ["生产省份_清洗", "年份", "月份", "年月", "供应链环节"]
    count_panel = make_risk_summary(df, grouping_time).rename(columns={"生产省份_清洗": "省份"})
    count_panel["数据完整性标记"] = np.where(
        count_panel[["省份", "年份", "月份", "年月", "供应链环节"]].isna().any(axis=1) | count_panel["省份"].eq(""),
        "关键索引缺失",
        "基本完整",
    )
    to_csv(count_panel, ROOT / "data/04_feature/peanut_count_panel.csv")
    count_panel.to_excel(ROOT / "data/04_feature/peanut_count_panel.xlsx", index=False)

    conc_cols = [
        "序号", "生产省份_清洗", "年份", "月份", "年月", "供应链环节", "产品二级类", "产品名称",
        "污染物名称_标准化", "是否AFB1相关", "原始检测数值", "初检浓度值", "复检浓度值", "最终采用浓度值",
        "浓度单位", "法规限量_数值", "法规限量_单位", "是否超标", "超标倍数", "浓度清洗状态",
        "是否建议人工复核", "复核原因",
    ]
    conc = df[df["是否AFB1相关"]].copy()
    conc_out = conc[[c for c in conc_cols if c in conc.columns]].rename(
        columns={"生产省份_清洗": "省份", "产品二级类": "产品分类"}
    )
    to_csv(conc_out, ROOT / "data/04_feature/peanut_concentration_clean_table.csv")

    if not conc_out.empty:
        def log_mu(s):
            vals = pd.to_numeric(s, errors="coerce")
            vals = vals[vals > 0]
            return float(np.log(vals).mean()) if len(vals) >= 2 else np.nan

        def log_sigma(s):
            vals = pd.to_numeric(s, errors="coerce")
            vals = vals[vals > 0]
            return float(np.log(vals).std(ddof=1)) if len(vals) >= 2 else np.nan

        dist = conc_out.groupby(["省份", "年份", "月份", "年月", "供应链环节"], dropna=False).agg(
            样本量=("最终采用浓度值", "count"),
            均值=("最终采用浓度值", "mean"),
            中位数=("最终采用浓度值", "median"),
            P75=("最终采用浓度值", lambda s: s.quantile(0.75)),
            P90=("最终采用浓度值", lambda s: s.quantile(0.90)),
            P95=("最终采用浓度值", lambda s: s.quantile(0.95)),
            最大值=("最终采用浓度值", "max"),
            lognormal_mu=("最终采用浓度值", log_mu),
            lognormal_sigma=("最终采用浓度值", log_sigma),
        ).reset_index()
        dist["是否适合MonteCarlo模拟"] = np.where(dist["样本量"] >= 10, "较适合", np.where(dist["样本量"] >= 3, "可探索但样本偏少", "不适合/样本不足"))
    else:
        dist = pd.DataFrame(columns=["省份", "年份", "月份", "年月", "供应链环节", "样本量", "均值", "中位数", "P75", "P90", "P95", "最大值", "lognormal_mu", "lognormal_sigma", "是否适合MonteCarlo模拟"])
    to_csv(dist, ROOT / "data/04_feature/peanut_concentration_distribution_summary.csv")

    by_year = make_risk_summary(df, ["年份"])
    by_region = make_risk_summary(df, ["生产省份_清洗"]).rename(columns={"生产省份_清洗": "省份"})
    by_stage = make_risk_summary(df, ["供应链环节"])
    by_category = make_risk_summary(df, ["产品二级类"]).rename(columns={"产品二级类": "产品类别"})
    to_csv(by_year, ROOT / "reports/tables/peanut_risk_summary_by_year.csv")
    to_csv(by_region, ROOT / "reports/tables/peanut_risk_summary_by_region.csv")
    to_csv(by_stage, ROOT / "reports/tables/peanut_risk_summary_by_stage.csv")
    to_csv(by_category, ROOT / "reports/tables/peanut_risk_summary_by_category.csv")
    quality = issue_log[["问题类型", "记录数", "处理方式"]].copy()
    to_csv(quality, ROOT / "reports/tables/peanut_data_quality_summary.csv")

    fig_paths = plot_outputs(df, count_panel, conc_out, quality)

    total = len(df)
    afb_n = int(df["是否AFB1相关"].sum())
    conc_avail = int(conc_out["最终采用浓度值"].notna().sum()) if "最终采用浓度值" in conc_out else 0
    month_ok = int(df["年月"].notna().sum())
    panel_cells = len(count_panel)

    cleaning_report = f"""# 花生抽检数据清洗报告

## 输入与范围

- 仅使用 `data/01_raw/PEANUT2023-20241.xlsx`，未读取或使用 `PEANUTwithProb0627.xlsx`。
- 原始数据 sheet：`{sheet}`；原始行数 {raw_meta['rows']}，原始列数 {raw_meta['cols']}。
- 隔离完全空列或 `Unnamed` 列：{len(drop_cols)} 个（{', '.join(map(str, drop_cols)) if drop_cols else '无'}）。
- exact duplicate 记录：{duplicate_count} 条，分析主表保留首条。

## 关键清洗逻辑

- 时间：优先解析 `通报时间`，失败时使用 `生产日期`，生成 `年份`、`月份`、`年月`。
- 地区：从 `生产省份`、`通报单位`、`抽样单位地址` 提取省级口径。
- 产品：保留 `原始产品分类` 与 `原始产品名称`，构建 `产品一级类`、`产品二级类`、`产品细分类`。
- 环节：将 `抽检环节` 和网抽信息映射为 `源头/初加工`、`生产`、`流通`、`餐饮`、`其他`。
- AFB1：从不合格项目、检验项目、检测数值、备注和标准中识别 `黄曲霉`、`AFB1`、`B1/B₁` 等写法；可疑但不确定的记录进入人工复核。
- 浓度：保留 `原始检测数值`；解析初检、复检和单一数值；当复检浓度可解析时最终采用复检值，否则采用初检/单一检测值。
- 限量：解析 `法规限制` 中的 `≤20μg/kg`、`20µg/kg`、`≤20` 等数值和单位。

## 标签构建结果

- `是否合格`、`是否不合格`、`是否检出`、`是否超标`
- `不合格项目_标准化`、`污染物名称_标准化`、`污染物类别_标准化`
- `是否AFB1相关`、`AFB1识别依据`、`AFB1相关不合格`
- `是否建议人工复核`、`复核原因`

## 规模摘要

- 清洗后主表记录数：{total}
- AFB1 相关记录数：{afb_n}
- AFB1 浓度可用记录数：{conc_avail}
- 可解析年月记录数：{month_ok}

## 数据质量问题

{safe_markdown_table(issue_log, 20)}
"""
    write_text(ROOT / "reports/peanut_cleaning_report.md", cleaning_report)

    panel_report = f"""# 花生风险计数面板报告

## 面板索引

- 优先索引：`生产省份_清洗` × `年月` × `供应链环节`
- 输出中保留：`省份`、`年份`、`月份`、`年月`、`供应链环节`

## 面板规模

- 面板单元数：{panel_cells}
- 抽检总批次数合计：{int(count_panel['抽检总批次数'].sum())}
- AFB1 相关记录数合计：{int(count_panel['AFB1相关记录数'].sum())}
- 浓度可用记录数合计：{int(count_panel['浓度可用记录数'].sum())}

## 面板字段

包括 `抽检总批次数`、`合格批次数`、`不合格批次数`、`不合格率`、`AFB1相关记录数`、`AFB1相关不合格批次数`、`AFB1相关不合格率`、`浓度可用记录数`、`数据完整性标记`。

## 样例

{safe_markdown_table(count_panel.head(20), 20)}
"""
    write_text(ROOT / "reports/peanut_count_panel_report.md", panel_report)

    conc_report = f"""# AFB1 浓度清洗报告

## 记录范围

- 仅包含 `是否AFB1相关=True` 的记录。
- AFB1 相关记录数：{afb_n}
- 浓度可用记录数：{conc_avail}

## 浓度采用逻辑

- 含 `初检结果` 与 `复检结果` 的记录分别提取初检值和复检值。
- 若复检浓度可解析，`最终采用浓度值` 使用复检值。
- 若无复检但初检或单一检测值可解析，使用初检/单一值。
- 无法解析记录不删除，保留原始值并写入 issue log。

## 分布摘要样例

{safe_markdown_table(dist.head(20), 20)}
"""
    write_text(ROOT / "reports/peanut_concentration_cleaning_report.md", conc_report)

    eda_report = f"""# 花生风险 EDA 报告

## 年份汇总

{safe_markdown_table(by_year.sort_values('年份'), 20)}

## 省份风险 Top 20

{safe_markdown_table(by_region.sort_values(['不合格率','抽检总批次数'], ascending=False).head(20), 20)}

## 供应链环节

{safe_markdown_table(by_stage.sort_values('抽检总批次数', ascending=False), 20)}

## 产品类别 Top 20

{safe_markdown_table(by_category.sort_values(['不合格率','抽检总批次数'], ascending=False).head(20), 20)}

## AFB1 与浓度可用性

- AFB1 相关记录数：{afb_n}
- AFB1 相关不合格批次数：{int(df['AFB1相关不合格'].sum())}
- AFB1 浓度可用记录数：{conc_avail}

## 数据质量

{safe_markdown_table(quality, 20)}
"""
    write_text(ROOT / "reports/peanut_eda_report.md", eda_report)

    moe = "部分具备。已有 AFB1 浓度清洗表和限量/超标信息，但缺少消费量、体重、人口权重和 BMDL 等外部参数。"
    beta = "基本具备。已形成省份—年月—供应链环节抽检总数与 AFB1/不合格计数，可用于 Beta-Binomial 后验更新；遗忘因子需外部设定。"
    pomdp = "部分具备。已有观测计数、时间、地区和环节，但仍需定义状态转移、成本、预算、产能、动作档位和环境响应函数。"
    dqn = "暂不具备直接训练条件。缺少仿真环境、奖励参数、动作约束参数、预算/产能、成本/召回损失和供应链响应参数。"
    feasibility = f"""# 建模与优化可行性判断

## 监督学习目标

- `是否不合格`：具备，可作为二分类目标。
- `是否AFB1相关`：具备，但更适合作为污染物识别/筛选标签。
- `是否超标`：部分具备；当浓度和法规限量均可解析时较可靠，无法解析记录需复核。
- `AFB1相关不合格`：具备，可用于 AFB1 风险识别。
- 风险等级：可进一步由超标倍数、MOE/EDI 或分位数构建，本轮未强行分级。

## Random Forest / XGBoost 条件

可在清洗主表上尝试，但应先处理类别编码、时间切分、类别不平衡与泄漏风险。不建议把 `判定结果` 或直接由判定派生的字段作为特征预测自身。

## 时间序列或面板建模条件

已具备初步省份—年月—供应链环节面板；若月份缺失较多，可退化到年份面板。

## Beta-Binomial

{beta}

## MOE/EDI

{moe}

## POMDP / belief-MDP

{pomdp}

## DQN / 强化学习

{dqn}

## 当前缺少外部参数

- 消费量
- 人口
- 体重
- BMDL 或毒理学基准
- 抽检成本
- 预算
- 产能上限
- 处置/召回损失参数
- 供应链响应参数
- 遗忘因子设定

## 下一阶段建议

1. 基于 `peanut_count_panel.csv` 实现 Beta-Binomial 信念更新原型。
2. 补充消费量、人口、体重与 BMDL，计算 EDI/MOE。
3. 设计 POMDP/belief-MDP 的状态、动作、观测、转移和奖励，并先做小规模仿真。
4. 在约束与奖励参数明确后，再考虑 DQN/Double DQN。
"""
    write_text(ROOT / "reports/peanut_modeling_optimization_feasibility.md", feasibility)

    vis_report = f"""# 花生风险基础可视化报告

## 已生成图表

{chr(10).join('- `' + p + '`' for p in fig_paths)}

## 说明

图表采用中文标题和标签，服务于首轮 EDA 与数据质量诊断；后续论文级图表可在补充建模结果后进一步统一配色、排序和注释。
"""
    write_text(ROOT / "reports/peanut_visualization_report.md", vis_report)

    full_summary = f"""# 花生 AFB1 风险监管首轮自动科研流程总结

## 数据来源与结构

- 原始数据：`data/01_raw/PEANUT2023-20241.xlsx`
- sheet：`{sheet}`
- 原始规模：{raw_meta['rows']} 行 × {raw_meta['cols']} 列
- 清洗后主表：{total} 行

## 研究计划对本轮数据结构的要求

研究计划要求先构建省份—时间—环节计数面板，用于 Beta-Binomial、POMDP/belief-MDP 和 DQN；同时构建 AFB1 浓度清洗表，用于 MOE/EDI、浓度分布参数和 Monte Carlo 风险模拟。

## 数据清洗过程

{cleaning_report}

## 计数面板构建结果

{panel_report}

## 浓度清洗表与分布摘要

{conc_report}

## EDA 结果

{eda_report}

## 建模与优化可行性判断

{feasibility}

## 可视化输出

{vis_report}

## 当前局限

- 原始字段含义仍需结合监管通报口径人工确认。
- 部分浓度和限量文本无法完全自动解析。
- MOE/EDI 与 DQN 所需外部参数尚未补齐。

## 下一步建议

使用本轮生成的计数面板实现 Beta-Binomial 信念更新原型，并补充 MOE/EDI 外部参数。
"""
    write_text(ROOT / "reports/peanut_full_workflow_summary.md", full_summary)

    outputs = [
        "references/processed_summaries/peanut_research_plan_summary.md",
        "reports/schema_inventory_PEANUT2023-20241.md",
        "reports/tables/schema_inventory_PEANUT2023-20241.csv",
        "data/03_primary/peanut_cleaned_analysis_ready.xlsx",
        "data/03_primary/peanut_cleaned_analysis_ready.csv",
        "reports/peanut_cleaning_report.md",
        "reports/tables/peanut_variable_dictionary.csv",
        "reports/tables/peanut_label_dictionary.csv",
        "reports/tables/peanut_cleaning_issue_log.csv",
        "data/04_feature/peanut_count_panel.csv",
        "data/04_feature/peanut_count_panel.xlsx",
        "reports/peanut_count_panel_report.md",
        "data/04_feature/peanut_concentration_clean_table.csv",
        "data/04_feature/peanut_concentration_distribution_summary.csv",
        "reports/peanut_concentration_cleaning_report.md",
        "reports/peanut_eda_report.md",
        "reports/tables/peanut_risk_summary_by_year.csv",
        "reports/tables/peanut_risk_summary_by_region.csv",
        "reports/tables/peanut_risk_summary_by_stage.csv",
        "reports/tables/peanut_risk_summary_by_category.csv",
        "reports/tables/peanut_data_quality_summary.csv",
        "reports/peanut_modeling_optimization_feasibility.md",
        "reports/peanut_visualization_report.md",
        "reports/peanut_full_workflow_summary.md",
        *fig_paths,
    ]
    update_project_state(outputs, {"moe": moe, "beta": beta, "pomdp": pomdp, "dqn": dqn})

    run_summary = {
        "rows_cleaned": total,
        "afb1_records": afb_n,
        "concentration_available": conc_avail,
        "panel_cells": panel_cells,
        "figures": len(fig_paths),
        "doc_issue": doc_summary_note,
        "outputs": outputs,
    }
    write_text(ROOT / "reports/peanut_workflow_run_summary.json", pd.Series(run_summary).to_json(force_ascii=False, indent=2))
    print(pd.Series(run_summary).to_json(force_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        err = f"""# PEANUT workflow failed

错误类型：{type(exc).__name__}

错误位置：`scripts/run_peanut_risk_workflow.py`

错误信息：{exc}

已尝试的修复方式：脚本执行前已检查原始 Excel 存在性、xlsx zip 结构、sheet 可读性与 Word 文档目录扫描。

需要用户提供：请根据错误信息补充缺失文件、修复损坏 Excel/Word，或说明关键字段替代口径。
"""
        ensure_dirs()
        write_text(ROOT / "reports/peanut_workflow_error.md", err)
        print(err, file=sys.stderr)
        raise
