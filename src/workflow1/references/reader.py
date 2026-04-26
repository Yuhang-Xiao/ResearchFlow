"""Lightweight reference document text extraction."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


SUPPORTED_TEXT_EXTENSIONS = {".md", ".txt"}
SUPPORTED_TABLE_EXTENSIONS = {".csv", ".xlsx"}
SUPPORTED_DOCUMENT_EXTENSIONS = {".docx", ".pdf"}
SUPPORTED_EXTENSIONS = SUPPORTED_TEXT_EXTENSIONS | SUPPORTED_TABLE_EXTENSIONS | SUPPORTED_DOCUMENT_EXTENSIONS


@dataclass(frozen=True)
class ReferenceDocument:
    """Reference document metadata."""

    path: Path
    suffix: str


@dataclass(frozen=True)
class ReferenceReadResult:
    """Result of a safe reference read attempt."""

    document: ReferenceDocument
    status: str
    text: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: tuple[str, ...] = ()


def detect_reference(path: str | Path) -> ReferenceDocument:
    """Return basic reference metadata."""

    ref_path = Path(path)
    return ReferenceDocument(path=ref_path, suffix=ref_path.suffix.lower())


def read_reference(path: str | Path, max_chars: int = 20000) -> ReferenceReadResult:
    """Extract lightweight plain text from a supported reference file."""

    document = detect_reference(path)
    if not document.path.exists():
        return ReferenceReadResult(
            document=document,
            status="missing",
            warnings=(f"Reference file not found: {document.path}",),
        )
    if document.suffix not in SUPPORTED_EXTENSIONS:
        return ReferenceReadResult(
            document=document,
            status="unsupported",
            warnings=(f"Unsupported reference type: {document.suffix}",),
        )

    try:
        if document.suffix in SUPPORTED_TEXT_EXTENSIONS:
            text = document.path.read_text(encoding="utf-8", errors="replace")
            return _trim_result(document, "ok", text, max_chars)
        if document.suffix == ".docx":
            return _read_docx(document, max_chars)
        if document.suffix == ".pdf":
            return _read_pdf(document, max_chars)
        if document.suffix in SUPPORTED_TABLE_EXTENSIONS:
            return ReferenceReadResult(
                document=document,
                status="metadata_only",
                metadata={"message": "Tabular reference detected; inspect sheets/columns in a task-specific step."},
                warnings=("Tabular references are not fully extracted by the lightweight reader.",),
            )
    except Exception as exc:  # pragma: no cover - defensive boundary for damaged user files.
        return ReferenceReadResult(document=document, status="error", warnings=(str(exc),))

    return ReferenceReadResult(document=document, status="unsupported")


def _read_docx(document: ReferenceDocument, max_chars: int) -> ReferenceReadResult:
    try:
        from docx import Document
    except ImportError:
        return ReferenceReadResult(
            document=document,
            status="dependency_missing",
            warnings=("python-docx is required to read .docx references.",),
        )

    doc = Document(str(document.path))
    chunks: list[str] = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return _trim_result(document, "ok", "\n".join(chunks), max_chars)


def _read_pdf(document: ReferenceDocument, max_chars: int) -> ReferenceReadResult:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ReferenceReadResult(
            document=document,
            status="dependency_missing",
            warnings=("pypdf is required to read text-based .pdf references.",),
        )

    reader = PdfReader(str(document.path))
    chunks: list[str] = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    text = "\n".join(chunk for chunk in chunks if chunk.strip())
    warnings: tuple[str, ...] = ()
    if not text.strip():
        warnings = ("No text extracted. The PDF may be scanned or image-only; OCR is not enabled by default.",)
        return ReferenceReadResult(document=document, status="no_text_extracted", warnings=warnings)
    return _trim_result(document, "ok", text, max_chars)


def _trim_result(
    document: ReferenceDocument,
    status: str,
    text: str,
    max_chars: int,
    warnings: tuple[str, ...] = (),
) -> ReferenceReadResult:
    trimmed = text[:max_chars]
    if len(text) > max_chars:
        warnings = warnings + (f"Text was truncated to {max_chars} characters.",)
    return ReferenceReadResult(
        document=document,
        status=status,
        text=trimmed,
        metadata={"characters": len(trimmed)},
        warnings=warnings,
    )

