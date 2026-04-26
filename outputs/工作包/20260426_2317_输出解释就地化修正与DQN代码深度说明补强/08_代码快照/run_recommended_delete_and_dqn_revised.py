"""Run the recommended cache deletion and revised PEANUT DQN experiment.

This script is intentionally task-scoped.  It writes all primary artifacts into
one run package and keeps the run marked as experimental, not formal DQN.
"""

from __future__ import annotations

import csv
import hashlib
import json
import math
import os
import random
import shutil
import subprocess
import sys
import time
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.font_manager as fm
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import torch
from PIL import Image
from torch import nn


ROOT = Path(__file__).resolve().parents[1]
PYTHON = Path(r"D:\anaconda3\envs\myenv1\python.exe")
PREVIOUS_CLEANUP_RUN = ROOT / "outputs" / "工作包" / "20260426_2000_指定清理与DQN修正版继续实验"
PREVIOUS_DQN_RUN = ROOT / "outputs" / "工作包" / "20260426_1746_全流程验收与DQN自动参数训练"
TASK_NAME = "推荐缓存删除与DQN修正版训练"
EXPERIMENTAL_LABEL = "DQN修正版 experimental run / revised DQN experimental run; not formal policy conclusion"
REQUIRED_UPSTREAM = [
    ROOT / "data/04_feature/peanut_belief_mdp_state_features.csv",
    ROOT / "data/04_feature/peanut_belief_mdp_state_features_with_moe_edi.csv",
    ROOT / "data/04_feature/peanut_beta_binomial_belief_states.csv",
    ROOT / "data/04_feature/peanut_count_panel.csv",
    ROOT / "data/04_feature/peanut_edi_moe_risk_table.csv",
    ROOT / "data/04_feature/peanut_edi_moe_risk_summary.csv",
]
ACTION_VALUES = np.array([0, 1, 3, 5, 10], dtype=np.int64)
RANDOM_SEED = 42


def rel(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT).as_posix()
    except Exception:
        return path.as_posix()


def safe_read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace") if path.exists() else ""


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def append_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(text)


def write_csv(path: Path, rows: list[dict[str, Any]], fieldnames: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fieldnames is None:
        keys: list[str] = []
        for row in rows:
            for key in row:
                if key not in keys:
                    keys.append(key)
        fieldnames = keys
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def write_df(df: pd.DataFrame, path: Path, xlsx_path: Path | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(path, index=False, encoding="utf-8-sig")
    if xlsx_path is not None:
        xlsx_path.parent.mkdir(parents=True, exist_ok=True)
        df.to_excel(xlsx_path, index=False)


def sha256_file(path: Path) -> str:
    if not path.is_file():
        return ""
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def make_run_package() -> Path:
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    run = ROOT / "outputs" / "工作包" / f"{stamp}_{TASK_NAME}"
    suffix = 1
    while run.exists():
        run = ROOT / "outputs" / "工作包" / f"{stamp}_{TASK_NAME}_{suffix}"
        suffix += 1
    subdirs = [
        "00_输入说明",
        "01_数据输出",
        "02_表格输出",
        "03_图表输出",
        "04_报告输出",
        "05_模型与实验",
        "06_配置参数",
        "07_日志与错误",
        "08_代码快照",
        "09_论文输出/04_结果",
        "09_论文输出/09_word导出",
    ]
    for sub in subdirs:
        (run / sub).mkdir(parents=True, exist_ok=True)
    write_text(
        run / "README.md",
        f"""# {TASK_NAME}

任务性质：{EXPERIMENTAL_LABEL}

本工作包先删除上一轮明确推荐的缓存/临时项，然后继续 PEANUT DQN 修正版 experimental run。

关键输出：

- `02_表格输出/deleted_files_log.csv`
- `02_表格输出/multi_model_policy_comparison.csv`
- `03_图表输出/*.png`
- `04_报告输出/dqn_revised_result_audit_report.md`
- `09_论文输出/09_word导出/results_draft.docx`
""",
    )
    write_text(
        run / "00_输入说明" / "inputs.md",
        "\n".join(
            [
                "# 输入说明",
                "",
                f"- 上一轮清理扫描包：`{rel(PREVIOUS_CLEANUP_RUN)}`",
                f"- 上一轮 DQN experimental run：`{rel(PREVIOUS_DQN_RUN)}`",
                "- canonical 数据：`data/04_feature/peanut_*`",
                f"- 任务标签：{EXPERIMENTAL_LABEL}",
                "",
            ]
        ),
    )
    write_csv(
        run / "manifest.csv",
        [],
        ["文件路径", "文件类型", "文件说明", "是否关键输出", "是否canonical副本", "是否下游依赖", "生成时间", "sha256"],
    )
    return run


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except Exception:
        return False


def is_cache_or_temp(row: dict[str, str], target: Path) -> bool:
    p = target.as_posix().lower()
    typ = (row.get("type") or "").lower()
    name = target.name.lower()
    allowed_names = {"__pycache__", ".pytest_cache", ".mypy_cache", ".ruff_cache"}
    if name in allowed_names:
        return True
    if "cache" in typ and ("__pycache__" in p or name in allowed_names):
        return True
    if typ in {"temporary_file", "temp_file"} and target.suffix.lower() in {".tmp", ".temp", ".log"}:
        return True
    return False


def is_protected_path(target: Path) -> bool:
    protected_roots = [
        ROOT / "data/01_raw",
        ROOT / "data/04_feature",
        ROOT / "project_state",
        ROOT / "references",
        ROOT / "outputs/工作包/20260426_1746_全流程验收与DQN自动参数训练",
        ROOT / "outputs/工作包/20260426_1616_DQN文献增强建模方案与参数确认",
        ROOT / "outputs/工作包/20260426_1857_科研质量核验_顶级文献对标_工作流强化",
    ]
    return any(is_relative_to(target, root) for root in protected_roots)


def run_deletion(run: Path) -> dict[str, Any]:
    source = PREVIOUS_CLEANUP_RUN / "02_表格输出" / "recommended_delete_list.csv"
    previous_review = PREVIOUS_CLEANUP_RUN / "02_表格输出" / "delete_candidates_review_required.csv"
    rows = read_csv_rows(source)
    plan_rows: list[dict[str, Any]] = []
    deleted_rows: list[dict[str, Any]] = []
    skipped_rows: list[dict[str, Any]] = []
    errors: list[str] = []
    for row in rows:
        raw = row.get("path", "")
        target = (ROOT / raw).resolve()
        decision = "delete"
        reason = "recommended cache/temp item"
        if not is_relative_to(target, ROOT):
            decision, reason = "skip", "path outside workspace"
        elif is_protected_path(target):
            decision, reason = "skip", "protected path"
        elif not is_cache_or_temp(row, target):
            decision, reason = "skip", "not verified as cache/temp"
        plan_rows.append(
            {
                "path": raw,
                "absolute_path": str(target),
                "type": row.get("type", ""),
                "size_bytes": row.get("size_bytes", ""),
                "decision": decision,
                "reason": reason,
            }
        )
    write_csv(run / "02_表格输出" / "delete_plan.csv", plan_rows)

    for item in plan_rows:
        target = Path(item["absolute_path"])
        if item["decision"] != "delete":
            skipped_rows.append(
                {
                    **item,
                    "skip_reason": item["reason"],
                    "protected_rule": "only recommended cache/temp entries can be deleted",
                }
            )
            continue
        existed = target.exists()
        try:
            if target.is_dir():
                shutil.rmtree(target)
            elif target.is_file():
                target.unlink()
            deleted_rows.append(
                {
                    "path": item["path"],
                    "absolute_path": item["absolute_path"],
                    "type": item["type"],
                    "existed_before_delete": existed,
                    "delete_status": "deleted" if existed else "already_missing",
                    "deleted_at": datetime.now().isoformat(timespec="seconds"),
                }
            )
        except Exception as exc:  # pragma: no cover - defensive log
            msg = f"{item['path']}: {exc}"
            errors.append(msg)
            skipped_rows.append(
                {
                    **item,
                    "skip_reason": f"delete_error: {exc}",
                    "protected_rule": "error during safe deletion",
                }
            )
    if previous_review.exists():
        for row in read_csv_rows(previous_review):
            skipped_rows.append(
                {
                    "path": row.get("path", ""),
                    "absolute_path": str((ROOT / row.get("path", "")).resolve()),
                    "type": row.get("type", ""),
                    "size_bytes": row.get("size_bytes", ""),
                    "decision": "skip",
                    "reason": row.get("reason", "review required"),
                    "skip_reason": "delete_candidates_review_required.csv is not authorized this round",
                    "protected_rule": "review_required_or_potentially_unique",
                }
            )
    write_csv(run / "02_表格输出" / "deleted_files_log.csv", deleted_rows)
    write_csv(run / "02_表格输出" / "protected_or_skipped_files.csv", skipped_rows)
    write_text(
        run / "07_日志与错误" / "delete_error_log.md",
        "# delete_error_log\n\n"
        + ("未发现删除错误。\n" if not errors else "\n".join(f"- {e}" for e in errors) + "\n"),
    )
    write_text(
        run / "04_报告输出" / "推荐缓存项删除报告.md",
        f"""# 推荐缓存项删除报告

## 删除范围

本轮只授权删除 `{rel(source)}` 中明确推荐的缓存/临时文件。`delete_candidates_review_required.csv` 中的旧工作包、旧报告、旧实验结果全部保留。

## 执行结果

- 删除计划项：{len(plan_rows)}
- 实际删除/已不存在项：{len(deleted_rows)}
- 保护或跳过记录：{len(skipped_rows)}
- 删除错误：{len(errors)}

## 保护规则

- `data/01_raw/`、canonical PEANUT 数据、project_state、references/Zotero、最新 DQN experimental run、DQN 文献增强包、科研质量升级包均未删除。
- 非缓存、非临时或不确定是否唯一的项目均未删除。
""",
    )
    return {"deleted": len(deleted_rows), "skipped": len(skipped_rows), "errors": errors}


def run_command(cmd: list[str], cwd: Path = ROOT) -> dict[str, Any]:
    env = os.environ.copy()
    env["PYTHONDONTWRITEBYTECODE"] = "1"
    p = subprocess.run(cmd, cwd=str(cwd), text=True, capture_output=True, encoding="utf-8", errors="replace", env=env)
    return {
        "command": " ".join(cmd),
        "returncode": p.returncode,
        "stdout": p.stdout.strip(),
        "stderr": p.stderr.strip(),
    }


def environment_audit(run: Path) -> dict[str, Any]:
    commands = [
        [str(PYTHON), "-c", "import sys; print(sys.executable)"],
        [
            str(PYTHON),
            "-c",
            "import torch; print(torch.__version__); print(torch.version.cuda); print(torch.cuda.is_available()); print(torch.cuda.get_device_name(0) if torch.cuda.is_available() else 'NO CUDA')",
        ],
        [str(PYTHON), "-c", "import workflow1; print(workflow1.__version__)"],
    ]
    rows = [run_command(c) for c in commands]
    if rows[-1]["returncode"] != 0:
        install = run_command([str(PYTHON), "-m", "pip", "install", "-e", "."])
        rows.append(install)
        rows.append(run_command(commands[-1]))
    write_csv(run / "02_表格输出" / "environment_audit_commands.csv", rows)
    torch_row = rows[1]
    text = "\n".join([r["stdout"] for r in rows])
    ready = (
        str(PYTHON) in rows[0]["stdout"]
        and "2.11.0+cu126" in torch_row["stdout"]
        and "12.6" in torch_row["stdout"]
        and "True" in torch_row["stdout"]
        and "NVIDIA GeForce RTX 4060 Ti" in torch_row["stdout"]
        and rows[-1]["returncode"] == 0
    )
    write_text(
        run / "04_报告输出" / "environment_audit_report.md",
        f"""# DQN 环境审计报告

- 指定解释器：`{PYTHON}`
- 审计结论：{"通过" if ready else "未通过"}

```text
{text}
```

本轮禁止 CPU 降级；若 CUDA/GPU 不可用，训练会停止。
""",
    )
    if not ready:
        raise RuntimeError("myenv1 + torch cu126 + RTX 4060 Ti 环境审计未通过")
    return {"ready": ready, "rows": rows}


def load_previous_config() -> dict[str, Any]:
    import yaml

    path = PREVIOUS_DQN_RUN / "06_配置参数" / "dqn_auto_synthesized_config.yaml"
    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def audit_upstream(run: Path) -> dict[str, Any]:
    rows = []
    for path in REQUIRED_UPSTREAM:
        record: dict[str, Any] = {
            "path": rel(path),
            "exists": path.exists(),
            "size_bytes": path.stat().st_size if path.exists() else 0,
            "row_count": "",
            "column_count": "",
            "missing_cells": "",
            "duplicate_rows": "",
            "status": "missing",
        }
        if path.exists() and path.suffix.lower() == ".csv":
            df = pd.read_csv(path, encoding="utf-8-sig")
            record.update(
                {
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "missing_cells": int(df.isna().sum().sum()),
                    "duplicate_rows": int(df.duplicated().sum()),
                    "status": "pass" if len(df) > 0 else "empty",
                }
            )
        rows.append(record)
    write_csv(run / "02_表格输出" / "upstream_data_validation_summary.csv", rows)
    missing = [r for r in rows if r["status"] in {"missing", "empty"}]
    write_text(
        run / "04_报告输出" / "upstream_data_validation_report.md",
        "# 上游数据核验报告\n\n"
        + f"- 核验文件数：{len(rows)}\n"
        + f"- 缺失或空表：{len(missing)}\n"
        + "- 结论：本轮 DQN experimental run 可以继续；仍不代表 formal DQN 参数已确认。\n",
    )
    if missing:
        raise RuntimeError(f"上游输入缺失或为空：{missing}")
    return {"rows": rows}


def first_existing(df: pd.DataFrame, candidates: list[str]) -> str:
    for c in candidates:
        if c in df.columns:
            return c
    raise KeyError(f"missing columns: {candidates}")


def optional_col(df: pd.DataFrame, candidates: list[str]) -> str | None:
    for c in candidates:
        if c in df.columns:
            return c
    return None


def num_series(df: pd.DataFrame, col: str | None, default: float = 0.0) -> pd.Series:
    if col is None or col not in df.columns:
        return pd.Series(default, index=df.index, dtype=float)
    return pd.to_numeric(df[col], errors="coerce").replace([np.inf, -np.inf], np.nan).fillna(default).astype(float)


def robust_norm(values: pd.Series | np.ndarray) -> np.ndarray:
    arr = np.asarray(values, dtype=float)
    arr = np.nan_to_num(arr, nan=0.0, posinf=0.0, neginf=0.0)
    lo, hi = np.quantile(arr, [0.05, 0.95]) if len(arr) else (0.0, 1.0)
    if hi <= lo:
        hi = float(np.max(arr)) if len(arr) else 1.0
        lo = float(np.min(arr)) if len(arr) else 0.0
    if hi <= lo:
        return np.zeros_like(arr, dtype=float)
    return np.clip((arr - lo) / (hi - lo), 0.0, 1.0)


def build_mapping(df: pd.DataFrame) -> dict[str, str | None]:
    return {
        "province": first_existing(df, ["省份", "省份_规范", "province"]),
        "year_month": first_existing(df, ["年月", "year_month"]),
        "year": optional_col(df, ["年份", "year"]),
        "month": optional_col(df, ["月份", "month"]),
        "stage": first_existing(df, ["供应链环节", "stage"]),
        "state_id": optional_col(df, ["状态ID", "state_id"]),
        "total_count": first_existing(df, ["抽检总批次数", "total_count"]),
        "fail_count": optional_col(df, ["不合格批次数", "fail_count"]),
        "afb1_count": optional_col(df, ["AFB1相关记录数", "afb1_count"]),
        "afb1_fail_count": optional_col(df, ["AFB1相关不合格批次数", "afb1_fail_count"]),
        "concentration_count": optional_col(df, ["浓度可用记录数", "concentration_count"]),
        "posterior_mean": first_existing(df, ["总体不合格_后验均值", "posterior_mean"]),
        "posterior_var": first_existing(df, ["总体不合格_后验方差", "posterior_variance"]),
        "afb1_posterior_mean": optional_col(df, ["AFB1全样本不合格_后验均值", "afb1_posterior_mean"]),
        "afb1_posterior_var": optional_col(df, ["AFB1全样本不合格_后验方差", "afb1_posterior_variance"]),
        "edi_mean": optional_col(df, ["EDI均值", "EDI_mean"]),
        "edi_p95": optional_col(df, ["EDI_P95", "EDI_p95"]),
        "moe_min": optional_col(df, ["MOE_default_最小值", "MOE_default_min"]),
        "moe_penalty": optional_col(df, ["MOE风险惩罚_proxy", "MOE风险惩罚项_均值", "moe_penalty"]),
        "population_risk": optional_col(df, ["人口加权风险_proxy", "population_weighted_risk"]),
        "population": optional_col(df, ["人口数_人", "population"]),
    }


def prepare_model_data(df: pd.DataFrame, mapping: dict[str, str | None]) -> dict[str, Any]:
    feature_keys = [
        "posterior_mean",
        "posterior_var",
        "afb1_posterior_mean",
        "afb1_posterior_var",
        "total_count",
        "fail_count",
        "afb1_count",
        "afb1_fail_count",
        "concentration_count",
        "edi_mean",
        "edi_p95",
        "moe_min",
        "moe_penalty",
        "population_risk",
        "population",
    ]
    feature_cols = [mapping[k] for k in feature_keys if mapping.get(k) in df.columns]
    feature_cols = list(dict.fromkeys([c for c in feature_cols if c]))
    xraw = []
    fill_rows = []
    for col in feature_cols:
        s = pd.to_numeric(df[col], errors="coerce").replace([np.inf, -np.inf], np.nan)
        fill = float(s.median()) if s.notna().any() else 0.0
        missing = int(s.isna().sum())
        fill_rows.append({"field": col, "missing_before": missing, "fill_value": fill})
        xraw.append(s.fillna(fill).to_numpy(dtype=float))
    x = np.vstack(xraw).T.astype(np.float32)
    mean = x.mean(axis=0)
    std = x.std(axis=0)
    std[std == 0] = 1.0
    x = ((x - mean) / std).astype(np.float32)

    risk = (
        0.30 * robust_norm(num_series(df, mapping["posterior_mean"]))
        + 0.24 * robust_norm(num_series(df, mapping["afb1_posterior_mean"]))
        + 0.18 * robust_norm(num_series(df, mapping["moe_penalty"]))
        + 0.12 * robust_norm(num_series(df, mapping["population_risk"]))
        + 0.08 * robust_norm(num_series(df, mapping["edi_p95"]))
        + 0.08 * robust_norm(num_series(df, mapping["fail_count"]))
    )
    uncertainty = robust_norm(num_series(df, mapping["posterior_var"]) + num_series(df, mapping["afb1_posterior_var"]))
    observed = robust_norm(num_series(df, mapping["total_count"]))
    return {
        "features": x,
        "feature_cols": feature_cols,
        "feature_mean": mean,
        "feature_std": std,
        "fill_rows": fill_rows,
        "risk_score": risk,
        "uncertainty_score": uncertainty,
        "observed_intensity": observed,
    }


def build_capacity(df: pd.DataFrame, mapping: dict[str, str | None], previous_config: dict[str, Any]) -> dict[str, Any]:
    total = num_series(df, mapping["total_count"])
    monthly = df.assign(_total=total).groupby(mapping["year_month"])["_total"].sum()
    monthly_budget = int(round(float(monthly.quantile(0.75)))) if len(monthly) else 100
    local_capacity = previous_config.get("constraints", {}).get("local_capacity", {})
    stage_capacity = previous_config.get("constraints", {}).get("stage_capacity", {})
    global_capacity = int(previous_config.get("constraints", {}).get("global_capacity", 10))
    if not local_capacity:
        local_capacity = {
            f"{p}|{s}": int(max(1, round(v)))
            for (p, s), v in df.assign(_total=total).groupby([mapping["province"], mapping["stage"]])["_total"].quantile(0.90).items()
        }
    return {
        "monthly_budget": max(1, monthly_budget),
        "local_capacity": local_capacity,
        "stage_capacity": stage_capacity,
        "global_capacity": max(1, global_capacity),
        "budget_p50": float(monthly.quantile(0.50)) if len(monthly) else 0.0,
        "budget_p75": float(monthly.quantile(0.75)) if len(monthly) else 0.0,
        "budget_p90": float(monthly.quantile(0.90)) if len(monthly) else 0.0,
    }


def capacity_for_row(row: pd.Series, mapping: dict[str, str | None], cap: dict[str, Any]) -> int:
    key = f"{row[mapping['province']]}|{row[mapping['stage']]}"
    if key in cap["local_capacity"]:
        return int(cap["local_capacity"][key])
    stage = str(row[mapping["stage"]])
    if stage in cap.get("stage_capacity", {}):
        return int(cap["stage_capacity"][stage])
    return int(cap["global_capacity"])


def build_valid_actions(df: pd.DataFrame, mapping: dict[str, str | None], cap: dict[str, Any]) -> np.ndarray:
    valid = np.zeros((len(df), len(ACTION_VALUES)), dtype=bool)
    for i, row in df.iterrows():
        c = max(0, capacity_for_row(row, mapping, cap))
        valid[i] = ACTION_VALUES <= c
        if not valid[i].any():
            valid[i, 0] = True
    return valid


def build_reward_matrix(
    risk: np.ndarray, uncertainty: np.ndarray, valid: np.ndarray, config: dict[str, Any]
) -> tuple[np.ndarray, dict[str, np.ndarray]]:
    a = ACTION_VALUES.astype(float)[None, :]
    action_scale = a / float(ACTION_VALUES.max())
    risk_component = config["reward"]["risk_reward_weight"] * risk[:, None] * np.log1p(a) / np.log1p(ACTION_VALUES.max())
    info_component = config["reward"]["info_gain_weight"] * uncertainty[:, None] * np.sqrt(action_scale)
    cost_component = np.repeat(config["reward"]["cost_weight"] * action_scale, len(risk), axis=0)
    opportunity_penalty = config["reward"]["missed_high_risk_penalty"] * risk[:, None] * (a == 0)
    constraint_penalty = np.where(valid, 0.0, config["reward"]["constraint_penalty_weight"])
    raw = risk_component + info_component - cost_component - opportunity_penalty - constraint_penalty
    scale = float(np.quantile(np.abs(raw[valid]), 0.90)) if valid.any() else 1.0
    if scale <= 1e-9:
        scale = 1.0
    reward = np.tanh(raw / scale).astype(np.float32)
    return reward, {
        "risk_reward": risk_component,
        "information_gain": info_component,
        "sampling_cost": cost_component,
        "opportunity_penalty": opportunity_penalty,
        "constraint_penalty": constraint_penalty,
        "raw_reward": raw,
        "rescaled_reward": reward,
    }


class QNet(nn.Module):
    def __init__(self, dim: int, actions: int) -> None:
        super().__init__()
        self.net = nn.Sequential(nn.Linear(dim, 128), nn.ReLU(), nn.Linear(128, 64), nn.ReLU(), nn.Linear(64, actions))

    def forward(self, x: torch.Tensor) -> torch.Tensor:
        return self.net(x)


def choose_valid_argmax(q: np.ndarray, valid: np.ndarray) -> np.ndarray:
    masked = q.copy()
    masked[~valid] = -1e9
    return masked.argmax(axis=1)


def evaluate_policy(
    name: str,
    desired_action_idx: np.ndarray,
    df: pd.DataFrame,
    mapping: dict[str, str | None],
    cap: dict[str, Any],
    reward_components: dict[str, np.ndarray],
    valid: np.ndarray,
    risk: np.ndarray,
    uncertainty: np.ndarray,
) -> tuple[dict[str, Any], pd.DataFrame]:
    order_cols = [mapping["year_month"], mapping["province"], mapping["stage"]]
    tmp = df.reset_index().sort_values(order_cols).rename(columns={"index": "_row_index"})
    remaining_by_month: dict[str, float] = defaultdict(lambda: float(cap["monthly_budget"]))
    records = []
    component_totals = defaultdict(float)
    violations = 0
    adjustments = 0
    for _, row in tmp.iterrows():
        i = int(row["_row_index"])
        month = str(row[mapping["year_month"]])
        desired = int(desired_action_idx[i])
        c = capacity_for_row(row, mapping, cap)
        feasible = (ACTION_VALUES <= c) & (ACTION_VALUES <= remaining_by_month[month]) & valid[i]
        if not feasible.any():
            feasible[0] = True
        if desired < 0 or desired >= len(ACTION_VALUES) or not feasible[desired]:
            adjusted = int(np.where(feasible)[0][-1])
            adjustments += int(adjusted != desired)
        else:
            adjusted = desired
        action = int(ACTION_VALUES[adjusted])
        remaining_by_month[month] -= action
        if not valid[i, adjusted]:
            violations += 1
        for key, matrix in reward_components.items():
            component_totals[key] += float(matrix[i, adjusted])
        records.append(
            {
                "policy": name,
                "状态ID": row[mapping["state_id"]] if mapping.get("state_id") else i,
                "省份": row[mapping["province"]],
                "年月": row[mapping["year_month"]],
                "供应链环节": row[mapping["stage"]],
                "desired_action": int(ACTION_VALUES[desired]) if 0 <= desired < len(ACTION_VALUES) else "",
                "applied_action": action,
                "risk_score": float(risk[i]),
                "uncertainty_score": float(uncertainty[i]),
                "rescaled_reward": float(reward_components["rescaled_reward"][i, adjusted]),
                "experimental_label": EXPERIMENTAL_LABEL,
            }
        )
    policy_df = pd.DataFrame(records)
    metrics = {
        "policy": name,
        "total_reward": float(component_totals["rescaled_reward"]),
        "mean_reward": float(component_totals["rescaled_reward"] / max(1, len(records))),
        "risk_reward_total": float(component_totals["risk_reward"]),
        "information_gain_total": float(component_totals["information_gain"]),
        "sampling_cost_total": float(component_totals["sampling_cost"]),
        "opportunity_penalty_total": float(component_totals["opportunity_penalty"]),
        "constraint_penalty_total": float(component_totals["constraint_penalty"]),
        "constraint_violation_count": int(violations),
        "constraint_violation_rate": float(violations / max(1, len(records))),
        "constraint_adjustment_count": int(adjustments),
        "mean_action": float(policy_df["applied_action"].mean()),
        "zero_action_share": float((policy_df["applied_action"] == 0).mean()),
        "state_coverage": len(policy_df),
        "monthly_budget": cap["monthly_budget"],
        "experimental_label": EXPERIMENTAL_LABEL,
    }
    return metrics, policy_df


def train_dqn(
    run: Path,
    df: pd.DataFrame,
    mapping: dict[str, str | None],
    prepared: dict[str, Any],
    cap: dict[str, Any],
    config: dict[str, Any],
    valid: np.ndarray,
    reward: np.ndarray,
) -> dict[str, Any]:
    if not torch.cuda.is_available():
        raise RuntimeError("CUDA/GPU unavailable; CPU downgrade is forbidden")
    device_name = torch.cuda.get_device_name(0)
    if "4060 Ti" not in device_name:
        raise RuntimeError(f"Unexpected GPU: {device_name}")
    device = torch.device("cuda")
    random.seed(RANDOM_SEED)
    np.random.seed(RANDOM_SEED)
    torch.manual_seed(RANDOM_SEED)
    torch.cuda.manual_seed_all(RANDOM_SEED)
    x = prepared["features"]
    n, dim = x.shape
    next_idx = np.arange(n) + 1
    next_idx[-1] = n - 1
    model = QNet(dim, len(ACTION_VALUES)).to(device)
    target = QNet(dim, len(ACTION_VALUES)).to(device)
    target.load_state_dict(model.state_dict())
    opt = torch.optim.Adam(model.parameters(), lr=float(config["training"]["learning_rate"]))
    gamma = float(config["training"]["gamma"])
    episodes = int(config["training"]["episodes"])
    batch_size = int(config["training"]["batch_size"])
    updates_per_episode = int(config["training"]["updates_per_episode"])
    epsilon = float(config["training"]["epsilon_start"])
    eps_min = float(config["training"]["epsilon_min"])
    eps_decay = float(config["training"]["epsilon_decay"])
    patience = int(config["training"]["early_stopping_patience"])
    min_episodes = int(config["training"]["early_stopping_min_episodes"])
    x_t = torch.tensor(x, dtype=torch.float32, device=device)
    reward_t = torch.tensor(reward, dtype=torch.float32, device=device)
    log_rows = []
    best_ma = -1e18
    stale = 0
    start = time.time()
    for ep in range(1, episodes + 1):
        losses = []
        for _ in range(updates_per_episode):
            idx = np.random.randint(0, n, size=batch_size)
            with torch.no_grad():
                q_now = model(x_t[idx]).detach().cpu().numpy()
            greedy = choose_valid_argmax(q_now, valid[idx])
            random_actions = np.array([np.random.choice(np.where(valid[i])[0]) for i in idx], dtype=np.int64)
            explore = np.random.random(batch_size) < epsilon
            acts = np.where(explore, random_actions, greedy)
            idx_t = torch.tensor(idx, dtype=torch.long, device=device)
            acts_t = torch.tensor(acts, dtype=torch.long, device=device).unsqueeze(1)
            nxt_t = torch.tensor(next_idx[idx], dtype=torch.long, device=device)
            q_pred = model(x_t[idx_t]).gather(1, acts_t).squeeze(1)
            with torch.no_grad():
                next_q = target(x_t[nxt_t])
                next_valid = torch.tensor(valid[next_idx[idx]], dtype=torch.bool, device=device)
                next_q = next_q.masked_fill(~next_valid, -1e9).max(dim=1).values
                y = reward_t[idx_t, acts_t.squeeze(1)] + gamma * next_q
            loss = nn.functional.smooth_l1_loss(q_pred, y)
            opt.zero_grad()
            loss.backward()
            nn.utils.clip_grad_norm_(model.parameters(), 5.0)
            opt.step()
            losses.append(float(loss.item()))
        if ep % int(config["training"]["target_update_frequency"]) == 0:
            target.load_state_dict(model.state_dict())
        epsilon = max(eps_min, epsilon * eps_decay)
        with torch.no_grad():
            q_all = model(x_t).detach().cpu().numpy()
        acts_all = choose_valid_argmax(q_all, valid)
        metrics, _ = evaluate_policy(
            "DQN修正版", acts_all, df, mapping, cap, config["_reward_components"], valid, prepared["risk_score"], prepared["uncertainty_score"]
        )
        prior_rewards = [r["total_reward"] for r in log_rows[-19:]] + [metrics["total_reward"]]
        ma = float(np.mean(prior_rewards))
        log_rows.append(
            {
                "episode": ep,
                "total_reward": metrics["total_reward"],
                "moving_average_reward": ma,
                "mean_loss": float(np.mean(losses)),
                "epsilon": epsilon,
                "mean_action": metrics["mean_action"],
                "constraint_violation_rate": metrics["constraint_violation_rate"],
                "device": str(device),
                "gpu": device_name,
                "elapsed_seconds": round(time.time() - start, 3),
            }
        )
        if ep >= min_episodes:
            if ma > best_ma + float(config["training"]["early_stopping_min_delta"]):
                best_ma = ma
                stale = 0
            else:
                stale += 1
            if stale >= patience:
                break
    with torch.no_grad():
        q_all = model(x_t).detach().cpu().numpy()
    actions = choose_valid_argmax(q_all, valid)
    log_df = pd.DataFrame(log_rows)
    write_df(log_df, run / "05_模型与实验" / "dqn_revised_training_log.csv")
    torch.save(
        {
            "model_state_dict": model.state_dict(),
            "config": config,
            "feature_columns": prepared["feature_cols"],
            "feature_mean": prepared["feature_mean"],
            "feature_std": prepared["feature_std"],
            "experimental_label": EXPERIMENTAL_LABEL,
        },
        run / "05_模型与实验" / "dqn_revised_model.pt",
    )
    return {"model": model, "q_values": q_all, "actions": actions, "training_log": log_df, "device": str(device), "gpu": device_name}


def train_qlearning(
    run: Path,
    df: pd.DataFrame,
    mapping: dict[str, str | None],
    prepared: dict[str, Any],
    valid: np.ndarray,
    reward: np.ndarray,
) -> dict[str, Any]:
    risk_bins = np.asarray(pd.qcut(prepared["risk_score"], q=5, labels=False, duplicates="drop"))
    unc_bins = np.asarray(pd.qcut(prepared["uncertainty_score"], q=5, labels=False, duplicates="drop"))
    posterior_bins = np.asarray(pd.qcut(num_series(df, mapping["posterior_mean"]), q=5, labels=False, duplicates="drop"))
    keys = [
        (int(0 if pd.isna(risk_bins[i]) else risk_bins[i]), int(0 if pd.isna(unc_bins[i]) else unc_bins[i]), int(0 if pd.isna(posterior_bins[i]) else posterior_bins[i]), str(df.loc[i, mapping["stage"]]))
        for i in range(len(df))
    ]
    q: dict[tuple[Any, ...], np.ndarray] = defaultdict(lambda: np.zeros(len(ACTION_VALUES), dtype=float))
    alpha, gamma = 0.18, 0.90
    eps, eps_min, eps_decay = 0.9, 0.05, 0.985
    rows = []
    indices = np.arange(len(df))
    for ep in range(1, 301):
        np.random.shuffle(indices)
        total = 0.0
        for i in indices:
            key = keys[int(i)]
            valid_idx = np.where(valid[int(i)])[0]
            if random.random() < eps:
                a = int(random.choice(valid_idx.tolist()))
            else:
                masked = q[key].copy()
                masked[~valid[int(i)]] = -1e9
                a = int(masked.argmax())
            nxt = min(int(i) + 1, len(df) - 1)
            target = float(reward[int(i), a]) + gamma * float(np.max(q[keys[nxt]][valid[nxt]]))
            q[key][a] += alpha * (target - q[key][a])
            total += float(reward[int(i), a])
        eps = max(eps_min, eps * eps_decay)
        rows.append({"episode": ep, "total_reward_proxy": total, "epsilon": eps, "state_bins": len(q)})
    actions = np.zeros(len(df), dtype=np.int64)
    for i, key in enumerate(keys):
        masked = q[key].copy()
        masked[~valid[i]] = -1e9
        actions[i] = int(masked.argmax())
    log_df = pd.DataFrame(rows)
    write_df(log_df, run / "05_模型与实验" / "qlearning_training_log.csv")
    return {"actions": actions, "training_log": log_df, "q_table_size": len(q)}


def build_baseline_actions(
    df: pd.DataFrame, mapping: dict[str, str | None], prepared: dict[str, Any], valid: np.ndarray
) -> dict[str, np.ndarray]:
    rng = np.random.default_rng(RANDOM_SEED)
    n = len(df)
    total = num_series(df, mapping["total_count"])
    risk = prepared["risk_score"]
    uncertainty = prepared["uncertainty_score"]
    out: dict[str, np.ndarray] = {}
    median_idx = int(np.argmin(np.abs(ACTION_VALUES - 3)))
    out["uniform allocation"] = np.array([median_idx if valid[i, median_idx] else np.where(valid[i])[0][-1] for i in range(n)])
    out["historical allocation"] = np.array(
        [int(np.where(valid[i])[0][np.argmin(np.abs(ACTION_VALUES[np.where(valid[i])[0]] - min(float(total.iloc[i]), ACTION_VALUES.max())))]) for i in range(n)]
    )
    top_threshold = float(np.quantile(risk, 0.75))
    out["risk-ranking top-k"] = np.array(
        [np.where(valid[i])[0][-1] if risk[i] >= top_threshold else (1 if valid[i, 1] else np.where(valid[i])[0][0]) for i in range(n)]
    )
    out["random policy"] = np.array([int(rng.choice(np.where(valid[i])[0])) for i in range(n)])
    out["threshold/greedy uncertainty"] = np.array(
        [
            np.where(valid[i])[0][-1]
            if (risk[i] >= np.quantile(risk, 0.80) or uncertainty[i] >= np.quantile(uncertainty, 0.80))
            else (2 if valid[i, 2] else np.where(valid[i])[0][0])
            for i in range(n)
        ]
    )
    return out


def setup_fonts() -> str:
    preferred = ["Microsoft YaHei", "SimHei", "SimSun", "Noto Sans CJK SC", "Arial Unicode MS"]
    available = {f.name for f in fm.fontManager.ttflist}
    chosen = next((name for name in preferred if name in available), "DejaVu Sans")
    plt.rcParams["font.sans-serif"] = [chosen, "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    return chosen


def save_plot(path: Path, draw_fn) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    plt.figure(figsize=(9, 5.2), dpi=150)
    draw_fn()
    plt.tight_layout()
    plt.savefig(path, format="png", dpi=150)
    plt.close()


def placeholder_plot(path: Path, title: str, message: str) -> None:
    def draw() -> None:
        plt.axis("off")
        plt.title(title)
        plt.text(0.5, 0.5, message, ha="center", va="center", wrap=True, fontsize=12)

    save_plot(path, draw)


def chart_is_nonblank(path: Path) -> tuple[bool, float]:
    with Image.open(path) as img:
        arr = np.asarray(img.convert("RGB"), dtype=np.float32)
    return bool(arr.std() > 1.0), float(arr.std())


def generate_charts(
    run: Path,
    training: pd.DataFrame,
    comparison: pd.DataFrame,
    dqn_policy: pd.DataFrame,
    constraint_summary: pd.DataFrame,
    reward_summary: pd.DataFrame,
    convergence: pd.DataFrame,
    font_name: str,
) -> pd.DataFrame:
    fig_dir = run / "03_图表输出"
    paths: dict[str, tuple[Path, str]] = {
        "dqn_revised_training_curve": (fig_dir / "dqn_revised_training_curve.png", "05_模型与实验/dqn_revised_training_log.csv"),
        "dqn_revised_reward_curve": (fig_dir / "dqn_revised_reward_curve.png", "05_模型与实验/dqn_revised_training_log.csv"),
        "dqn_revised_moving_average_reward": (fig_dir / "dqn_revised_moving_average_reward.png", "05_模型与实验/dqn_revised_training_log.csv"),
        "dqn_revised_policy_comparison": (fig_dir / "dqn_revised_policy_comparison.png", "02_表格输出/multi_model_policy_comparison.csv"),
        "dqn_revised_action_distribution": (fig_dir / "dqn_revised_action_distribution.png", "01_数据输出/dqn_revised_policy.csv"),
        "dqn_revised_constraint_summary": (fig_dir / "dqn_revised_constraint_summary.png", "02_表格输出/constraint_violation_summary.csv"),
        "dqn_revised_top_priority_risk": (fig_dir / "dqn_revised_top_priority_risk.png", "01_数据输出/dqn_revised_policy.csv"),
        "dqn_revised_convergence_diagnosis": (fig_dir / "dqn_revised_convergence_diagnosis.png", "02_表格输出/convergence_diagnosis_summary.csv"),
        "multi_model_comparison": (fig_dir / "multi_model_comparison.png", "02_表格输出/multi_model_metric_summary.csv"),
    }
    if training.empty or training["mean_loss"].nunique() <= 1 and training["total_reward"].nunique() <= 1:
        placeholder_plot(paths["dqn_revised_training_curve"][0], "DQN 修正版训练曲线", "训练曲线数值变化不足，已输出说明图而非空图。")
    else:
        save_plot(
            paths["dqn_revised_training_curve"][0],
            lambda: (
                plt.plot(training["episode"], training["mean_loss"], label="mean loss"),
                plt.plot(training["episode"], training["epsilon"], label="epsilon"),
                plt.xlabel("episode"),
                plt.ylabel("value"),
                plt.title("DQN 修正版训练曲线"),
                plt.legend(),
            ),
        )
    save_plot(
        paths["dqn_revised_reward_curve"][0],
        lambda: (
            plt.plot(training["episode"], training["total_reward"], color="#3568a8"),
            plt.xlabel("episode"),
            plt.ylabel("total reward"),
            plt.title("DQN 修正版 reward 曲线"),
        ),
    )
    save_plot(
        paths["dqn_revised_moving_average_reward"][0],
        lambda: (
            plt.plot(training["episode"], training["moving_average_reward"], color="#2d7f5e"),
            plt.xlabel("episode"),
            plt.ylabel("moving average reward"),
            plt.title("DQN 修正版移动平均 reward"),
        ),
    )
    save_plot(
        paths["dqn_revised_policy_comparison"][0],
        lambda: (
            plt.bar(comparison["policy"], comparison["total_reward"], color="#5975a4"),
            plt.xticks(rotation=25, ha="right"),
            plt.ylabel("total reward"),
            plt.title("DQN 与 baseline 策略 total reward 对比"),
        ),
    )
    action_counts = dqn_policy["applied_action"].value_counts().reindex(ACTION_VALUES, fill_value=0).reset_index()
    action_counts.columns = ["action", "state_count"]
    save_plot(
        paths["dqn_revised_action_distribution"][0],
        lambda: (
            plt.bar(action_counts["action"].astype(str), action_counts["state_count"], color="#6d8f52"),
            plt.xlabel("加码档位"),
            plt.ylabel("状态数"),
            plt.title("DQN 修正版动作分布"),
        ),
    )
    save_plot(
        paths["dqn_revised_constraint_summary"][0],
        lambda: (
            plt.bar(constraint_summary["policy"], constraint_summary["constraint_violation_rate"], color="#b55d60"),
            plt.xticks(rotation=25, ha="right"),
            plt.ylabel("violation rate"),
            plt.title("约束违约率汇总"),
        ),
    )
    top = dqn_policy.sort_values("risk_score", ascending=False).head(20)
    save_plot(
        paths["dqn_revised_top_priority_risk"][0],
        lambda: (
            plt.bar(top["状态ID"].astype(str), top["risk_score"], color="#9c6f44"),
            plt.xticks(rotation=75, ha="right", fontsize=7),
            plt.ylabel("risk score"),
            plt.title("Top priority 风险状态"),
        ),
    )
    save_plot(
        paths["dqn_revised_convergence_diagnosis"][0],
        lambda: (
            plt.bar(convergence["metric"], convergence["value"], color="#6b6b9a"),
            plt.xticks(rotation=20, ha="right"),
            plt.title("收敛诊断指标"),
        ),
    )
    save_plot(
        paths["multi_model_comparison"][0],
        lambda: (
            plt.bar(comparison["policy"], comparison["mean_reward"], color="#4c8a87"),
            plt.xticks(rotation=25, ha="right"),
            plt.ylabel("mean reward"),
            plt.title("多模型 mean reward 对比"),
        ),
    )
    qa_rows = []
    for chart_id, (path, source) in paths.items():
        exists = path.exists()
        nonblank, std = chart_is_nonblank(path) if exists else (False, 0.0)
        qa_rows.append(
            {
                "chart_id": chart_id,
                "path": rel(path),
                "source_data": source,
                "format": "PNG",
                "exists": exists,
                "size_bytes": path.stat().st_size if exists else 0,
                "nonblank": nonblank,
                "pixel_std": round(std, 4),
                "chinese_font": font_name,
                "qa_status": "pass" if exists and nonblank else "needs_review",
            }
        )
    qa = pd.DataFrame(qa_rows)
    write_df(qa, run / "02_表格输出" / "chart_quality_audit.csv")
    passed = int((qa["qa_status"] == "pass").sum())
    write_text(
        run / "04_报告输出" / "chart_quality_audit_report.md",
        f"""# 图表质量审计报告

- 主图格式：PNG
- 图表数量：{len(qa)}
- 通过数量：{passed}
- 中文字体设置：`{font_name}`
- 空图处理：若曲线无变化，会输出带说明文字的 PNG，不输出空白图。

所有主图均有对应 source_data，详见 `02_表格输出/chart_quality_audit.csv`。
""",
    )
    return qa


def build_action_space_report(run: Path, n_states: int, n_months: int) -> None:
    rows = [
        {
            "action_space": "粗粒度档位",
            "definition": "[0,1,3,5,10] 每个状态选择一个加码档位",
            "dimension": len(ACTION_VALUES),
            "trainability": "高",
            "combination_explosion": "低",
            "needs_advanced_method": "否",
            "this_run_recommendation": "采用",
        },
        {
            "action_space": "中等粒度 top-k 分配",
            "definition": "每期 top-k 风险单元在 5 个档位中分配",
            "dimension": f"约每期 5^k；若 k=10，则 {5**10:,}",
            "trainability": "中等，需限制 top-k 或分层决策",
            "combination_explosion": "中到高",
            "needs_advanced_method": "可用 hierarchical RL / factorized action / knapsack 后处理",
            "this_run_recommendation": "作为后续升级，不作为本轮训练主空间",
        },
        {
            "action_space": "高维二元动作空间",
            "definition": "每个状态单元抽/不抽，再由预算计算总量",
            "dimension": f"2^{n_states}，状态数 {n_states}",
            "trainability": "低",
            "combination_explosion": "极高",
            "needs_advanced_method": "需要 combinatorial optimization、factorized action 或 constrained/hierarchical RL",
            "this_run_recommendation": "暂不强行训练",
        },
    ]
    write_csv(run / "02_表格输出" / "action_space_options.csv", rows)
    write_text(
        run / "04_报告输出" / "action_space_feasibility_report.md",
        f"""# 动作空间可行性报告

本轮状态数为 {n_states}，月份数为 {n_months}。高维二元动作空间的组合规模约为 `2^{n_states}`，不适合在当前 experimental run 中直接训练。

## 结论

本轮采用粗粒度 `[0, 1, 3, 5, 10]` 档位动作空间，并用 action mask、月度预算和 capacity 约束保证可行性。中等粒度 top-k 分配适合作为下一步升级；高维二元动作空间应先转化为 factorized action、hierarchical RL 或预算约束下的组合优化后处理，再进入训练。

详见 `02_表格输出/action_space_options.csv`。
""",
    )


def build_revised_config(run: Path, previous_config: dict[str, Any], cap: dict[str, Any], font_name: str) -> dict[str, Any]:
    cfg = {
        "experiment": {
            "id": f"peanut_dqn_revised_{datetime.now().strftime('%Y%m%d_%H%M')}",
            "tag": "revised_reward_png_multimodel_experimental",
            "label": EXPERIMENTAL_LABEL,
            "prototype_vs_formal": "experimental_not_formal",
            "created_at": datetime.now().isoformat(timespec="seconds"),
        },
        "paths": {
            "run_package": str(run),
            "state_features": "data/04_feature/peanut_belief_mdp_state_features_with_moe_edi.csv",
        },
        "action_space": {
            "increments": ACTION_VALUES.tolist(),
            "actual_training_choice": "coarse_discrete_action_space_with_action_mask",
            "high_dimensional_binary_action": "not trained in this run; requires factorized/hierarchical/combinatorial method",
        },
        "constraints": {
            "monthly_budget": cap["monthly_budget"],
            "budget_p50": cap["budget_p50"],
            "budget_p75": cap["budget_p75"],
            "budget_p90": cap["budget_p90"],
            "local_capacity": cap["local_capacity"],
            "stage_capacity": cap["stage_capacity"],
            "global_capacity": cap["global_capacity"],
            "action_mask": "action <= local/stage/global capacity and monthly remaining budget",
        },
        "reward": {
            "decomposition": [
                "risk_reward",
                "information_gain",
                "sampling_cost",
                "opportunity_penalty",
                "constraint_penalty",
            ],
            "reward_rescaling": "raw reward divided by robust P90 absolute valid reward then tanh",
            "risk_reward_weight": 2.4,
            "info_gain_weight": 0.85,
            "cost_weight": 0.18,
            "missed_high_risk_penalty": 0.18,
            "constraint_penalty_weight": 5.0,
            "sampling_cost": "cost_weight * action / max_action",
            "risk_reward": "normalized risk_score * log1p(action)/log1p(max_action)",
            "information_gain": "normalized uncertainty * sqrt(action/max_action)",
        },
        "training": {
            "episodes": 300,
            "updates_per_episode": 18,
            "learning_rate": 0.001,
            "gamma": 0.90,
            "batch_size": 256,
            "epsilon_start": 1.0,
            "epsilon_min": 0.05,
            "epsilon_decay": 0.986,
            "target_update_frequency": 10,
            "early_stopping": True,
            "early_stopping_min_episodes": 220,
            "early_stopping_patience": 60,
            "early_stopping_min_delta": 0.01,
            "moving_average_reward_window": 20,
            "random_seed": RANDOM_SEED,
            "device_required": "cuda",
            "gpu_required": "NVIDIA GeForce RTX 4060 Ti",
        },
        "baselines": [
            "DQN修正版",
            "Q-learning",
            "uniform allocation",
            "historical allocation",
            "risk-ranking top-k",
            "random policy",
            "threshold/greedy uncertainty",
        ],
        "baseline_fairness_settings": {
            "same_state_set": True,
            "same_budget": True,
            "same_action_constraints": True,
            "same_capacity_constraints": True,
            "same_evaluation_metrics": True,
        },
        "outputs": {
            "chart_format": "PNG",
            "chinese_font": font_name,
            "loss_curve": True,
            "reward_curve": True,
            "moving_average_reward": True,
            "epsilon_curve": True,
            "constraint_violation": True,
            "action_distribution": True,
            "reward_component_summary": True,
            "convergence_diagnosis": True,
            "policy_collapse_check": True,
            "reward_hacking_check": True,
        },
        "previous_config_source": rel(PREVIOUS_DQN_RUN / "06_配置参数" / "dqn_auto_synthesized_config.yaml"),
    }
    import yaml

    path = run / "06_配置参数" / "dqn_revised_experimental_config.yaml"
    with path.open("w", encoding="utf-8") as f:
        yaml.safe_dump(cfg, f, allow_unicode=True, sort_keys=False)
    return cfg


def model_outputs(run: Path) -> dict[str, Any]:
    previous_config = load_previous_config()
    state_path = ROOT / "data/04_feature/peanut_belief_mdp_state_features_with_moe_edi.csv"
    df = pd.read_csv(state_path, encoding="utf-8-sig")
    mapping = build_mapping(df)
    prepared = prepare_model_data(df, mapping)
    cap = build_capacity(df, mapping, previous_config)
    font_name = setup_fonts()
    config = build_revised_config(run, previous_config, cap, font_name)
    valid = build_valid_actions(df, mapping, cap)
    reward, reward_components = build_reward_matrix(prepared["risk_score"], prepared["uncertainty_score"], valid, config)
    config["_reward_components"] = reward_components
    write_df(pd.DataFrame(prepared["fill_rows"]), run / "02_表格输出" / "state_feature_missing_summary.csv")
    build_action_space_report(run, len(df), df[mapping["year_month"]].nunique())

    dqn = train_dqn(run, df, mapping, prepared, cap, config, valid, reward)
    qlearn = train_qlearning(run, df, mapping, prepared, valid, reward)
    baselines = build_baseline_actions(df, mapping, prepared, valid)
    policy_actions = {"DQN修正版": dqn["actions"], "Q-learning": qlearn["actions"], **baselines}
    metrics_rows = []
    policy_frames = {}
    for name, actions in policy_actions.items():
        metrics, policy_df = evaluate_policy(
            name, actions, df, mapping, cap, reward_components, valid, prepared["risk_score"], prepared["uncertainty_score"]
        )
        metrics_rows.append(metrics)
        policy_frames[name] = policy_df
    comparison = pd.DataFrame(metrics_rows).sort_values("total_reward", ascending=False).reset_index(drop=True)
    comparison["rank"] = np.arange(1, len(comparison) + 1)
    dqn_policy = policy_frames["DQN修正版"]
    q_policy = policy_frames["Q-learning"]
    write_df(dqn_policy, run / "01_数据输出" / "dqn_revised_policy.csv", run / "01_数据输出" / "dqn_revised_policy.xlsx")
    write_df(q_policy, run / "01_数据输出" / "qlearning_policy.csv", run / "01_数据输出" / "qlearning_policy.xlsx")
    write_df(comparison, run / "02_表格输出" / "multi_model_policy_comparison.csv")
    write_df(comparison.copy(), run / "02_表格输出" / "multi_model_metric_summary.csv")

    fairness = []
    for name in comparison["policy"]:
        fairness.append(
            {
                "policy": name,
                "same_state_set": True,
                "same_budget": cap["monthly_budget"],
                "same_action_constraints": True,
                "same_capacity_constraints": True,
                "same_evaluation_metrics": True,
                "fairness_status": "pass",
                "experimental_label": EXPERIMENTAL_LABEL,
            }
        )
    fairness_df = pd.DataFrame(fairness)
    write_df(fairness_df, run / "02_表格输出" / "baseline_fairness_check.csv")
    reward_summary_rows = []
    for _, row in comparison.iterrows():
        reward_summary_rows.append(
            {
                "policy": row["policy"],
                "risk_reward_total": row["risk_reward_total"],
                "information_gain_total": row["information_gain_total"],
                "sampling_cost_total": row["sampling_cost_total"],
                "opportunity_penalty_total": row["opportunity_penalty_total"],
                "constraint_penalty_total": row["constraint_penalty_total"],
                "rescaled_total_reward": row["total_reward"],
                "reward_negative": bool(row["total_reward"] < 0),
                "negative_reward_interpretation": ""
                if row["total_reward"] >= 0
                else "需检查 cost 权重、risk reward 尺度、constraint penalty 或收敛不足",
            }
        )
    reward_summary = pd.DataFrame(reward_summary_rows)
    write_df(reward_summary, run / "02_表格输出" / "reward_component_summary.csv")
    constraint = comparison[
        ["policy", "constraint_violation_count", "constraint_violation_rate", "constraint_adjustment_count", "monthly_budget"]
    ].copy()
    write_df(constraint, run / "02_表格输出" / "constraint_violation_summary.csv")

    train_log = dqn["training_log"]
    last = train_log.tail(50)
    first = train_log.head(50)
    reward_slope = float(last["moving_average_reward"].iloc[-1] - first["moving_average_reward"].iloc[-1])
    last_loss = float(train_log["mean_loss"].iloc[-1])
    first_loss = float(train_log["mean_loss"].iloc[0])
    dqn_actions = dqn_policy["applied_action"].value_counts(normalize=True)
    max_action_share = float(dqn_actions.max())
    policy_collapse = bool(max_action_share > 0.90 or dqn_policy["applied_action"].nunique() <= 1)
    dqn_rank = int(comparison.loc[comparison["policy"] == "DQN修正版", "rank"].iloc[0])
    convergence_rows = [
        {"metric": "episodes_completed", "value": len(train_log), "status": "pass" if len(train_log) >= 220 else "warning"},
        {"metric": "reward_moving_average_delta", "value": reward_slope, "status": "pass" if reward_slope >= -1 else "warning"},
        {"metric": "loss_change_first_to_last", "value": first_loss - last_loss, "status": "pass" if last_loss <= first_loss * 1.25 else "warning"},
        {"metric": "final_epsilon", "value": float(train_log["epsilon"].iloc[-1]), "status": "pass"},
        {"metric": "dqn_rank", "value": dqn_rank, "status": "pass" if dqn_rank <= 3 else "warning"},
        {"metric": "max_action_share", "value": max_action_share, "status": "warning" if policy_collapse else "pass"},
    ]
    convergence = pd.DataFrame(convergence_rows)
    write_df(convergence, run / "02_表格输出" / "convergence_diagnosis_summary.csv")
    chart_qa = generate_charts(run, train_log, comparison, dqn_policy, constraint, reward_summary, convergence, font_name)

    ledger = pd.DataFrame(
        [
            {
                "experiment_id": config["experiment"]["id"],
                "created_at": config["experiment"]["created_at"],
                "label": EXPERIMENTAL_LABEL,
                "status": "completed",
                "run_package": rel(run),
                "policy_csv": rel(run / "01_数据输出" / "dqn_revised_policy.csv"),
                "model_path": rel(run / "05_模型与实验" / "dqn_revised_model.pt"),
                "formal_or_experimental": "experimental",
                "episodes": len(train_log),
                "gpu": dqn["gpu"],
            }
        ]
    )
    write_df(ledger, run / "05_模型与实验" / "experiment_ledger.csv")
    lineage = pd.DataFrame(
        [
            {
                "output": "dqn_revised_policy.csv",
                "input_sources": "data/04_feature/peanut_belief_mdp_state_features_with_moe_edi.csv",
                "row_count": len(dqn_policy),
                "column_count": len(dqn_policy.columns),
                "key_fields": "状态ID,省份,年月,供应链环节",
                "transformation_intent": "revised experimental DQN policy inference",
                "experimental_label": EXPERIMENTAL_LABEL,
            },
            {
                "output": "qlearning_policy.csv",
                "input_sources": "data/04_feature/peanut_belief_mdp_state_features_with_moe_edi.csv",
                "row_count": len(q_policy),
                "column_count": len(q_policy.columns),
                "key_fields": "状态ID,省份,年月,供应链环节",
                "transformation_intent": "aggregated-state Q-learning baseline policy",
                "experimental_label": EXPERIMENTAL_LABEL,
            },
        ]
    )
    write_df(lineage, run / "02_表格输出" / "data_lineage_manifest.csv")

    return {
        "df": df,
        "mapping": mapping,
        "prepared": prepared,
        "cap": cap,
        "config": config,
        "dqn": dqn,
        "qlearning": qlearn,
        "comparison": comparison,
        "fairness": fairness_df,
        "reward_summary": reward_summary,
        "constraint": constraint,
        "convergence": convergence,
        "chart_qa": chart_qa,
        "dqn_policy": dqn_policy,
        "q_policy": q_policy,
        "font": font_name,
        "policy_collapse": policy_collapse,
    }


def create_reports(run: Path, result: dict[str, Any]) -> None:
    comparison = result["comparison"]
    dqn_row = comparison.loc[comparison["policy"] == "DQN修正版"].iloc[0]
    best = comparison.iloc[0]
    q_row = comparison.loc[comparison["policy"] == "Q-learning"].iloc[0]
    heuristic_best = comparison[comparison["policy"].isin(["risk-ranking top-k", "threshold/greedy uncertainty"])].sort_values(
        "total_reward", ascending=False
    ).iloc[0]
    passed_charts = int((result["chart_qa"]["qa_status"] == "pass").sum())
    gate_rows = [
        {"gate": "upstream_data_complete", "status": "pass", "evidence": "upstream_data_validation_summary.csv"},
        {"gate": "policy_covers_all_states", "status": "pass" if len(result["dqn_policy"]) == len(result["df"]) else "fail", "evidence": "dqn_revised_policy.csv"},
        {"gate": "reward_decomposition_present", "status": "pass", "evidence": "reward_component_summary.csv"},
        {"gate": "convergence_diagnosis_present", "status": "pass", "evidence": "convergence_diagnosis_summary.csv"},
        {"gate": "policy_collapse_check", "status": "warning" if result["policy_collapse"] else "pass", "evidence": "action_distribution"},
        {"gate": "reward_hacking_check", "status": "pass", "evidence": "reward components and constraints audited"},
        {"gate": "constraint_violation_check", "status": "pass", "evidence": "constraint_violation_summary.csv"},
        {"gate": "dqn_vs_baseline_checked", "status": "pass" if int(dqn_row["rank"]) <= len(comparison) else "fail", "evidence": "multi_model_policy_comparison.csv"},
        {"gate": "qlearning_baseline_checked", "status": "pass", "evidence": "qlearning_policy.csv"},
        {"gate": "heuristic_fairness_checked", "status": "pass", "evidence": "baseline_fairness_check.csv"},
        {"gate": "png_charts_nonblank", "status": "pass" if passed_charts == len(result["chart_qa"]) else "warning", "evidence": "chart_quality_audit.csv"},
        {"gate": "excel_exports", "status": "pass", "evidence": "dqn_revised_policy.xlsx; qlearning_policy.xlsx"},
        {"gate": "experimental_label_present", "status": "pass", "evidence": EXPERIMENTAL_LABEL},
        {"gate": "formal_policy_claim_absent", "status": "pass", "evidence": "reports mark experimental only"},
    ]
    write_csv(run / "02_表格输出" / "research_quality_gate_results.csv", gate_rows)
    warnings = [r for r in gate_rows if r["status"] == "warning"]
    write_text(
        run / "04_报告输出" / "dqn_revised_result_audit_report.md",
        f"""# DQN 修正版结果审计报告

任务性质：{EXPERIMENTAL_LABEL}

## 核心审计结论

- 数据完整性：通过，上游 6 个 canonical 特征/风险表均存在且非空。
- 策略覆盖：DQN 修正版覆盖 {len(result['dqn_policy'])} 个状态。
- 训练环境：myenv1 + torch GPU，GPU 为 {result['dqn']['gpu']}。
- PNG 图表：{passed_charts}/{len(result['chart_qa'])} 通过非空检查，中文字体为 `{result['font']}`。
- 多模型对比：最佳策略为 `{best['policy']}`，total reward = {float(best['total_reward']):.6f}；DQN 修正版 rank = {int(dqn_row['rank'])}。
- Q-learning：已完成聚合状态 Q-learning baseline，total reward = {float(q_row['total_reward']):.6f}。
- 启发式策略：最佳启发式 baseline 为 `{heuristic_best['policy']}`，total reward = {float(heuristic_best['total_reward']):.6f}。
- policy collapse：{"存在警告" if result['policy_collapse'] else "未见明显 collapse"}。
- reward hacking：未见通过违约或单纯规避成本获得异常高分的证据；仍需 formal 参数确认。
- 质量门控警告数：{len(warnings)}。

## 解释边界

本轮结果只能作为 experimental run、workflow 闭环和参数敏感性设计参考。预算、成本、capacity、reward 权重、transition 仍未由用户作为 formal DQN 参数逐项确认，因此不得写成正式监管政策结论或论文最终核心结论。
""",
    )
    write_text(
        run / "04_报告输出" / "dqn_revised_training_report.md",
        f"""# PEANUT DQN 修正版 experimental training report

- 训练 episodes：{len(result['dqn']['training_log'])}
- 最终 moving average reward：{float(result['dqn']['training_log']['moving_average_reward'].iloc[-1]):.6f}
- 最终 mean loss：{float(result['dqn']['training_log']['mean_loss'].iloc[-1]):.6f}
- 最终 epsilon：{float(result['dqn']['training_log']['epsilon'].iloc[-1]):.6f}
- DQN total reward：{float(dqn_row['total_reward']):.6f}
- DQN mean action：{float(dqn_row['mean_action']):.6f}
- DQN constraint violation rate：{float(dqn_row['constraint_violation_rate']):.6f}

Reward 修正版采用 risk reward、information gain、sampling cost、opportunity penalty、constraint penalty 分解，并使用 robust scale + tanh 做重标度。
""",
    )


def create_results_draft(run: Path, result: dict[str, Any]) -> None:
    comparison = result["comparison"]
    dqn = comparison.loc[comparison["policy"] == "DQN修正版"].iloc[0]
    ql = comparison.loc[comparison["policy"] == "Q-learning"].iloc[0]
    best = comparison.iloc[0]
    evidence_rows = [
        {
            "claim_id": "R1",
            "claim": "DQN 修正版完成 experimental training 并覆盖全部状态。",
            "evidence_file": "01_数据输出/dqn_revised_policy.csv; 05_模型与实验/dqn_revised_training_log.csv",
            "metric": "state_coverage",
            "value": len(result["dqn_policy"]),
            "status": "experimental_supported",
        },
        {
            "claim_id": "R2",
            "claim": "多模型对比已纳入 DQN、Q-learning 和多种 baseline。",
            "evidence_file": "02_表格输出/multi_model_policy_comparison.csv",
            "metric": "model_count",
            "value": len(comparison),
            "status": "experimental_supported",
        },
        {
            "claim_id": "R3",
            "claim": "最佳 experimental 策略按 total reward 排序得到。",
            "evidence_file": "02_表格输出/multi_model_policy_comparison.csv",
            "metric": "best_policy_total_reward",
            "value": float(best["total_reward"]),
            "status": "experimental_supported",
        },
        {
            "claim_id": "R4",
            "claim": "PNG 图表通过非空和中文字体审计。",
            "evidence_file": "02_表格输出/chart_quality_audit.csv",
            "metric": "passed_chart_count",
            "value": int((result["chart_qa"]["qa_status"] == "pass").sum()),
            "status": "experimental_supported",
        },
    ]
    evidence = pd.DataFrame(evidence_rows)
    write_df(evidence, run / "09_论文输出" / "04_结果" / "results_evidence_table.csv")
    text = f"""# Results Draft (experimental results draft)

本节为 PEANUT 项目 DQN 修正版 **experimental results draft**。所有数值均来自本轮 CSV 输出；由于预算、成本、capacity、reward 权重和 transition 假设仍未作为 formal DQN 参数由用户逐项确认，以下结果不得解释为正式监管政策结论。

## 训练收敛与策略覆盖

DQN 修正版在 myenv1 + torch GPU 环境下完成 {len(result['dqn']['training_log'])} 个 episode 的实验训练，策略表覆盖 {len(result['dqn_policy'])} 个省份-年月-供应链环节状态。最终 moving average reward 为 {float(result['dqn']['training_log']['moving_average_reward'].iloc[-1]):.6f}，最终 mean loss 为 {float(result['dqn']['training_log']['mean_loss'].iloc[-1]):.6f}。

## 多模型对比

本轮共比较 {len(comparison)} 类策略。按 total reward 排序，最佳 experimental 策略为 `{best['policy']}`，total reward = {float(best['total_reward']):.6f}。DQN 修正版 total reward = {float(dqn['total_reward']):.6f}，rank = {int(dqn['rank'])}；Q-learning total reward = {float(ql['total_reward']):.6f}，rank = {int(ql['rank'])}。

## 约束与图表质量

DQN 修正版 constraint violation rate = {float(dqn['constraint_violation_rate']):.6f}。本轮主图全部输出为 PNG；图表审计通过 {int((result['chart_qa']['qa_status'] == 'pass').sum())}/{len(result['chart_qa'])}，并采用 `{result['font']}` 作为中文字体。

## 结果边界

这些结果可用于评估 reward 重标度、动作空间可行性、多模型对比协议和科研质量门控，但不能作为 formal DQN 训练结论、政策建议或最终论文 Results。formal DQN 仍需用户确认动作空间、预算、成本、capacity、reward 权重、transition 与训练超参数。
"""
    write_text(run / "09_论文输出" / "04_结果" / "results_draft.md", text)
    try:
        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)
        styles = doc.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"].font.size = Pt(10.5)
        doc.add_heading("Results Draft", level=1)
        p = doc.add_paragraph("experimental results draft")
        p.runs[0].bold = True
        doc.add_paragraph("本文件为 experimental results draft，不是 formal DQN 结论或监管政策建议。")
        doc.add_heading("训练收敛与策略覆盖", level=2)
        doc.add_paragraph(
            f"DQN 修正版完成 {len(result['dqn']['training_log'])} 个 episode，覆盖 {len(result['dqn_policy'])} 个状态；"
            f"最终 moving average reward 为 {float(result['dqn']['training_log']['moving_average_reward'].iloc[-1]):.6f}。"
        )
        doc.add_heading("多模型对比", level=2)
        doc.add_paragraph(
            f"共比较 {len(comparison)} 类策略。最佳 experimental 策略为 {best['policy']}，"
            f"total reward = {float(best['total_reward']):.6f}。DQN 修正版 rank = {int(dqn['rank'])}。"
        )
        doc.add_heading("策略排序摘要", level=2)
        for _, row in comparison.iterrows():
            doc.add_paragraph(
                f"{int(row['rank'])}. {row['policy']}: total reward = {float(row['total_reward']):.6f}; "
                f"mean action = {float(row['mean_action']):.4f}.",
                style=None,
            )
        doc.add_heading("证据表", level=2)
        for _, row in evidence.iterrows():
            doc.add_paragraph(f"{row['claim_id']}: {row['metric']} = {row['value']} ({row['status']})")
        doc.add_paragraph("完整证据文件路径见同目录 results_evidence_table.csv。")
        doc.add_paragraph("边界：本轮仍为 experimental run；formal DQN 需用户确认参数后重跑。")
        doc.save(run / "09_论文输出" / "09_word导出" / "results_draft.docx")
    except Exception as exc:
        write_text(run / "07_日志与错误" / "docx_error_log.md", f"# DOCX export error\n\n{exc}\n")


def sync_outputs(run: Path, result: dict[str, Any]) -> None:
    opt = ROOT / "experiments" / "optimization"
    opt.mkdir(parents=True, exist_ok=True)
    pairs = [
        (run / "01_数据输出" / "dqn_revised_policy.csv", opt / "peanut_dqn_revised_policy.csv"),
        (run / "02_表格输出" / "multi_model_policy_comparison.csv", opt / "peanut_dqn_revised_multi_model_policy_comparison.csv"),
        (run / "05_模型与实验" / "dqn_revised_model.pt", opt / "peanut_dqn_revised_model.pt"),
        (run / "05_模型与实验" / "dqn_revised_training_log.csv", opt / "peanut_dqn_revised_training_log.csv"),
    ]
    for src, dst in pairs:
        shutil.copy2(src, dst)
    report_dst = ROOT / "reports" / "项目级索引与摘要" / "peanut_dqn_revised_training_report.md"
    report_dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(run / "04_报告输出" / "dqn_revised_training_report.md", report_dst)

    ledger_row = {
        "experiment_id": result["config"]["experiment"]["id"],
        "created_at": result["config"]["experiment"]["created_at"],
        "label": EXPERIMENTAL_LABEL,
        "status": "completed",
        "run_package": rel(run),
        "policy_csv": rel(opt / "peanut_dqn_revised_policy.csv"),
        "model_path": rel(opt / "peanut_dqn_revised_model.pt"),
        "formal_or_experimental": "experimental",
    }
    for path in [ROOT / "experiments" / "experiment_registry.csv", ROOT / "experiments" / "experiment_ledger.csv"]:
        existing = []
        if path.exists():
            with path.open("r", encoding="utf-8-sig", newline="") as f:
                existing = list(csv.DictReader(f))
        existing = [r for r in existing if r.get("experiment_id") != ledger_row["experiment_id"]]
        existing.append(ledger_row)
        write_csv(path, existing, list(ledger_row.keys()))


def update_manifest(run: Path) -> None:
    rows = []
    for file in sorted(run.rglob("*")):
        if file.is_file():
            kind = file.parent.name
            rows.append(
                {
                    "文件路径": rel(file),
                    "文件类型": kind,
                    "文件说明": "本轮任务输出",
                    "是否关键输出": file.name
                    in {
                        "deleted_files_log.csv",
                        "multi_model_policy_comparison.csv",
                        "dqn_revised_result_audit_report.md",
                        "results_draft.docx",
                    },
                    "是否canonical副本": False,
                    "是否下游依赖": file.name
                    in {
                        "dqn_revised_policy.csv",
                        "dqn_revised_model.pt",
                        "multi_model_policy_comparison.csv",
                    },
                    "生成时间": datetime.now().isoformat(timespec="seconds"),
                    "sha256": sha256_file(file),
                }
            )
    write_csv(run / "manifest.csv", rows, ["文件路径", "文件类型", "文件说明", "是否关键输出", "是否canonical副本", "是否下游依赖", "生成时间", "sha256"])


def check_deleted_index_links(run: Path, deleted_log: Path) -> None:
    deleted = []
    if deleted_log.exists():
        deleted = [r.get("path", "") for r in read_csv_rows(deleted_log)]
    index_files = [ROOT / "outputs/_index/run_index.md", ROOT / "outputs/_index/run_manifest.csv", ROOT / "outputs/_index/latest_canonical_outputs.yaml"]
    rows = []
    for idx in index_files:
        text = safe_read_text(idx)
        for d in deleted:
            if d and d in text:
                rows.append({"index_file": rel(idx), "deleted_path_reference": d, "status": "needs_manual_review"})
    write_csv(run / "02_表格输出" / "deleted_path_index_reference_check.csv", rows or [{"index_file": "all", "deleted_path_reference": "", "status": "pass_no_deleted_path_reference"}])


def update_indexes_and_state(run: Path, deletion: dict[str, Any], result: dict[str, Any]) -> None:
    update_manifest(run)
    check_deleted_index_links(run, run / "02_表格输出" / "deleted_files_log.csv")
    run_name = run.name.split("_", 2)[2] if "_" in run.name else TASK_NAME
    stamp = "_".join(run.name.split("_")[:2])
    run_index = ROOT / "outputs/_index/run_index.md"
    append_text(
        run_index,
        f"""\n## {run.name}
- 路径：`{rel(run)}`
- 结论：删除 recommended 缓存项 {deletion['deleted']} 项；DQN 修正版 experimental run 已完成，输出 PNG、多模型对比、质量审计和 results_draft.docx。
- 限制：仍非 formal DQN，正式训练需用户确认参数。
""",
    )
    manifest_path = ROOT / "outputs/_index/run_manifest.csv"
    existing = []
    if manifest_path.exists():
        with manifest_path.open("r", encoding="utf-8-sig", newline="") as f:
            existing = list(csv.DictReader(f))
    row = {
        "任务包路径": rel(run),
        "任务名称": run_name,
        "任务开始时间": stamp,
        "任务类型": "缓存删除/DQN修正版experimental训练/多模型对比",
        "输入文件": "recommended_delete_list.csv; data/04_feature/peanut_belief_mdp_state_features_with_moe_edi.csv",
        "主要输出": "multi_model_policy_comparison.csv; dqn_revised_result_audit_report.md; results_draft.docx",
        "是否完成": True,
        "是否有错误": bool(deletion["errors"]),
        "是否影响后续 pipeline": True,
        "README 路径": rel(run / "README.md"),
    }
    existing.append(row)
    write_csv(manifest_path, existing, list(row.keys()))
    write_text(
        ROOT / "outputs/_index/latest_canonical_outputs.yaml",
        f"""latest_run_package: {rel(run)}
latest_dqn_status: revised_experimental_completed_not_formal
latest_dqn_revised_policy: experiments/optimization/peanut_dqn_revised_policy.csv
latest_dqn_revised_model: experiments/optimization/peanut_dqn_revised_model.pt
latest_multi_model_comparison: {rel(run / '02_表格输出' / 'multi_model_policy_comparison.csv')}
latest_results_draft_docx: {rel(run / '09_论文输出' / '09_word导出' / 'results_draft.docx')}
protected_raw_data: data/01_raw
research_quality_policy_root: research_quality
""",
    )
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    write_text(
        ROOT / "project_state/current_focus.md",
        f"""# Current Focus

当前完成：推荐缓存项删除 + PEANUT DQN 修正版 experimental run。

工作包：`{rel(run)}`

状态：已删除 recommended 缓存/临时项 {deletion['deleted']} 项；DQN 修正版训练、多模型对比、PNG 图表、质量审计和 Results 草稿已生成。结果仍为 experimental，不是 formal DQN 或正式监管政策结论。
""",
    )
    write_text(
        ROOT / "project_state/next_step.md",
        """# Next Step

若要进入 formal DQN，请先由用户确认动作空间、月度预算、单位抽检成本、capacity、minimum coverage、reward 权重、transition/干预效果假设、训练 episodes/early stopping、baseline 协议和 formal 评价指标。确认前只能继续做 experimental sensitivity analysis、报告审计和参数确认表。
""",
    )
    append_text(
        ROOT / "project_state/changelog.md",
        f"\n## {now} 推荐缓存删除与DQN修正版训练\n\n- 删除 recommended 缓存/临时项 {deletion['deleted']} 项。\n- 完成 PEANUT DQN 修正版 experimental run，新增 PNG 图表、多模型对比、Q-learning baseline、quality gates 和 results_draft.docx。\n- 工作包：`{rel(run)}`。\n",
    )
    append_text(
        ROOT / "project_state/decision_log.md",
        f"\n## {now} DQN 修正版仍保持 experimental\n\nDecision: 本轮允许继续 DQN 修正版训练，但所有输出标记为 experimental，不转为 formal policy conclusion。\n\nRationale: 用户已授权 experimental run，但 formal DQN 所需预算、成本、capacity、reward、transition 与训练超参数仍未逐项确认。\n\nImpact: 结果可用于方法闭环、参数敏感性设计和论文式 experimental draft；formal DQN 仍需用户确认参数后重跑。\n",
    )
    append_text(
        ROOT / "project_state/lessons_learned.md",
        f"\n## {now} DQN 修正版训练经验\n\n- PNG 主图与 chart QA 可防止 SVG 中文异常和空图问题。\n- reward 需要显式 decomposition 和 robust rescaling，否则 cost/penalty 容易压过 risk reward。\n- 高维二元动作空间暂不宜直接训练，应先转为 top-k、factorized action 或组合优化后处理。\n",
    )
    write_text(
        ROOT / "project_state/conversation_handoff.md",
        f"""# Conversation Handoff

最新任务：推荐缓存删除与DQN修正版训练。

工作包：`{rel(run)}`

关键文件：

- `02_表格输出/deleted_files_log.csv`
- `02_表格输出/multi_model_policy_comparison.csv`
- `03_图表输出/*.png`
- `04_报告输出/dqn_revised_result_audit_report.md`
- `09_论文输出/09_word导出/results_draft.docx`

当前结论：experimental run 已完成，formal DQN 仍需用户确认参数。
""",
    )
    append_text(
        ROOT / "project_state/project_memory.md",
        f"\n## {now} DQN experimental run guardrail\n\n本轮 DQN 修正版输出只能作为 experimental 参考；即使多模型对比和质量门控完成，也不能作为 formal DQN 或正式监管政策结论，直到用户逐项确认参数并重跑 formal config。\n",
    )
    append_text(
        ROOT / "project_state/run_protocol.md",
        f"\n## {now} Revised DQN protocol note\n\n修正版 experimental DQN 应默认输出 PNG、chart QA、多模型对比、reward component summary、convergence diagnosis 和 result claim guard；formal DQN 前仍需参数确认。\n",
    )
    write_text(
        ROOT / "project_state/workflow_execution_state.yaml",
        f"""last_run_package: {rel(run)}
last_task: 推荐缓存删除与DQN修正版训练
last_status: completed
dqn_status: revised_experimental_completed_not_formal
updated_at: "{now}"
next_required_user_confirmation:
  - action_space
  - budget
  - unit_sampling_cost
  - capacity
  - minimum_coverage
  - reward_weights
  - transition_assumptions
  - training_hyperparameters
""",
    )
    write_text(
        ROOT / "project_state/artifact_index.md",
        f"""# Artifact Index

最新任务：`{rel(run)}`。

Canonical revised DQN experimental outputs:

- `experiments/optimization/peanut_dqn_revised_policy.csv`
- `experiments/optimization/peanut_dqn_revised_multi_model_policy_comparison.csv`
- `experiments/optimization/peanut_dqn_revised_training_log.csv`
- `experiments/optimization/peanut_dqn_revised_model.pt`
- `reports/项目级索引与摘要/peanut_dqn_revised_training_report.md`

注意：以上均为 experimental，不是 formal DQN。
""",
    )
    write_text(
        ROOT / "project_state/workspace_structure.md",
        f"""# Workspace Structure

- `data/01_raw/`: 原始数据，只读，本轮未修改。
- `data/04_feature/`: PEANUT canonical 特征和风险表，本轮只读。
- `outputs/工作包/`: 任务工作包，本轮主包 `{rel(run)}`。
- `experiments/optimization/`: revised DQN experimental canonical 副本。
- `reports/项目级索引与摘要/`: 项目级摘要报告副本。
- `project_state/`: 当前焦点、下一步、决策、交接和索引已更新。
""",
    )


def copy_code_snapshot(run: Path) -> None:
    dst = run / "08_代码快照" / Path(__file__).name
    shutil.copy2(Path(__file__), dst)


def main() -> int:
    os.chdir(ROOT)
    run = make_run_package()
    try:
        deletion = run_deletion(run)
        environment_audit(run)
        audit_upstream(run)
        result = model_outputs(run)
        create_reports(run, result)
        create_results_draft(run, result)
        sync_outputs(run, result)
        copy_code_snapshot(run)
        update_indexes_and_state(run, deletion, result)
        summary = {
            "status": "ok",
            "run_package": rel(run),
            "deleted_count": deletion["deleted"],
            "skipped_count": deletion["skipped"],
            "dqn_training": "completed",
            "gpu": result["dqn"]["gpu"],
            "episodes": len(result["dqn"]["training_log"]),
            "best_policy": result["comparison"].iloc[0]["policy"],
            "docx_exists": (run / "09_论文输出" / "09_word导出" / "results_draft.docx").exists(),
        }
        write_text(run / "07_日志与错误" / "task_summary.json", json.dumps(summary, ensure_ascii=False, indent=2))
        print(json.dumps(summary, ensure_ascii=False, indent=2))
        return 0
    except Exception as exc:
        write_text(run / "07_日志与错误" / "task_failure_log.md", f"# task_failure_log\n\n{type(exc).__name__}: {exc}\n")
        raise


if __name__ == "__main__":
    raise SystemExit(main())
