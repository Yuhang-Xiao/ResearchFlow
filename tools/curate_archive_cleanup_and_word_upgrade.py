from __future__ import annotations

import hashlib
import json
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path

import pandas as pd
import yaml


ROOT = Path(__file__).resolve().parents[1]
PYTHON = r"D:\anaconda3\envs\myenv1\python.exe"
RUN = ROOT / "outputs" / "工作包" / f"{datetime.now():%Y%m%d_%H%M}_输出解释精简_项目存档与Word学术升级"
SRC_DQN = ROOT / "outputs" / "工作包" / "20260426_2056_推荐缓存删除与DQN修正版训练"
SRC_EXPLAIN = ROOT / "outputs" / "工作包" / "20260426_2254_DQN输出复核_解释体系与论文输出升级"
SRC_LOCAL = ROOT / "outputs" / "工作包" / "20260426_2317_输出解释就地化修正与DQN代码深度说明补强"
ARCHIVE_DIR = ROOT / "archive" / "project_snapshots"


DIRS = [
    "00_输入说明",
    "02_表格输出",
    "03_图表输出",
    "04_报告输出",
    "07_日志与错误",
    "08_代码快照",
    "09_论文输出/04_结果",
    "09_论文输出/09_word导出",
    "10_输出解释与索引",
]


def rel(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT)).replace("\\", "/")
    except ValueError:
        return str(path).replace("\\", "/")


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip() + "\n", encoding="utf-8")


def sha(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def init_run() -> None:
    for d in DIRS:
        (RUN / d).mkdir(parents=True, exist_ok=True)
    write(
        RUN / "00_输入说明" / "inputs.md",
        f"""
        # 输入说明

        本轮处理用户反馈：上一轮解释文件过多、Word 不够像学术论文且图表未充分写入、需要项目存档与缓存/冗余输出清理。

        保护范围：
        - `data/01_raw/`
        - latest DQN training run `{rel(SRC_DQN)}`
        - project_state、outputs/_index、canonical experiments
        - 旧 baseline run `{rel(ROOT / 'outputs/工作包/20260426_1746_全流程验收与DQN自动参数训练')}`

        删除范围仅限：缓存、pycache、我方生成的失败/中间 run 包、过度生成的 `.explanation.md` 侧车和重复 local md。
        """,
    )


def copy_key_inputs() -> None:
    key_files = [
        SRC_DQN / "02_表格输出" / "multi_model_policy_comparison.csv",
        SRC_DQN / "02_表格输出" / "constraint_violation_summary.csv",
        SRC_DQN / "02_表格输出" / "reward_component_summary.csv",
        SRC_DQN / "02_表格输出" / "research_quality_gate_results.csv",
        SRC_DQN / "06_配置参数" / "dqn_revised_experimental_config.yaml",
        SRC_LOCAL / "08_代码快照" / "dqn_code_deep_explanation.md",
        SRC_LOCAL / "08_代码快照" / "dqn_code_to_model_setting_map.csv",
        SRC_LOCAL / "08_代码快照" / "dqn_code_to_outputs_map.csv",
        SRC_LOCAL / "08_代码快照" / "README_DQN代码总览.md",
        SRC_EXPLAIN / "04_报告输出" / "dqn_model_setting_detail_report.md",
        SRC_EXPLAIN / "04_报告输出" / "dqn_result_interpretation_report.md",
        SRC_EXPLAIN / "02_表格输出" / "dqn_model_component_literature_map.csv",
    ]
    for src in key_files:
        if src.exists():
            if "08_代码快照" in rel(src):
                dst = RUN / "08_代码快照" / src.name
            elif src.suffix.lower() in {".yaml", ".yml"}:
                dst = RUN / "02_表格输出" / src.name
            elif "04_报告输出" in rel(src):
                dst = RUN / "04_报告输出" / src.name
            else:
                dst = RUN / "02_表格输出" / src.name
            shutil.copy2(src, dst)
    for src in [
        SRC_DQN / "03_图表输出" / "dqn_revised_policy_comparison.png",
        SRC_DQN / "03_图表输出" / "dqn_revised_reward_curve.png",
        SRC_DQN / "03_图表输出" / "dqn_revised_moving_average_reward.png",
        SRC_EXPLAIN / "03_图表输出" / "dqn_revised_constraint_summary_explained.png",
        SRC_DQN / "03_图表输出" / "dqn_revised_action_distribution.png",
        SRC_DQN / "03_图表输出" / "multi_model_comparison.png",
    ]:
        if src.exists():
            shutil.copy2(src, RUN / "03_图表输出" / src.name)


def inspect_explanation_sprawl() -> pd.DataFrame:
    rows = []
    for pkg in [SRC_LOCAL, SRC_EXPLAIN]:
        for d in sorted([p for p in pkg.rglob("*") if p.is_dir()]):
            files = [p for p in d.iterdir() if p.is_file()]
            md_files = [p for p in files if p.suffix.lower() == ".md"]
            sidecars = [p for p in md_files if p.name.endswith(".explanation.md")]
            local = [p for p in md_files if p.name.endswith("_local.md") or "local" in p.name.lower()]
            if files:
                rows.append(
                    {
                        "package": rel(pkg),
                        "directory": rel(d),
                        "file_count": len(files),
                        "md_count": len(md_files),
                        "same_name_explanation_count": len(sidecars),
                        "local_md_count": len(local),
                        "diagnosis": "over_explained" if len(sidecars) > 3 or len(local) > 1 else "ok_or_review",
                    }
                )
    df = pd.DataFrame(rows)
    df.to_csv(RUN / "02_表格输出" / "explanation_sprawl_audit.csv", index=False, encoding="utf-8-sig")
    return df


def delete_path(path: Path, reason: str, rows: list[dict[str, object]]) -> None:
    if not path.exists():
        return
    size = 0
    if path.is_file():
        size = path.stat().st_size
        path.unlink()
        kind = "file"
    else:
        size = sum(p.stat().st_size for p in path.rglob("*") if p.is_file())
        shutil.rmtree(path)
        kind = "directory"
    rows.append({"deleted_path": rel(path), "kind": kind, "size_bytes": size, "reason": reason})


def clean_over_explanations_and_caches() -> pd.DataFrame:
    rows: list[dict[str, object]] = []

    # Reduce the 2317 package: keep directory READMEs and core DQN code docs/maps, delete bulk sidecars and duplicate local files.
    keep_md_names = {
        "README.md",
        "inputs.md",
        "README_表格解释.md",
        "README_图表解释.md",
        "README_报告解释.md",
        "README_模型输出解释.md",
        "README_参数配置解释.md",
        "README_代码解释.md",
        "README_DQN代码总览.md",
        "dqn_code_deep_explanation.md",
        "dqn_code_method_notes.md",
        "dqn_code_reproducibility_notes.md",
        "README_论文输出解释.md",
        "README_Results解释.md",
        "README_Word导出解释.md",
        "README_总索引说明.md",
        "artifact_explanation_index.md",
        "local_explanation_repair_report.md",
        "local_explanation_dry_run_report.md",
    }
    if SRC_LOCAL.exists():
        for p in sorted(SRC_LOCAL.rglob("*.md")):
            if p.name.endswith(".explanation.md") or p.name.endswith("_local.md") or p.name in {"figure_explanations_local.md", "table_explanations_local.md"}:
                delete_path(p, "过度生成的同名/本地解释侧车，目录 README 与核心 DQN 说明已足够", rows)
            elif p.name not in keep_md_names and any(part in rel(p) for part in ["02_表格输出", "03_图表输出", "04_报告输出", "05_模型与实验", "06_配置参数", "09_论文输出", "10_输出解释与索引"]):
                # Keep code deep docs and README, remove copied older broad explanation docs in curated local package.
                if "08_代码快照" not in rel(p):
                    delete_path(p, "重复的上一轮解释报告副本，保留本轮精简报告和核心源数据", rows)

    # Delete failed/intermediate packages that are superseded by 2056/2254/2317/current package.
    redundant_runs = [
        "20260426_2048_推荐缓存删除与DQN修正版训练",
        "20260426_2049_推荐缓存删除与DQN修正版训练",
        "20260426_2050_推荐缓存删除与DQN修正版训练",
        "20260426_2054_推荐缓存删除与DQN修正版训练",
        "20260426_2244_DQN输出复核_解释体系与论文输出升级",
        "20260426_2245_DQN输出复核_解释体系与论文输出升级",
        "20260426_2249_DQN输出复核_解释体系与论文输出升级",
        "20260426_2252_DQN输出复核_解释体系与论文输出升级",
    ]
    for name in redundant_runs:
        delete_path(ROOT / "outputs" / "工作包" / name, "已被后续成功 run 取代的失败/中间工作包", rows)

    # Python/runtime caches.
    for cache in ROOT.rglob("*"):
        if cache.is_dir() and cache.name in {"__pycache__", ".pytest_cache", ".mypy_cache", ".ruff_cache"}:
            delete_path(cache, "Python/runtime cache，可自动再生", rows)

    df = pd.DataFrame(rows)
    df.to_csv(RUN / "02_表格输出" / "deleted_redundant_outputs_log.csv", index=False, encoding="utf-8-sig")
    return df


def create_academic_word() -> Path:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    comparison = pd.read_csv(SRC_DQN / "02_表格输出" / "multi_model_policy_comparison.csv")
    installed_fonts = {f.name for f in fm.fontManager.ttflist}
    for font_name in ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "Arial Unicode MS"]:
        if font_name in installed_fonts:
            plt.rcParams["font.sans-serif"] = [font_name]
            break
    plt.rcParams["axes.unicode_minus"] = False
    evidence = comparison[["policy", "total_reward", "risk_reward_total", "information_gain_total", "sampling_cost_total", "constraint_violation_rate", "rank"]].copy()
    evidence.to_csv(RUN / "09_论文输出" / "04_结果" / "academic_results_evidence_table.csv", index=False, encoding="utf-8-sig")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("PEANUT 风险监管 DQN 修正版实验结果", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("experimental results draft；不构成 formal 监管政策结论")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 结果边界与分析目标", level=2)
    doc.add_paragraph(
        "本节报告 PEANUT/AFB1 风险监管场景下 DQN 修正版 experimental run 的结果。模型目标是在给定预算、容量约束和动作空间下，比较 DQN、Q-learning 与多类启发式策略的风险覆盖、信息增益、抽检成本和约束满足情况。由于 reward 权重、转移近似、预算与容量参数尚未完成 formal 确认，本文档仅作为论文 Results 草稿和后续 formal DQN 参数确认依据。"
    )

    doc.add_heading("2. 多模型策略比较", level=2)
    table_img_df = evidence.copy()
    short_names = {
        "threshold/greedy uncertainty": "threshold",
        "risk-ranking top-k": "risk top-k",
        "uniform allocation": "uniform",
        "historical allocation": "historical",
        "random policy": "random",
    }
    table_img_df["policy"] = table_img_df["policy"].map(lambda x: short_names.get(str(x), str(x)))
    table_img_df.columns = ["策略", "总奖励", "风险收益", "信息增益", "抽检成本", "违约率", "排名"]
    for col in ["总奖励", "风险收益", "信息增益", "抽检成本", "违约率"]:
        table_img_df[col] = table_img_df[col].map(lambda x: f"{float(x):.3f}")
    table_img_df["排名"] = table_img_df["排名"].map(lambda x: str(int(x)))
    fig, ax = plt.subplots(figsize=(11.2, 3.3), dpi=180)
    ax.axis("off")
    ax.set_title("表 1 多模型策略比较（experimental）", fontsize=13, pad=10)
    tbl = ax.table(cellText=table_img_df.values, colLabels=table_img_df.columns, loc="center", cellLoc="center")
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(8.5)
    tbl.scale(1, 1.25)
    for (row, _col), cell in tbl.get_celld().items():
        if row == 0:
            cell.set_facecolor("#3b6699")
            cell.set_text_props(color="white", weight="bold")
        elif row % 2 == 0:
            cell.set_facecolor("#f4f7fb")
    table_img = RUN / "09_论文输出" / "09_word导出" / "academic_policy_comparison_table.png"
    fig.savefig(table_img, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    doc.add_picture(str(table_img), width=Cm(16.2))
    best = comparison.iloc[0]
    dqn = comparison[comparison["policy"].astype(str).str.contains("DQN", na=False)].iloc[0]
    doc.add_paragraph(
        f"表 1 显示，在统一评价协议下，{best['policy']} 取得最高 total reward（{best['total_reward']:.3f}），DQN 修正版排名第 {int(dqn['rank'])}（total reward={dqn['total_reward']:.3f}）。该结果说明修正版 DQN 已能产生接近最优对照的策略，但当前设定下尚未超过 Q-learning。"
    )

    def add_fig(path: Path, caption: str, width: float = 14.8) -> None:
        if path.exists():
            doc.add_picture(str(path), width=Cm(width))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3. 图表结果", level=2)
    add_fig(RUN / "03_图表输出" / "dqn_revised_policy_comparison.png", "图 1 DQN 与 baseline 策略 total reward 对比。")
    add_fig(RUN / "03_图表输出" / "multi_model_comparison.png", "图 2 多模型比较的图形化摘要。")
    add_fig(RUN / "03_图表输出" / "dqn_revised_reward_curve.png", "图 3 DQN 修正版训练 reward 曲线。")
    add_fig(RUN / "03_图表输出" / "dqn_revised_moving_average_reward.png", "图 4 DQN 修正版移动平均 reward，用于观察训练稳定性。")
    add_fig(RUN / "03_图表输出" / "dqn_revised_action_distribution.png", "图 5 DQN 修正版动作分布，反映不同加码动作在状态集合上的使用情况。")
    add_fig(RUN / "03_图表输出" / "dqn_revised_constraint_summary_explained.png", "图 6 约束违约为 0 的解释性图，说明当前 experimental 约束下无违约，不是缺失数据。")

    doc.add_heading("4. 约束、收敛与解释限制", level=2)
    doc.add_paragraph(
        "所有比较策略的 constraint violation rate 均为 0，表明当前 action mask、monthly budget 与 capacity 设定下未发生违约。但该发现依赖当前 experimental 约束参数，不能替代 formal 参数确认。训练曲线和移动平均 reward 可用于判断 DQN 训练过程是否具有基本稳定趋势，但仍需敏感性分析、外部验证和更长训练或更细粒度动作空间验证。"
    )
    doc.add_paragraph(
        "从监管问题解释角度看，当前结果支持“风险收益—信息增益—抽检成本—约束满足”的可解释比较框架；但不支持直接给出最终监管资源配置建议。Q-learning 在本轮 total reward 上领先 DQN，提示当前状态聚合、reward 尺度或训练设置可能使表格型对照更稳健。"
    )

    doc.add_heading("5. Evidence 与文献状态", level=2)
    doc.add_paragraph(
        "本节所有数值均来自 multi_model_policy_comparison.csv 和相关质量门控表。文献依据已按 DQN 方法、constrained RL、risk-based inspection、MOE/EDI/AFB1 和 model reporting 建立候选映射；部分条目仍为 metadata/abstract-level 或浏览受限候选，不能作为 formal 文献证据。"
    )

    out = RUN / "09_论文输出" / "09_word导出" / "dqn_results_academic_with_figures.docx"
    doc.save(out)
    write(
        RUN / "09_论文输出" / "04_结果" / "dqn_results_academic_with_figures.md",
        """
        # PEANUT 风险监管 DQN 修正版实验结果

        本 Markdown 与 Word 同步，Word 中已写入多模型比较表和核心 PNG 图表。全部结果仍为 experimental results draft。
        """,
    )
    return out


def create_archive() -> Path:
    ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    zip_path = ARCHIVE_DIR / f"workflow1_curated_project_snapshot_{datetime.now():%Y%m%d_%H%M}.zip"
    include_paths = [
        "AGENTS.md",
        "README.md",
        "START_HERE.md",
        "pyproject.toml",
        "requirements.txt",
        "src",
        "tools",
        "workflow_recipes",
        "model_registry",
        "research_quality",
        "workflow_improvement",
        "project_state",
        "outputs/_index",
        "outputs/工作包/20260426_2056_推荐缓存删除与DQN修正版训练",
        "outputs/工作包/20260426_2317_输出解释就地化修正与DQN代码深度说明补强",
        rel(RUN),
        "experiments/optimization",
        "reports/项目级索引与摘要",
        "references/processed_summaries",
        "references/top_journal_benchmark",
    ]
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for item in include_paths:
            path = ROOT / item
            if not path.exists():
                continue
            if path.is_file():
                z.write(path, rel(path))
            else:
                for p in path.rglob("*"):
                    if p.is_file() and "__pycache__" not in p.parts:
                        z.write(p, rel(p))
    write(
        RUN / "04_报告输出" / "project_archive_report.md",
        f"""
        # Project Archive Report

        已创建精选项目快照：`{rel(zip_path)}`。

        存档包含代码、配置、project_state、索引、最新 DQN 修正版训练包、精简后的 DQN 代码说明包、本轮 Word 学术升级包、canonical optimization 输出和 processed references。

        为控制体积，未打包 `data/01_raw/` 原始数据、Zotero SQLite、失败/中间工作包和缓存目录。原始数据仍保留在本地项目中，未修改。
        """,
    )
    return zip_path


def update_indexes(deleted: pd.DataFrame, archive_path: Path, docx_path: Path) -> None:
    ts = datetime.now().isoformat(timespec="seconds")
    run_index = ROOT / "outputs" / "_index" / "run_index.md"
    with run_index.open("a", encoding="utf-8") as f:
        f.write(
            f"\n## {RUN.name}\n\n- 时间：{ts}\n- 类型：输出解释精简、项目存档、缓存清理与 Word 学术升级\n- 路径：`{rel(RUN)}`\n- 存档：`{rel(archive_path)}`\n- Word：`{rel(docx_path)}`\n- 删除项：{len(deleted)}\n"
        )
    manifest = ROOT / "outputs" / "_index" / "run_manifest.csv"
    row = pd.DataFrame([{"run_id": RUN.name, "run_path": rel(RUN), "created_at": ts, "task": "输出解释精简_项目存档与Word学术升级", "status": "completed"}])
    if manifest.exists():
        old = pd.read_csv(manifest)
        pd.concat([old, row], ignore_index=True).to_csv(manifest, index=False, encoding="utf-8-sig")
    else:
        row.to_csv(manifest, index=False, encoding="utf-8-sig")
    latest_path = ROOT / "outputs" / "_index" / "latest_canonical_outputs.yaml"
    latest = yaml.safe_load(latest_path.read_text(encoding="utf-8")) if latest_path.exists() else {}
    latest = latest or {}
    latest.update(
        {
            "latest_cleanup_archive_word_upgrade_run": rel(RUN),
            "latest_project_archive_zip": rel(archive_path),
            "latest_academic_results_docx_with_figures": rel(docx_path),
            "latest_dqn_status": "experimental_audited_curated_not_formal",
        }
    )
    latest_path.write_text(yaml.safe_dump(latest, allow_unicode=True, sort_keys=False), encoding="utf-8")
    for file, addition in {
        "project_state/current_focus.md": f"\n\n## 输出精简与存档\n\n已精简过度解释文件、清理缓存和失败中间包，并生成学术化带图表 Word。任务包：`{rel(RUN)}`。\n",
        "project_state/next_step.md": "\n\n下一步建议：如果继续论文，请基于最新 `dqn_results_academic_with_figures.docx` 生成 Method/Discussion；如果继续 formal DQN，请先确认参数表。\n",
        "project_state/changelog.md": f"\n\n## {datetime.now():%Y-%m-%d}\n\n- 精简冗余解释文件，清理缓存/失败中间包，创建项目精选存档，生成含表和核心图的学术 Results Word。任务包：`{rel(RUN)}`。\n",
        "project_state/decision_log.md": "\n\n## Decision: Explanation Minimalism\n\n不再为每个辅助文件机械生成同名 explanation。保留目录 README、核心 DQN 说明、代码映射、关键报告和 Word 说明即可，避免阅读混淆。\n",
        "project_state/conversation_handoff.md": f"\n\n最新处理：`{rel(RUN)}` 已精简解释、存档项目、清理缓存和中间包，并生成学术化带图表 Word。DQN 仍为 experimental。\n",
    }.items():
        p = ROOT / file
        p.write_text(p.read_text(encoding="utf-8", errors="replace") + addition, encoding="utf-8")


def write_reports(sprawl: pd.DataFrame, deleted: pd.DataFrame, archive_path: Path, docx_path: Path) -> None:
    write(
        RUN / "04_报告输出" / "explanation_simplification_report.md",
        f"""
        # Explanation Simplification Report

        ## 为什么 08_代码快照和其他目录有很多 md

        上一轮为了满足“就地解释”和“同名 explanation”要求，机械地给几乎每个 CSV/PNG/MD/DOCX 生成 `.explanation.md`，还给每个目录生成 `_local.md`。这导致解释文件数量膨胀，读者需要在说明文件之间来回跳转，反而降低可读性。

        ## 本轮处理原则

        - 保留目录 README：让用户在目录内知道看什么。
        - 保留核心 DQN 代码说明：`README_DQN代码总览.md`、`dqn_code_deep_explanation.md`、`dqn_code_to_model_setting_map.csv`、`dqn_code_to_outputs_map.csv`。
        - 删除批量生成的同名 `.explanation.md` 和重复 `_local.md`。
        - 不删除核心数据、最新 DQN 训练包、project_state、raw data 或 canonical outputs。

        ## 删除概况

        - 删除记录数：{len(deleted)}
        - 释放空间：{deleted['size_bytes'].sum() / 1024 / 1024 if not deleted.empty else 0:.2f} MB
        - 删除日志：`02_表格输出/deleted_redundant_outputs_log.csv`

        ## Word 升级

        新 Word：`{rel(docx_path)}`，包含多模型比较表和核心 PNG 图表，写法调整为学术 Results 草稿，但仍明确标注 experimental。

        ## 项目存档

        存档：`{rel(archive_path)}`。
        """,
    )
    write(
        RUN / "README.md",
        f"""
        # 输出解释精简、项目存档与 Word 学术升级

        本轮完成：

        - 检查解释文件膨胀原因；
        - 删除过度生成的解释侧车和失败/中间 run 包；
        - 清理 Python 缓存；
        - 生成包含表格和核心图表的学术 Results Word；
        - 创建项目精选存档。

        核心输出：

        - `04_报告输出/explanation_simplification_report.md`
        - `02_表格输出/explanation_sprawl_audit.csv`
        - `02_表格输出/deleted_redundant_outputs_log.csv`
        - `09_论文输出/09_word导出/dqn_results_academic_with_figures.docx`
        - `{rel(archive_path)}`
        """,
    )


def manifest() -> None:
    rows = []
    for p in sorted(RUN.rglob("*")):
        if p.is_file():
            rows.append({"path": rel(p), "size_bytes": p.stat().st_size, "sha256": sha(p) if p.stat().st_size else ""})
    pd.DataFrame(rows).to_csv(RUN / "manifest.csv", index=False, encoding="utf-8-sig")


def main() -> None:
    init_run()
    copy_key_inputs()
    sprawl = inspect_explanation_sprawl()
    docx = create_academic_word()
    deleted = clean_over_explanations_and_caches()
    archive_path = create_archive()
    write_reports(sprawl, deleted, archive_path, docx)
    update_indexes(deleted, archive_path, docx)
    manifest()
    print(json.dumps({"run_dir": str(RUN), "archive": str(archive_path), "docx": str(docx), "deleted_count": int(len(deleted))}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
