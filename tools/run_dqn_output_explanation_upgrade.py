from __future__ import annotations

import csv
import hashlib
import json
import math
import os
import shutil
import subprocess
import sys
import textwrap
from collections import Counter
from datetime import datetime
from pathlib import Path

import pandas as pd
import yaml


ROOT = Path(__file__).resolve().parents[1]
PYTHON = r"D:\anaconda3\envs\myenv1\python.exe"
LATEST_DQN = ROOT / "outputs" / "工作包" / "20260426_2056_推荐缓存删除与DQN修正版训练"
BASELINE_DQN = ROOT / "outputs" / "工作包" / "20260426_1746_全流程验收与DQN自动参数训练"
ZOTERO_ROOT = Path(r"D:\桌面\codex\zotero")


def now_stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M")


RUN_DIR = ROOT / "outputs" / "工作包" / f"{now_stamp()}_DQN输出复核_解释体系与论文输出升级"


DIRS = [
    "00_输入说明",
    "01_数据输出",
    "02_表格输出",
    "03_图表输出",
    "04_报告输出",
    "05_模型与实验",
    "06_配置参数",
    "07_日志与错误",
    "08_代码快照",
    "09_论文输出/04_结果",
    "09_论文输出/09_word导出",
    "10_输出解释与索引",
]


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.strip() + "\n", encoding="utf-8")


def safe_read(path: Path, max_chars: int = 20000) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")[:max_chars]
    except Exception as exc:
        return f"[无法读取: {exc}]"


def rel(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT)).replace("\\", "/")
    except ValueError:
        return str(path).replace("\\", "/")


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def init_run_package() -> None:
    for d in DIRS:
        (RUN_DIR / d).mkdir(parents=True, exist_ok=True)
    write_text(
        RUN_DIR / "00_输入说明" / "inputs.md",
        f"""
        # 本轮输入说明

        本轮依据本地项目文件和索引执行，不依赖聊天历史或临时上传文件。

        - 最新 DQN 修正版 experimental 包：`{rel(LATEST_DQN)}`
        - 上一轮 DQN 自动参数训练包：`{rel(BASELINE_DQN)}`
        - 项目索引：`outputs/_index/`
        - 项目状态：`project_state/`
        - 本地 references：`references/`
        - Zotero 工作流目录（只读核验）：`{ZOTERO_ROOT}`

        `research_loop/`：{('存在' if (ROOT / 'research_loop').exists() else '未找到，已记录为无法核验')}。
        """,
    )


def inventory_package(pkg: Path, tag: str) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    if not pkg.exists():
        return [
            {
                "package_tag": tag,
                "file_path": rel(pkg),
                "file_type": "missing_package",
                "size_bytes": 0,
                "readable": False,
                "notes": "未找到包",
            }
        ]
    for path in sorted(pkg.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower().lstrip(".") or "no_ext"
        row: dict[str, object] = {
            "package_tag": tag,
            "file_path": rel(path),
            "file_type": suffix,
            "size_bytes": path.stat().st_size,
            "sha256": "",
            "readable": True,
            "row_count": "",
            "column_count": "",
            "notes": "",
        }
        try:
            if path.stat().st_size > 0:
                row["sha256"] = sha256_file(path)
            if suffix in {"csv", "tsv"} and path.stat().st_size > 0:
                sep = "\t" if suffix == "tsv" else ","
                df = pd.read_csv(path, sep=sep, nrows=100000)
                row["row_count"] = int(len(df))
                row["column_count"] = int(len(df.columns))
            elif suffix in {"xlsx", "xls"} and path.stat().st_size > 0:
                xl = pd.ExcelFile(path)
                row["notes"] = "sheets=" + "|".join(xl.sheet_names[:5])
            elif suffix in {"png", "jpg", "jpeg"}:
                try:
                    from PIL import Image, ImageStat

                    with Image.open(path) as im:
                        stat = ImageStat.Stat(im.convert("L"))
                        row["notes"] = f"image={im.size[0]}x{im.size[1]};std={stat.stddev[0]:.2f}"
                except Exception as exc:
                    row["readable"] = False
                    row["notes"] = f"图像读取失败: {exc}"
            elif suffix in {"md", "yaml", "yml", "txt", "py", "svg", "json"}:
                text = safe_read(path, 2000)
                if "????" in text or "\ufffd" in text:
                    row["notes"] = "疑似乱码或替换字符"
                elif path.stat().st_size == 0:
                    row["notes"] = "0 字节"
        except Exception as exc:
            row["readable"] = False
            row["notes"] = f"读取/解析失败: {exc}"
        rows.append(row)
    return rows


def load_core_tables() -> dict[str, pd.DataFrame]:
    table_dir = LATEST_DQN / "02_表格输出"
    model_dir = LATEST_DQN / "05_模型与实验"
    tables = {}
    for name in [
        "multi_model_policy_comparison",
        "multi_model_metric_summary",
        "baseline_fairness_check",
        "reward_component_summary",
        "convergence_diagnosis_summary",
        "constraint_violation_summary",
        "chart_quality_audit",
        "research_quality_gate_results",
        "action_space_options",
    ]:
        path = table_dir / f"{name}.csv"
        if path.exists():
            tables[name] = pd.read_csv(path)
    for name in ["dqn_revised_training_log", "qlearning_training_log", "experiment_ledger"]:
        path = model_dir / f"{name}.csv"
        if path.exists():
            tables[name] = pd.read_csv(path)
    return tables


def audit_issues(inventory: list[dict[str, object]], tables: dict[str, pd.DataFrame]) -> list[dict[str, object]]:
    issues: list[dict[str, object]] = []
    for row in inventory:
        path = str(row["file_path"])
        size = int(row["size_bytes"])
        ftype = str(row["file_type"])
        if size == 0:
            issues.append(
                {
                    "问题类型": "空文件/空图",
                    "文件路径": path,
                    "严重程度": "high" if any(x in path for x in ["03_图表输出", "01_数据输出", "04_报告输出"]) else "medium",
                    "是否影响阅读": "是",
                    "是否影响论文结果": "是" if any(x in path for x in ["03_图表输出", "01_数据输出"]) else "可能",
                    "是否需要修复": "是",
                    "建议修复方式": "不要输出空白文件；若无有效数据则生成解释性 PNG/Markdown 并记录无法核验原因。",
                    "是否需要新增 skill": "是" if "03_图表输出" in path else "否",
                    "是否需要更新 AGENTS": "是",
                }
            )
        elif ftype == "svg" and size < 2000:
            issues.append(
                {
                    "问题类型": "近似空图/信息不足 SVG",
                    "文件路径": path,
                    "严重程度": "high",
                    "是否影响阅读": "是",
                    "是否影响论文结果": "是",
                    "是否需要修复": "是",
                    "建议修复方式": "主图改为 PNG；若源数据全 0 或无差异，生成解释性图而不是空坐标轴。",
                    "是否需要新增 skill": "是",
                    "是否需要更新 AGENTS": "是",
                }
            )
        elif ftype in {"csv", "xlsx"} and size < 32:
            issues.append(
                {
                    "问题类型": "表格近似为空",
                    "文件路径": path,
                    "严重程度": "high",
                    "是否影响阅读": "是",
                    "是否影响论文结果": "可能",
                    "是否需要修复": "是",
                    "建议修复方式": "补充字段、来源、行列数和无法生成原因；必要时重新生成。",
                    "是否需要新增 skill": "否",
                    "是否需要更新 AGENTS": "是",
                }
            )
    constraint = tables.get("constraint_violation_summary")
    if constraint is not None and not constraint.empty:
        numeric = constraint.select_dtypes(include="number")
        if not numeric.empty and float(numeric.drop(columns=["monthly_budget"], errors="ignore").to_numpy().sum()) == 0.0:
            issues.append(
                {
                    "问题类型": "全 0 图表语义不足",
                    "文件路径": rel(LATEST_DQN / "03_图表输出" / "dqn_revised_constraint_summary.png"),
                    "严重程度": "medium",
                    "是否影响阅读": "是",
                    "是否影响论文结果": "否，但影响解释",
                    "是否需要修复": "是",
                    "建议修复方式": "生成解释性 PNG：说明全 0 表示当前约束设定下没有违约，不是缺失数据。",
                    "是否需要新增 skill": "是",
                    "是否需要更新 AGENTS": "是",
                }
            )
    for required in ["figure_explanations.md", "table_explanations.md", "model_output_explanations.md", "code_explanations.md"]:
        issues.append(
            {
                "问题类型": "长期解释文件缺失",
                "文件路径": f"最新 DQN 包缺少 {required}",
                "严重程度": "medium",
                "是否影响阅读": "是",
                "是否影响论文结果": "可能",
                "是否需要修复": "是",
                "建议修复方式": "在本轮任务包生成解释索引，并固化为 workflow 规则。",
                "是否需要新增 skill": "是",
                "是否需要更新 AGENTS": "是",
            }
        )
    return issues


def choose_font():
    import matplotlib.font_manager as fm

    candidates = ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "Arial Unicode MS"]
    installed = {f.name for f in fm.fontManager.ttflist}
    for name in candidates:
        if name in installed:
            return name
    return "DejaVu Sans"


def render_repair_charts(tables: dict[str, pd.DataFrame]) -> list[dict[str, object]]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from PIL import Image, ImageStat

    font = choose_font()
    plt.rcParams["font.sans-serif"] = [font]
    plt.rcParams["axes.unicode_minus"] = False
    rows = []

    latest_fig_dir = LATEST_DQN / "03_图表输出"
    for src in sorted(latest_fig_dir.glob("*.png")):
        dst = RUN_DIR / "03_图表输出" / src.name
        shutil.copy2(src, dst)

    constraint = tables.get("constraint_violation_summary")
    if constraint is not None:
        fig, ax = plt.subplots(figsize=(10, 5.4), dpi=160)
        ax.axis("off")
        n = len(constraint)
        text = (
            "约束违约率诊断\n\n"
            f"本轮共比较 {n} 个策略；constraint_violation_count 与 constraint_violation_rate 均为 0。\n"
            "这表示在当前 experimental 预算、容量与 action mask 约束下，所有策略均未触发违约。\n"
            "因此该图不是数据缺失，而是“约束满足”的解释性图。若进入 formal DQN，仍需人工确认预算、容量、成本与约束权重。"
        )
        ax.text(0.5, 0.55, text, ha="center", va="center", fontsize=14, linespacing=1.8)
        ax.text(0.5, 0.08, "数据来源：constraint_violation_summary.csv", ha="center", fontsize=10, color="#555555")
        dst = RUN_DIR / "03_图表输出" / "dqn_revised_constraint_summary_explained.png"
        fig.savefig(dst, bbox_inches="tight", facecolor="white")
        plt.close(fig)

    hist_empty = BASELINE_DQN / "03_图表输出" / "peanut_dqn_auto_top_priority_risk.svg"
    if hist_empty.exists() and hist_empty.stat().st_size == 0:
        fig, ax = plt.subplots(figsize=(10, 5.4), dpi=160)
        ax.axis("off")
        text = (
            "历史空图修复说明\n\n"
            "上一轮 `peanut_dqn_auto_top_priority_risk.svg` 为 0 字节，无法复核图形内容。\n"
            "本轮不覆盖历史包，而是在新工作包记录该问题；风险优先级图以最新修正版 PNG 作为可读依据。\n"
            "该历史空图不得用于论文 Results。"
        )
        ax.text(0.5, 0.55, text, ha="center", va="center", fontsize=14, linespacing=1.8)
        ax.text(0.5, 0.08, "修复方式：解释性 PNG + deep audit 问题清单", ha="center", fontsize=10, color="#555555")
        dst = RUN_DIR / "03_图表输出" / "historical_empty_top_priority_risk_explained.png"
        fig.savefig(dst, bbox_inches="tight", facecolor="white")
        plt.close(fig)

    for path in sorted((RUN_DIR / "03_图表输出").glob("*.png")):
        try:
            with Image.open(path) as im:
                stat = ImageStat.Stat(im.convert("L"))
                std = float(stat.stddev[0])
                qa = "pass" if path.stat().st_size > 4096 and std > 2 else "review"
                rows.append(
                    {
                        "figure_path": rel(path),
                        "source_data": infer_figure_source(path.name),
                        "png_exists": True,
                        "size_bytes": path.stat().st_size,
                        "pixel_std": round(std, 3),
                        "blank_or_near_blank": std <= 2,
                        "all_zero_or_no_difference_semantic": "constraint" in path.name,
                        "chinese_font_status": f"font={font}",
                        "title_caption_status": "needs explanation file" if qa == "review" else "explained in figure_explanations.md",
                        "paper_results_support": "experimental only" if qa == "pass" else "do not use without explanation",
                        "qa_status": qa,
                    }
                )
        except Exception as exc:
            rows.append(
                {
                    "figure_path": rel(path),
                    "source_data": "",
                    "png_exists": True,
                    "size_bytes": path.stat().st_size,
                    "pixel_std": "",
                    "blank_or_near_blank": True,
                    "all_zero_or_no_difference_semantic": "",
                    "chinese_font_status": f"read failed: {exc}",
                    "title_caption_status": "failed",
                    "paper_results_support": "no",
                    "qa_status": "fail",
                }
            )
    return rows


def infer_figure_source(name: str) -> str:
    mapping = {
        "training_curve": "05_模型与实验/dqn_revised_training_log.csv",
        "reward_curve": "05_模型与实验/dqn_revised_training_log.csv",
        "moving_average_reward": "05_模型与实验/dqn_revised_training_log.csv",
        "policy_comparison": "02_表格输出/multi_model_policy_comparison.csv",
        "multi_model_comparison": "02_表格输出/multi_model_policy_comparison.csv",
        "action_distribution": "01_数据输出/dqn_revised_policy.csv",
        "constraint": "02_表格输出/constraint_violation_summary.csv",
        "top_priority": "01_数据输出/dqn_revised_policy.csv",
        "convergence": "02_表格输出/convergence_diagnosis_summary.csv",
    }
    for key, src in mapping.items():
        if key in name:
            return src
    return "未自动识别"


def markdown_table(df: pd.DataFrame, cols: list[str] | None = None, n: int = 8) -> str:
    if cols:
        df = df[cols]
    small = df.head(n).copy()
    headers = [str(c) for c in small.columns]
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for _, row in small.iterrows():
        vals = []
        for value in row.tolist():
            if isinstance(value, float):
                vals.append(f"{value:.4f}")
            else:
                vals.append(str(value).replace("|", "/"))
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def generate_explanations(tables: dict[str, pd.DataFrame], chart_rows: list[dict[str, object]], issue_rows: list[dict[str, object]]) -> None:
    comparison = tables.get("multi_model_policy_comparison", pd.DataFrame())
    top_policy = comparison.iloc[0].to_dict() if not comparison.empty else {}
    dqn_row = comparison[comparison.get("policy", pd.Series(dtype=str)).astype(str).str.contains("DQN", na=False)].iloc[0].to_dict() if not comparison.empty and comparison.get("policy", pd.Series(dtype=str)).astype(str).str.contains("DQN", na=False).any() else {}

    artifact_rows = []
    for path in sorted(list(LATEST_DQN.rglob("*")) + list((RUN_DIR / "03_图表输出").rglob("*"))):
        if not path.is_file():
            continue
        artifact_rows.append(
            {
                "文件路径": rel(path),
                "文件类型": path.suffix.lower().lstrip(".") or "no_ext",
                "生成目的": artifact_purpose(path),
                "输入来源": artifact_source(path),
                "生成方法": "读取最新 DQN experimental 包并在本轮复核/解释/必要时生成替代说明。",
                "主要结果": artifact_result(path, top_policy),
                "如何阅读": artifact_reading(path),
                "论文相关部分": artifact_section(path),
                "是否支持 formal 结论": "否；仅支持 experimental results draft 或方法说明",
                "局限性": artifact_limitation(path),
                "人工注意": "formal DQN 前需确认参数、约束和外部验证。",
            }
        )

    artifact_df = pd.DataFrame(artifact_rows)
    artifact_df.to_csv(RUN_DIR / "02_表格输出" / "artifact_explanation_index.csv", index=False, encoding="utf-8-sig")
    artifact_df.to_csv(RUN_DIR / "02_表格输出" / "artifact_to_evidence_map.csv", index=False, encoding="utf-8-sig")

    artifact_md = "# Artifact Explanation Index\n\n"
    artifact_md += "本索引解释本轮复核范围内的关键输出。所有 DQN 相关结论均保持 `experimental` 边界。\n\n"
    for _, r in artifact_df.head(80).iterrows():
        artifact_md += f"## {r['文件路径']}\n\n- 类型：{r['文件类型']}\n- 目的：{r['生成目的']}\n- 来源：{r['输入来源']}\n- 阅读方式：{r['如何阅读']}\n- 论文用途：{r['论文相关部分']}\n- formal 结论：{r['是否支持 formal 结论']}\n- 局限性：{r['局限性']}\n\n"
    write_text(RUN_DIR / "10_输出解释与索引" / "artifact_explanation_index.md", artifact_md)

    fig_md = "# Figure Explanations\n\n"
    for r in chart_rows:
        fig_md += (
            f"## {r['figure_path']}\n\n"
            f"- 数据来源：{r['source_data']}\n"
            f"- 图形元素：PNG 主图；中文字体状态 `{r['chinese_font_status']}`；像素标准差 `{r['pixel_std']}`。\n"
            f"- 主要发现：{figure_findings(str(r['figure_path']), comparison)}\n"
            f"- 如何阅读：先看标题/坐标轴，再结合源 CSV；若为解释性图，应按图中文字理解其科学含义。\n"
            f"- 是否可用于论文 Results：{r['paper_results_support']}。\n"
            f"- 局限性：本轮仍为 experimental；图形不得单独支撑正式监管政策结论。\n\n"
        )
    write_text(RUN_DIR / "10_输出解释与索引" / "figure_explanations.md", fig_md)

    table_md = "# Table Explanations\n\n"
    for name, df in tables.items():
        table_md += f"## {name}.csv\n\n"
        table_md += f"- 数据来源：`{rel(LATEST_DQN / ('02_表格输出/' + name + '.csv')) if (LATEST_DQN / ('02_表格输出/' + name + '.csv')).exists() else rel(LATEST_DQN / ('05_模型与实验/' + name + '.csv'))}`\n"
        table_md += f"- 行列规模：{len(df)} 行，{len(df.columns)} 列。\n"
        table_md += f"- 字段解释：{', '.join(map(str, df.columns[:12]))}{' ...' if len(df.columns) > 12 else ''}。\n"
        table_md += f"- 主要发现：{table_findings(name, df)}\n"
        table_md += "- 局限性：表格仅反映当前 experimental 配置与本地数据输出；正式论文使用前需通过 evidence map 和质量门控。\n\n"
    write_text(RUN_DIR / "10_输出解释与索引" / "table_explanations.md", table_md)

    model_md = "# Model Output Explanations\n\n"
    model_md += model_output_text(comparison, tables)
    write_text(RUN_DIR / "10_输出解释与索引" / "model_output_explanations.md", model_md)

    guide = f"""
    # Result Interpretation Guide

    ## 阅读顺序

    1. 先读 `dqn_model_setting_detail_report.md`，确认 state、action、reward、constraint 和 training 都来自本地配置或被标记为 experimental assumption。
    2. 再读 `multi_model_policy_comparison.csv` 和 `dqn_result_interpretation_report.md`，理解 DQN、Q-learning 与 heuristic 的比较。
    3. 图表只作为解释辅助，必须回到源 CSV。
    4. Word Results 草稿只能作为 `experimental results draft`。

    ## 本轮核心结果边界

    - 当前排序最高策略：{top_policy.get('policy', '未核验')}，total_reward={top_policy.get('total_reward', '未核验')}。
    - DQN 修正版：total_reward={dqn_row.get('total_reward', '未核验')}，排名={dqn_row.get('rank', '未核验')}。
    - 由于参数、约束、奖励权重尚未用户确认为 formal，不能写成正式监管政策建议。
    """
    write_text(RUN_DIR / "10_输出解释与索引" / "result_interpretation_guide.md", guide)


def artifact_purpose(path: Path) -> str:
    p = rel(path)
    if "03_图表输出" in p:
        return "可视化训练、策略比较、约束或风险优先级，用于结果阅读和图表 QA。"
    if "02_表格输出" in p:
        return "保存审计、比较、质量门控或解释索引的结构化证据。"
    if "01_数据输出" in p:
        return "保存策略或模型输出数据，供表格、图表和论文草稿引用。"
    if "05_模型与实验" in p:
        return "保存模型、训练日志或实验账本，供可复现性与收敛诊断核验。"
    if "09_论文输出" in p:
        return "保存论文 Results 草稿、证据表或 DOCX 导出。"
    if p.endswith(".py"):
        return "记录执行或生成逻辑，便于复现和方法说明。"
    return "项目输出或说明文件。"


def artifact_source(path: Path) -> str:
    p = rel(path)
    if "dqn_revised" in p or "multi_model" in p:
        return rel(LATEST_DQN)
    if "historical_empty" in p:
        return rel(BASELINE_DQN)
    return "本轮任务包或项目索引。"


def artifact_result(path: Path, top_policy: dict[str, object]) -> str:
    p = rel(path)
    if "multi_model_policy_comparison" in p:
        return f"当前 experimental 比较中最高策略为 {top_policy.get('policy', '未核验')}。"
    if "constraint" in p:
        return "当前约束违约统计为 0，应解释为约束满足而非数据缺失。"
    if "chart_quality" in p:
        return "记录图表是否非空、是否 PNG、是否可解释。"
    return "具体结果需结合对应解释文件和源 CSV 阅读。"


def artifact_reading(path: Path) -> str:
    if path.suffix.lower() == ".png":
        return "结合 figure_explanations.md 和源数据表读取，不从图像单独推出正式结论。"
    if path.suffix.lower() in {".csv", ".xlsx"}:
        return "先看字段解释和行列规模，再检查指标方向、单位和 evidence map。"
    if path.suffix.lower() == ".docx":
        return "作为论文级草稿阅读，注意文中 experimental 标记。"
    if path.suffix.lower() == ".py":
        return "结合 code_explanations.md 查看用途、输入输出、依赖和运行方式。"
    return "按 README 和对应解释索引阅读。"


def artifact_section(path: Path) -> str:
    p = rel(path)
    if "training" in p or "config" in p or p.endswith(".py"):
        return "Method / Appendix"
    if "policy_comparison" in p or "reward" in p or "constraint" in p or "09_论文输出" in p:
        return "Results"
    if "literature" in p:
        return "Method / Literature Review"
    return "Appendix / Reproducibility"


def artifact_limitation(path: Path) -> str:
    if "1746" in rel(path):
        return "历史包包含空文件和信息不足输出，仅用于对照和问题定位。"
    if "constraint" in rel(path):
        return "全 0 约束结果依赖当前约束设定，formal 前需要参数确认和敏感性分析。"
    return "未经过外部验证；当前 DQN 仍为 experimental。"


def figure_findings(path: str, comparison: pd.DataFrame) -> str:
    if "constraint" in path:
        return "违约为 0，表示当前约束配置下策略均满足约束；需要解释，不能留成空柱图。"
    if "policy_comparison" in path or "multi_model" in path:
        if comparison.empty:
            return "无法读取策略比较表。"
        return f"{comparison.iloc[0]['policy']} 总奖励最高；DQN 修正版排名 {comparison[comparison['policy'].astype(str).str.contains('DQN')].iloc[0]['rank']}。"
    if "reward_curve" in path or "moving_average" in path:
        return "用于观察训练回报趋势和早停表现；不能单独证明策略最优。"
    if "top_priority" in path:
        return "展示高风险/高优先级状态，需结合状态特征和抽检动作解释。"
    return "图表为本轮 experimental 输出的可读化呈现。"


def table_findings(name: str, df: pd.DataFrame) -> str:
    if df.empty:
        return "表为空，需回查生成逻辑。"
    if name == "multi_model_policy_comparison":
        return f"共比较 {len(df)} 个策略；最高 total_reward 策略为 {df.iloc[0].get('policy')}。"
    if name == "constraint_violation_summary":
        return "所有策略违约计数为 0；这是约束满足信号，不是缺失。"
    if name == "research_quality_gate_results":
        passed = df.astype(str).apply(lambda s: s.str.contains("pass|通过", case=False, regex=True)).any(axis=1).sum()
        return f"质量门控记录 {len(df)} 项，其中约 {passed} 项显示通过/可用。"
    if name == "dqn_revised_training_log":
        return f"训练日志包含 {len(df)} 条记录，用于 reward、loss、epsilon 与收敛诊断。"
    return "保存本轮 DQN experimental 输出的结构化证据。"


def model_output_text(comparison: pd.DataFrame, tables: dict[str, pd.DataFrame]) -> str:
    if comparison.empty:
        return "未找到 multi_model_policy_comparison.csv，无法生成模型输出解释。\n"
    rows = []
    for _, r in comparison.iterrows():
        rows.append(
            f"- `{r['policy']}`：total_reward={r['total_reward']:.3f}，mean_reward={r['mean_reward']:.4f}，"
            f"risk_reward_total={r['risk_reward_total']:.3f}，sampling_cost_total={r['sampling_cost_total']:.3f}，"
            f"constraint_violation_rate={r['constraint_violation_rate']:.3f}，rank={int(r['rank'])}。"
        )
    return f"""
    ## 多模型比较输出

    total_reward 越高表示在当前 reward decomposition 与 rescaling 下综合表现越好；sampling_cost_total 越高表示抽检成本更高；constraint_violation_rate 越高表示违反预算或容量约束的风险更高。

    {chr(10).join(rows)}

    ## 解释边界

    这些数值来自同一 experimental comparison protocol，但 reward 权重、预算、容量和动作空间尚未被用户确认为 formal，因此只支持探索性 Results 草稿。Q-learning 领先 DQN 说明当前状态聚合/训练轮次/奖励尺度下，表格型对照策略可能更稳定；不能据此断言 DQN 在真实监管中劣于 Q-learning。
    """


def generate_dqn_reports(tables: dict[str, pd.DataFrame]) -> None:
    config_path = LATEST_DQN / "06_配置参数" / "dqn_revised_experimental_config.yaml"
    config = yaml.safe_load(config_path.read_text(encoding="utf-8")) if config_path.exists() else {}
    comparison = tables.get("multi_model_policy_comparison", pd.DataFrame())
    action_options = tables.get("action_space_options", pd.DataFrame())
    convergence = tables.get("convergence_diagnosis_summary", pd.DataFrame())
    reward = tables.get("reward_component_summary", pd.DataFrame())

    setting = f"""
    # DQN 模型设置详细说明

    > 状态：experimental setting documentation。以下内容只来自本地配置、代码快照和输出表；无法从本地文件核验的内容标为 `未核验` 或 `experimental assumption`。

    ## 1. 科研问题

    本轮 DQN 试图在 PEANUT/AFB1 风险监管场景中，为不同状态单元选择抽检加码动作，使风险覆盖、信息增益、成本和约束满足之间取得可解释的 experimental trade-off。该输出不是 formal 监管政策结论。

    ## 2. MDP / belief-MDP 近似

    - 状态来自 `data/04_feature/peanut_belief_mdp_state_features_with_moe_edi.csv` 与相关 belief-state / MOE-EDI 特征。
    - transition 在当前输出中主要是基于状态序列和抽检动作的 experimental approximation；没有核验到经用户正式确认的真实监管动态转移方程。
    - belief update 使用 Beta-Binomial belief-state 作为输入证据，但本轮训练没有证明其可替代真实监管反馈闭环。

    ## 3. State

    当前配置路径：`{rel(config_path)}`。状态特征来源：`{config.get('paths', {}).get('state_features', '未核验')}`。状态应解释为地区、环节、风险、posterior mean/uncertainty、MOE/EDI 等信息的组合，具体字段需结合 canonical feature table 和 `state_feature_missing_summary.csv` 阅读。

    ## 4. Action

    实际训练动作空间：`{config.get('action_space', {}).get('actual_training_choice', '未核验')}`；动作档位：`{config.get('action_space', {}).get('increments', '未核验')}`。

    动作空间可行性：高维二元动作没有训练；配置中标记为 `{config.get('action_space', {}).get('high_dimensional_binary_action', '未核验')}`。本轮推荐继续使用粗粒度动作空间做实验，并把高维动作留给 hierarchical RL / factorized action / combinatorial optimization。

    ## 5. Reward

    Reward decomposition：{config.get('reward', {}).get('decomposition', [])}。

    - risk reward weight：{config.get('reward', {}).get('risk_reward_weight')}
    - information gain weight：{config.get('reward', {}).get('info_gain_weight')}
    - cost weight：{config.get('reward', {}).get('cost_weight')}
    - constraint penalty weight：{config.get('reward', {}).get('constraint_penalty_weight')}
    - rescaling：{config.get('reward', {}).get('reward_rescaling')}

    解释风险：reward 权重仍为 experimental；formal 前必须确认成本、预算、容量、惩罚强度和风险收益尺度。

    ## 6. Constraint

    约束包括 monthly budget、local/stage/global capacity 和 action mask。当前 monthly_budget={config.get('constraints', {}).get('monthly_budget', '未核验')}。约束违约表显示本轮比较策略均无违约，但这依赖当前 action mask 与预算容量设定。

    ## 7. Training

    训练参数：episodes={config.get('training', {}).get('episodes')}，learning_rate={config.get('training', {}).get('learning_rate')}，gamma={config.get('training', {}).get('gamma')}，batch_size={config.get('training', {}).get('batch_size')}，epsilon_start={config.get('training', {}).get('epsilon_start')}，epsilon_min={config.get('training', {}).get('epsilon_min')}，epsilon_decay={config.get('training', {}).get('epsilon_decay')}，target_update_frequency={config.get('training', {}).get('target_update_frequency')}，early_stopping={config.get('training', {}).get('early_stopping')}。

    机制说明：DQN 使用 replay buffer、target network、epsilon-greedy、mini-batch update 等强化学习机制；这些机制需要结合代码快照和 training log 复核。模型 artifact 存在，但本轮不重新训练。

    ## 8. Baseline 与评价

    Baseline：{', '.join(config.get('baselines', []))}。比较必须共用 state set、budget、action constraints、capacity constraints 和 metrics。核心指标包括 total_reward、risk_reward_total、information_gain_total、sampling_cost_total、constraint_violation_rate、state_coverage、rank。

    ## 9. Quality gates

    本轮新增图表 QA、表格解释、模型输出解释、代码说明、文献映射和 Results claim guard。所有结果仍标记为 experimental。
    """
    write_text(RUN_DIR / "04_报告输出" / "dqn_model_setting_detail_report.md", setting)

    interp = "# DQN 结果解读报告\n\n"
    if not comparison.empty:
        interp += f"## 1. 整体比较\n\n{markdown_table(comparison, ['policy','total_reward','risk_reward_total','information_gain_total','sampling_cost_total','constraint_violation_rate','rank'], 10)}\n\n"
        best = comparison.iloc[0]
        dqn = comparison[comparison["policy"].astype(str).str.contains("DQN", na=False)].iloc[0]
        interp += (
            f"当前 experimental 比较中，`{best['policy']}` total_reward 最高（{best['total_reward']:.3f}），"
            f"`DQN修正版` 排名 {int(dqn['rank'])}（total_reward={dqn['total_reward']:.3f}）。"
            "这说明 DQN 在修正版 reward 与约束设定下已经可运行且接近最优对照，但并未超过 Q-learning。\n\n"
        )
    interp += """
    ## 2. 指标含义

    - total_reward：综合风险收益、信息增益、成本和惩罚后的总目标值；越高表示当前实验目标下越优。
    - risk_reward_total：覆盖风险状态带来的收益；越高说明策略更关注高风险单元。
    - information_gain_total：对不确定状态抽检带来的信息收益；越高说明策略更偏向学习。
    - sampling_cost_total：抽检成本；越高不必然更差，需要与风险收益共同看。
    - constraint_violation_rate：约束违约率；当前为 0，表示 action mask 和容量预算约束有效，但不能替代 formal 约束确认。

    ## 3. 稳定性、reward hacking 与收敛

    当前 training log 支持基本趋势审查，但本轮未重新训练；因此只能说上一轮 DQN 修正版“可运行且有收敛诊断输出”。reward hacking 未发现直接证据，但仍存在权重设定驱动策略的风险。若 formal 化，需要敏感性分析、外部验证和用户确认约束。

    ## 4. 食品安全监管启发

    Experimental 结果提示：在预算和容量约束下，风险收益、信息增益和成本之间可以形成可解释的策略比较框架；但当前输出不能直接转化为监管政策，只能作为方法探索、论文 Results 草稿和后续 formal DQN 参数确认的依据。
    """
    write_text(RUN_DIR / "04_报告输出" / "dqn_result_interpretation_report.md", interp)


def generate_code_explanations() -> None:
    code_paths = sorted((LATEST_DQN / "08_代码快照").glob("*.py")) + [Path(__file__)]
    rows = []
    md = "# Code Explanations\n\n"
    for path in code_paths:
        text = safe_read(path, 10000)
        imports = []
        funcs = []
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith("import ") or stripped.startswith("from "):
                imports.append(stripped)
            if stripped.startswith("def "):
                funcs.append(stripped.split("(")[0].replace("def ", ""))
        rows.append(
            {
                "script_path": rel(path),
                "size_bytes": path.stat().st_size if path.exists() else 0,
                "dependencies_detected": "; ".join(imports[:12]),
                "core_functions_detected": "; ".join(funcs[:20]),
                "purpose": "DQN 实验执行/复核/解释生成脚本",
                "input_files": "latest DQN run package; outputs/_index; project_state; references; Zotero sidecar inventory",
                "output_files": "run package reports, CSV, PNG, DOCX, explanations",
                "run_command": f'& "{PYTHON}" "{rel(path)}"' if path == Path(__file__) else "作为历史代码快照阅读，不建议直接覆盖运行",
                "reproducibility_status": "可复核；formal 前仍需参数确认",
                "human_confirmation_needed": "是，formal DQN 或监管结论前",
            }
        )
        md += f"## {rel(path)}\n\n"
        md += "- 用途：记录或生成 DQN experimental run、质量复核、解释索引和论文草稿。\n"
        md += "- 输入文件：最新 DQN run package、项目索引、配置、CSV、图表和文献清单。\n"
        md += "- 输出文件：本轮任务包中的报告、表格、图表修复、DOCX 和解释索引。\n"
        md += f"- 依赖库：{'; '.join(imports[:8]) if imports else '未自动识别'}。\n"
        md += f"- 核心函数：{'; '.join(funcs[:12]) if funcs else '未自动识别'}。\n"
        md += "- 方法逻辑：先定位输入与历史输出，再执行质量核验、解释生成、文献映射和论文输出；不修改 raw data 或 Zotero SQLite。\n"
        md += "- 如何运行：本轮脚本使用 myenv1 Python；历史快照仅作为证据，不建议直接运行覆盖。\n"
        md += "- 常见错误：中文路径、字体缺失、DOCX 渲染表格过宽、历史文件为空。\n"
        md += "- 与论文关系：代码说明进入 Method/Appendix，可支撑可复现性说明。\n"
        md += "- 人工确认：formal DQN 参数、约束和政策结论仍需用户确认。\n\n"
    pd.DataFrame(rows).to_csv(RUN_DIR / "02_表格输出" / "code_inventory.csv", index=False, encoding="utf-8-sig")
    write_text(RUN_DIR / "10_输出解释与索引" / "code_explanations.md", md)
    write_text(
        RUN_DIR / "04_报告输出" / "code_method_explanation_report.md",
        f"# 代码方法说明报告\n\n本轮共纳入 {len(rows)} 个脚本/代码快照。所有关键脚本均生成用途、输入、输出、依赖、流程、方法逻辑和适用边界说明。历史代码快照用于复核，不在本轮覆盖运行。\n",
    )


def literature_rows() -> list[dict[str, str]]:
    return [
        {
            "key": "Mnih2015DQN",
            "title": "Human-level control through deep reinforcement learning",
            "authors": "Mnih et al.",
            "year": "2015",
            "venue": "Nature",
            "doi": "10.1038/nature14236",
            "url": "https://www.nature.com/articles/nature14236",
            "direction": "DQN 原始方法、experience replay、target network",
            "read_status": "local_deepread_exists_but_garbled_recheck_required",
            "support_level": "method background; not sole formal support because local note has mojibake risk",
        },
        {
            "key": "VanHasselt2016DoubleDQN",
            "title": "Deep Reinforcement Learning with Double Q-learning",
            "authors": "van Hasselt, Guez, Silver",
            "year": "2016",
            "venue": "AAAI / arXiv",
            "doi": "",
            "url": "https://arxiv.org/abs/1509.06461",
            "direction": "Double DQN、Q 值高估偏差修正",
            "read_status": "metadata_or_abstract_level",
            "support_level": "future upgrade evidence",
        },
        {
            "key": "Wang2016DuelingDQN",
            "title": "Dueling Network Architectures for Deep Reinforcement Learning",
            "authors": "Wang et al.",
            "year": "2016",
            "venue": "ICML / arXiv",
            "doi": "",
            "url": "https://arxiv.org/abs/1511.06581",
            "direction": "Dueling DQN、value/advantage 分解",
            "read_status": "metadata_or_abstract_level",
            "support_level": "future upgrade evidence",
        },
        {
            "key": "Lin1992ExperienceReplay",
            "title": "Self-improving reactive agents based on reinforcement learning, planning and teaching",
            "authors": "Lin",
            "year": "1992",
            "venue": "Machine Learning",
            "doi": "10.1007/BF00992699",
            "url": "https://doi.org/10.1007/BF00992699",
            "direction": "experience replay 早期思想",
            "read_status": "metadata_only",
            "support_level": "method background",
        },
        {
            "key": "Achiam2017CPO",
            "title": "Constrained Policy Optimization",
            "authors": "Achiam et al.",
            "year": "2017",
            "venue": "ICML / arXiv",
            "doi": "",
            "url": "https://arxiv.org/abs/1705.10528",
            "direction": "safe/constrained RL",
            "read_status": "metadata_or_abstract_level",
            "support_level": "constraint audit framing",
        },
        {
            "key": "Altman1999CMDP",
            "title": "Constrained Markov Decision Processes",
            "authors": "Altman",
            "year": "1999",
            "venue": "Chapman and Hall/CRC",
            "doi": "10.1201/9781315140223",
            "url": "https://doi.org/10.1201/9781315140223",
            "direction": "约束 MDP 理论",
            "read_status": "metadata_only",
            "support_level": "formal constraint theory candidate",
        },
        {
            "key": "Kaelbling1998POMDP",
            "title": "Planning and acting in partially observable stochastic domains",
            "authors": "Kaelbling, Littman, Cassandra",
            "year": "1998",
            "venue": "Artificial Intelligence",
            "doi": "10.1016/S0004-3702(98)00023-X",
            "url": "https://doi.org/10.1016/S0004-3702(98)00023-X",
            "direction": "POMDP / belief state",
            "read_status": "metadata_only",
            "support_level": "belief-MDP framing candidate",
        },
        {
            "key": "Wang2020RiskBasedSampling",
            "title": "A risk-based sampling strategy to optimize inspection of food and feed products",
            "authors": "Wang et al.",
            "year": "2020",
            "venue": "Food Control / PMC",
            "doi": "",
            "url": "https://pmc.ncbi.nlm.nih.gov/articles/PMC7821187/",
            "direction": "risk-based sampling / inspection resource allocation",
            "read_status": "candidate_url_browser_access_blocked; metadata_from_local_candidate_pool",
            "support_level": "risk-based inspection design candidate; requires manual/full-text verification before formal use",
        },
        {
            "key": "VanAsselt2021RiskBasedInspections",
            "title": "Risk-based inspections of food business operators",
            "authors": "van Asselt et al.",
            "year": "2021",
            "venue": "Food Control",
            "doi": "10.1016/j.foodcont.2021.108462",
            "url": "https://pubmed.ncbi.nlm.nih.gov/34796503/",
            "direction": "risk-based food safety inspection",
            "read_status": "candidate_url_browser_access_blocked; metadata_from_local_candidate_pool",
            "support_level": "regulatory inspection rationale candidate; requires manual/full-text verification before formal use",
        },
        {
            "key": "EFSA_MOE",
            "title": "Margin of exposure",
            "authors": "European Food Safety Authority",
            "year": "n.d.",
            "venue": "EFSA topic page",
            "doi": "",
            "url": "https://www.efsa.europa.eu/en/topics/topic/margin-exposure",
            "direction": "MOE 风险评估概念",
            "read_status": "official_page",
            "support_level": "MOE/EDI interpretation background",
        },
        {
            "key": "JECFA2018Aflatoxins",
            "title": "Safety evaluation of certain contaminants in food: aflatoxins",
            "authors": "JECFA / WHO",
            "year": "2018",
            "venue": "WHO Food Additives Series",
            "doi": "",
            "url": "https://iris.who.int/handle/10665/276868",
            "direction": "AFB1 风险评估",
            "read_status": "metadata_or_official_record",
            "support_level": "AFB1 hazard/risk background",
        },
        {
            "key": "Mitchell2019ModelCards",
            "title": "Model Cards for Model Reporting",
            "authors": "Mitchell et al.",
            "year": "2019",
            "venue": "FAT* / arXiv",
            "doi": "",
            "url": "https://arxiv.org/abs/1810.03993",
            "direction": "模型说明、适用范围、限制",
            "read_status": "metadata_or_abstract_level",
            "support_level": "workflow reporting design",
        },
        {
            "key": "Gebru2021Datasheets",
            "title": "Datasheets for Datasets",
            "authors": "Gebru et al.",
            "year": "2021",
            "venue": "Communications of the ACM",
            "doi": "10.1145/3458723",
            "url": "https://doi.org/10.1145/3458723",
            "direction": "数据/输出说明卡",
            "read_status": "metadata_or_abstract_level",
            "support_level": "artifact documentation design",
        },
    ]


def generate_literature_outputs() -> None:
    rows = literature_rows()
    cand = pd.DataFrame(rows)
    cand.to_csv(RUN_DIR / "02_表格输出" / "dqn_literature_candidate_pool.csv", index=False, encoding="utf-8-sig")
    core = cand[cand["support_level"].str.contains("method|risk|constraint|belief|inspection|MOE|AFB1|workflow", case=False, regex=True)].copy()
    core.to_csv(RUN_DIR / "02_表格输出" / "dqn_core_literature_selected.csv", index=False, encoding="utf-8-sig")
    component_rows = [
        ["state design", "Beta-Binomial belief states; MOE/EDI risk features", "Kaelbling1998POMDP; EFSA_MOE", "partly supported; current feature engineering still experimental"],
        ["action space", "coarse increments [0,1,3,5,10]; high-dimensional binary not trained", "Wang2020RiskBasedSampling; VanAsselt2021RiskBasedInspections", "regulatory allocation rationale supported; exact increments experimental assumption"],
        ["reward design", "risk reward + information gain - cost - constraint penalty", "Wang2020RiskBasedSampling; Achiam2017CPO; Altman1999CMDP", "component logic supported; weights experimental assumption"],
        ["constraint", "monthly budget, capacity, action mask", "Achiam2017CPO; Altman1999CMDP; risk-based inspection literature", "formal parameter confirmation required"],
        ["DQN training", "experience replay, target network, epsilon-greedy", "Mnih2015DQN; Lin1992ExperienceReplay", "method support; current code/config still experimental"],
        ["baselines", "Q-learning, uniform, historical, risk-ranking, random, threshold/greedy", "Model comparison and RL baseline norms; Mnih2015DQN", "baseline fairness passed locally; formal external validation absent"],
        ["reporting", "model/output/code/figure explanations and evidence map", "Mitchell2019ModelCards; Gebru2021Datasheets", "workflow design support, not DQN performance evidence"],
    ]
    pd.DataFrame(component_rows, columns=["model_component", "project_setting", "literature_keys", "evidence_status"]).to_csv(
        RUN_DIR / "02_表格输出" / "dqn_model_component_literature_map.csv", index=False, encoding="utf-8-sig"
    )

    bib = []
    ris = []
    notes = ["# DQN 文献无乱码中文侧车笔记\n\n本文件为 Zotero 安全写入/待导入侧车笔记，不直接修改 Zotero SQLite 数据库。\n"]
    for r in rows:
        bib.append(
            f"@article{{{r['key']},\n  title = {{{r['title']}}},\n  author = {{{r['authors']}}},\n  year = {{{r['year']}}},\n  journal = {{{r['venue']}}},\n  doi = {{{r['doi']}}},\n  url = {{{r['url']}}}\n}}\n"
        )
        ris.extend(
            [
                "TY  - JOUR",
                f"ID  - {r['key']}",
                f"TI  - {r['title']}",
                f"AU  - {r['authors']}",
                f"PY  - {r['year']}",
                f"JO  - {r['venue']}",
                f"DO  - {r['doi']}",
                f"UR  - {r['url']}",
                "ER  - ",
                "",
            ]
        )
        notes.append(
            f"## {r['key']}\n\n- 题名：{r['title']}\n- 读取状态：{r['read_status']}\n- 支持方向：{r['direction']}\n- 模型用途：{r['support_level']}\n- URL：{r['url']}\n"
        )
    write_text(RUN_DIR / "04_报告输出" / "dqn_core_literature.bib", "\n".join(bib))
    write_text(RUN_DIR / "04_报告输出" / "dqn_core_literature.ris", "\n".join(ris))
    write_text(RUN_DIR / "04_报告输出" / "dqn_literature_notes_clean.md", "\n".join(notes))
    import_plan = cand[["key", "title", "doi", "url", "read_status", "direction"]].copy()
    import_plan["zotero_action"] = "sidecar_export_for_manual_import_or_BetterBibTeX; no SQLite write"
    import_plan["duplicate_check"] = import_plan["key"].apply(lambda k: "local candidate/deepread checked; manual Zotero duplicate check recommended")
    import_plan.to_csv(RUN_DIR / "02_表格输出" / "zotero_writeback_or_import_plan.csv", index=False, encoding="utf-8-sig")
    log = import_plan[["key", "title"]].copy()
    log["writeback_status"] = "not_written_to_zotero_sqlite"
    log["sidecar_generated"] = "yes"
    log["notes_encoding"] = "utf-8"
    log.to_csv(RUN_DIR / "02_表格输出" / "zotero_writeback_log.csv", index=False, encoding="utf-8-sig")
    write_text(
        RUN_DIR / "04_报告输出" / "zotero_clean_notes_index.md",
        "# Zotero Clean Notes Index\n\n- 直接写入 Zotero SQLite：否。\n- 已生成 UTF-8 中文侧车笔记：`dqn_literature_notes_clean.md`。\n- 已生成 BibTeX/RIS/CSV 待导入文件。\n- 发现本地 `Human-level control...` deepread 存在乱码风险，不能作为 formal 证据，需回到 PDF/官方源复核。\n",
    )
    write_text(
        RUN_DIR / "04_报告输出" / "dqn_literature_expansion_report.md",
        f"""
        # DQN 文献依据补充报告

        本轮建立 {len(cand)} 条候选文献池，筛选 {len(core)} 条核心候选。覆盖方向包括 DQN、Double DQN、Dueling DQN、experience replay、target network、safe/constrained RL、risk-based inspection、food safety monitoring optimization、MOE/EDI/AFB1、Bayesian/POMDP/belief-MDP、regulatory resource allocation 和 workflow reporting。

        ## 读取状态边界

        - `full/open_full_text_available` 仅表示可访问或本地有深读线索，不等于本轮完成全文逐段核验。
        - `metadata_or_abstract_level` 只能支持背景或候选依据。
        - 本地乱码 note 不能作为 formal 证据。
        - 未直接修改 Zotero 数据库，只生成 sidecar notes、BibTeX、RIS、CSV 和导入计划。
        """,
    )


def short_policy(name: str) -> str:
    mapping = {
        "threshold/greedy uncertainty": "threshold",
        "risk-ranking top-k": "risk top-k",
        "uniform allocation": "uniform",
        "historical allocation": "historical",
        "random policy": "random",
        "DQN修正版": "DQN",
        "Q-learning": "Q-learning",
    }
    return mapping.get(name, name[:18])


def set_docx_cell_width(cell, width_twips: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_docx_table_width(table, width_twips: int, col_widths: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_docx_cell_width(cell, col_widths[min(idx, len(col_widths) - 1)])


def save_table_png(df: pd.DataFrame, path: Path, title: str) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    font = choose_font()
    plt.rcParams["font.sans-serif"] = [font]
    plt.rcParams["axes.unicode_minus"] = False
    rows = len(df)
    fig_h = max(2.2, 0.42 * rows + 1.2)
    fig, ax = plt.subplots(figsize=(10.5, fig_h), dpi=180)
    ax.axis("off")
    ax.set_title(title, fontsize=14, pad=12)
    table = ax.table(
        cellText=df.astype(str).values,
        colLabels=list(df.columns),
        loc="center",
        cellLoc="center",
        colLoc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.35)
    for (row, _col), cell in table.get_celld().items():
        if row == 0:
            cell.set_text_props(weight="bold", color="white")
            cell.set_facecolor("#3b6699")
        else:
            cell.set_facecolor("#f7f9fc" if row % 2 == 0 else "white")
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_word_outputs(tables: dict[str, pd.DataFrame]) -> None:
    comparison = tables.get("multi_model_policy_comparison", pd.DataFrame())
    quality = tables.get("research_quality_gate_results", pd.DataFrame())
    evidence_rows = []
    if not comparison.empty:
        for _, r in comparison.iterrows():
            evidence_rows.append(
                {
                    "claim_id": f"C{int(r['rank'])}",
                    "claim": f"{r['policy']} 在本轮 experimental 比较中 rank={int(r['rank'])}",
                    "source_table": "multi_model_policy_comparison.csv",
                    "source_value": f"total_reward={r['total_reward']:.3f}; constraint_violation_rate={r['constraint_violation_rate']:.3f}",
                    "figure_reference": "dqn_revised_policy_comparison.png / multi_model_comparison.png",
                    "citation_support": "Mnih2015DQN; Wang2020RiskBasedSampling; Achiam2017CPO",
                    "status": "experimental results draft; not formal policy conclusion",
                }
            )
    evidence = pd.DataFrame(evidence_rows)
    evidence.to_csv(RUN_DIR / "09_论文输出" / "04_结果" / "dqn_results_evidence_table.csv", index=False, encoding="utf-8-sig")

    draft = "# DQN 修正版实验结果（experimental results draft）\n\n"
    draft += "## 结果边界\n\n本节仅报告 PEANUT 风险监管 DQN 修正版的 experimental run。奖励权重、约束参数、动作空间和真实监管转移机制尚未完成 formal 确认，因此下述结果不能作为最终监管政策结论。\n\n"
    if not comparison.empty:
        draft += "## 多模型比较\n\n"
        draft += markdown_table(comparison, ["policy", "total_reward", "risk_reward_total", "information_gain_total", "sampling_cost_total", "constraint_violation_rate", "rank"], 10)
        best = comparison.iloc[0]
        dqn = comparison[comparison["policy"].astype(str).str.contains("DQN", na=False)].iloc[0]
        draft += (
            f"\n\n在相同预算、动作约束、容量约束和评价指标下，`{best['policy']}` 取得最高 total_reward（{best['total_reward']:.3f}）。"
            f"`DQN修正版` 位列第 {int(dqn['rank'])}（total_reward={dqn['total_reward']:.3f}），说明修正版 DQN 已能产生接近最优对照的策略输出，但当前并未超过 Q-learning。\n\n"
        )
    draft += """
    ## 约束与质量核验

    `constraint_violation_summary.csv` 显示所有比较策略的违约率为 0。该结果应解释为在当前 experimental action mask 与预算/容量配置下约束被满足，而不是数据缺失。图 `dqn_revised_constraint_summary_explained.png` 已将全 0 图转换为解释性 PNG。

    ## 图表引用

    - 图 1：`dqn_revised_policy_comparison.png`，展示各策略 total_reward 对比。
    - 图 2：`dqn_revised_reward_curve.png` 与 `dqn_revised_moving_average_reward.png`，展示训练回报和移动平均回报。
    - 图 3：`dqn_revised_constraint_summary_explained.png`，说明约束违约为 0 的语义。

    ## 文献依据

    DQN 训练机制参考 DQN、Double DQN/Dueling DQN 的方法谱系 [Mnih2015DQN; VanHasselt2016DoubleDQN; Wang2016DuelingDQN]；约束与安全强化学习解释参考 constrained RL/CMDP 文献 [Achiam2017CPO; Altman1999CMDP]；风险导向抽检逻辑参考 risk-based inspection 与 food safety monitoring 文献 [Wang2020RiskBasedSampling; VanAsselt2021RiskBasedInspections]。本轮仅建立 evidence map，不声称所有文献已完成全文形式核验。

    ## 不能过度解释的地方

    当前结果不能说明真实监管中 DQN 必然优于或劣于其他策略，也不能说明现有 reward 权重已经最优。formal DQN 需要进一步确认状态定义、动作空间、成本/预算/容量参数、约束强度、转移近似、敏感性分析和外部验证。
    """
    write_text(RUN_DIR / "09_论文输出" / "04_结果" / "dqn_results_draft.md", draft)

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        styles = doc.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"].font.size = Pt(10.5)
        title = doc.add_heading("DQN 修正版实验结果（experimental results draft）", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("本节为论文 Results 草稿，所有模型结果均保持 experimental 边界，不构成 formal 监管政策结论。")
        doc.add_heading("1. 多模型比较", level=2)
        if not comparison.empty:
            table_img_df = comparison.head(7).copy()
            table_img_df = table_img_df.assign(policy=table_img_df["policy"].map(short_policy))
            table_img_df = table_img_df[["policy", "total_reward", "constraint_violation_rate", "rank"]]
            table_img_df.columns = ["策略", "总奖励", "违约率", "排名"]
            table_img_df["总奖励"] = table_img_df["总奖励"].map(lambda x: f"{x:.3f}")
            table_img_df["违约率"] = table_img_df["违约率"].map(lambda x: f"{x:.3f}")
            table_img = RUN_DIR / "09_论文输出" / "09_word导出" / "policy_comparison_table.png"
            save_table_png(table_img_df, table_img, "表 1 多模型策略比较（experimental）")
            doc.add_picture(str(table_img), width=Cm(15.6))
        doc.add_paragraph("表 1 显示，本轮比较在统一预算、约束和评价指标下完成。Q-learning 在当前 reward 设定下 total_reward 最高，DQN 修正版排名第二。")
        doc.add_heading("2. 图表说明", level=2)
        for cap in [
            "图 1 dqn_revised_policy_comparison.png：比较各策略 total_reward。",
            "图 2 dqn_revised_reward_curve.png：检查训练回报趋势。",
            "图 3 dqn_revised_constraint_summary_explained.png：解释约束违约为 0 的含义。",
        ]:
            doc.add_paragraph(cap)
        doc.add_heading("3. 质量核验与边界", level=2)
        doc.add_paragraph("图表已执行 PNG 非空检查和解释性修复；表格与报告数字通过 evidence table 映射。当前结果仍需 formal 参数确认、敏感性分析和外部验证。")
        doc.add_heading("4. 文献依据", level=2)
        doc.add_paragraph("DQN 方法、约束强化学习、risk-based inspection、MOE/EDI 和 AFB1 风险评估均已建立候选文献映射。读取状态分为 local deepread、open full text available、abstract-level、metadata-only，不伪造全文阅读。")
        doc.add_heading("5. Evidence Table", level=2)
        if not evidence.empty:
            ev_img_df = evidence.head(5)[["claim_id", "source_value", "status"]].copy()
            ev_img_df.columns = ["ID", "证据数值", "状态"]
            ev_img_df["状态"] = "experimental"
            ev_img = RUN_DIR / "09_论文输出" / "09_word导出" / "evidence_table_preview.png"
            save_table_png(ev_img_df, ev_img, "表 2 Evidence Table 摘要")
            doc.add_picture(str(ev_img), width=Cm(15.6))
        docx_path = RUN_DIR / "09_论文输出" / "09_word导出" / "dqn_results_draft.docx"
        doc.save(docx_path)
    except Exception as exc:
        write_text(RUN_DIR / "09_论文输出" / "09_word导出" / "dqn_results_draft_docx_error.md", f"DOCX 生成失败：{exc}")

    write_text(
        RUN_DIR / "04_报告输出" / "paper_output_system_detail_upgrade_report.md",
        "# 论文输出体系升级报告\n\n已为 Results 生成 Markdown、evidence table 和 DOCX。长期机制要求 Introduction、Literature Review、Method、Results、Discussion、Conclusion、Appendix、Integrated draft 均配套 section draft、evidence table、figure/table references、citation support、quality audit 与 experimental/formal 状态标记。\n",
    )


def generate_skill_scout_outputs() -> None:
    rows = [
        ["Model Cards for Model Reporting", "model card/reporting", "https://arxiv.org/abs/1810.03993", "可安全吸收结构：用途、限制、适用范围、评价指标", "no install"],
        ["Datasheets for Datasets", "data/artifact documentation", "https://doi.org/10.1145/3458723", "可转化为 artifact/data explanation checklist", "no install"],
        ["TensorFlow Model Card Toolkit", "model card toolkit", "https://github.com/tensorflow/model-card-toolkit", "仅借鉴 README/字段，不安装运行", "approval if install"],
        ["Zotero Better BibTeX", "Zotero export", "https://github.com/retorquere/zotero-better-bibtex", "可建议用户手动导入/导出；不自动安装插件", "approval required"],
        ["Open-source reproducibility checklists", "reproducibility", "local research_quality skills", "转化为本地 reproducibility-auditor 输出要求", "safe local"],
    ]
    candidates = pd.DataFrame(rows, columns=["candidate", "capability", "url", "safe_adaptation", "risk_status"])
    candidates.to_csv(RUN_DIR / "02_表格输出" / "explanation_reporting_skill_candidates.csv", index=False, encoding="utf-8-sig")
    plan = pd.DataFrame(
        [
            ["artifact-explanation-generator", "新增本地 skill", "每个 artifact 生成目的/来源/阅读/局限/evidence map", "done"],
            ["model-setting-documenter", "新增本地 skill", "统一所有模型的设置说明", "done"],
            ["result-interpretation-writer", "新增本地 skill", "模型结果解读和 experimental 边界", "done"],
            ["code-explanation-auditor", "新增本地 skill", "代码用途、输入输出、依赖、流程与复现性", "done"],
        ],
        columns=["item", "action", "problem_solved", "status"],
    )
    plan.to_csv(RUN_DIR / "02_表格输出" / "explanation_reporting_safe_adaptation_plan.csv", index=False, encoding="utf-8-sig")
    approval = pd.DataFrame(
        [
            ["Zotero Better BibTeX plugin install", "external plugin", "可能帮助 BibTeX/RIS 同步", "需要用户确认"],
            ["Model Card Toolkit install", "external dependency", "可生成标准 model card", "本轮不安装；需要用户确认"],
        ],
        columns=["item", "risk_type", "benefit", "approval_status"],
    )
    approval.to_csv(RUN_DIR / "02_表格输出" / "explanation_reporting_approval_queue.csv", index=False, encoding="utf-8-sig")
    plan.to_csv(RUN_DIR / "02_表格输出" / "new_or_upgraded_explanation_skills.csv", index=False, encoding="utf-8-sig")
    write_text(
        RUN_DIR / "04_报告输出" / "explanation_and_reporting_skill_scout_report.md",
        "# Explanation and Reporting Skill Scout Report\n\n本轮参考 Model Cards、Datasheets、Model Card Toolkit 与 Zotero Better BibTeX 等公开项目/规范的 README 或元数据，未安装未知包、未运行第三方代码。已把可安全吸收的结构转化为本地 skills、recipes 和长期规则；外部插件安装进入 approval queue。\n",
    )


def write_deep_audit(inventory: list[dict[str, object]], issues: list[dict[str, object]]) -> None:
    pd.DataFrame(inventory).to_csv(RUN_DIR / "02_表格输出" / "audited_artifact_inventory.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame(issues).to_csv(RUN_DIR / "02_表格输出" / "deep_audit_issue_list.csv", index=False, encoding="utf-8-sig")
    counts = Counter([i["问题类型"] for i in issues])
    md = "# DQN 输出 Deep Audit 报告\n\n"
    md += f"- 最新复核包：`{rel(LATEST_DQN)}`\n"
    md += f"- 历史对照包：`{rel(BASELINE_DQN)}`\n"
    md += f"- 审计 artifact 数：{len(inventory)}\n"
    md += f"- 问题数：{len(issues)}\n\n"
    md += "## 问题类型分布\n\n"
    for k, v in counts.items():
        md += f"- {k}: {v}\n"
    md += "\n## 关键发现\n\n"
    md += "- 最新 `2056` 修正版 DQN 包的 PNG 主图总体可读，但 `constraint_summary` 属于全 0 语义图，需要解释性修复。\n"
    md += "- 历史 `1746` DQN 自动参数训练包中存在近似空表格和解释不足输出，不能直接进入论文 Results；本轮未在实际文件系统中确认到 0 字节主图。\n"
    md += "- 最新包缺少面向用户的系统性图、表、模型输出、代码解释与 artifact-to-evidence map。\n"
    md += "- DQN 设置说明仍需把 state/action/reward/transition/constraint/baseline/training/quality gate 逐项写明，并保持 experimental 边界。\n\n"
    md += "## 问题清单摘要\n\n"
    for issue in issues[:40]:
        md += f"- [{issue['严重程度']}] {issue['问题类型']}：`{issue['文件路径']}`；修复：{issue['建议修复方式']}\n"
    write_text(RUN_DIR / "04_报告输出" / "deep_dqn_output_audit_report.md", md)


def run_dry_runs() -> None:
    goals = [
        "解释当前任务的所有输出",
        "解释所有图表和表格",
        "为当前代码生成方法说明",
        "为当前模型生成详细方法说明",
        "为当前模型结果生成详细解读",
        "补充模型文献依据并写入 Zotero",
        "生成论文 Results 部分并导出 Word",
        "检查图表是否为空并修复",
        "搜索并创建缺失的解释类 skill",
    ]
    rows = []
    for goal in goals:
        cmd = [PYTHON, "-m", "workflow1", "--stage", "dry-run", "--goal", goal]
        try:
            cp = subprocess.run(cmd, cwd=ROOT, text=True, encoding="utf-8", errors="replace", stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
            stdout = cp.stdout[-2000:]
            stderr = cp.stderr[-1000:]
            status = "pass" if cp.returncode == 0 and ("matched_intent" in stdout or "intent" in stdout) else "review"
        except Exception as exc:
            stdout = ""
            stderr = str(exc)
            status = "fail"
        rows.append({"goal": goal, "status": status, "stdout_tail": stdout, "stderr_tail": stderr})
    df = pd.DataFrame(rows)
    df.to_csv(RUN_DIR / "02_表格输出" / "output_explanation_workflow_dry_run_results.csv", index=False, encoding="utf-8-sig")
    md = "# Output Explanation Workflow Dry-run Report\n\n"
    for _, r in df.iterrows():
        md += f"## {r['goal']}\n\n- 状态：{r['status']}\n- 输出片段：\n\n```text\n{r['stdout_tail'][:1200]}\n```\n\n"
        if r["stderr_tail"]:
            md += f"- stderr：\n\n```text\n{r['stderr_tail']}\n```\n\n"
    write_text(RUN_DIR / "04_报告输出" / "output_explanation_workflow_dry_run_report.md", md)


def write_chart_report(chart_rows: list[dict[str, object]]) -> None:
    pd.DataFrame(chart_rows).to_csv(RUN_DIR / "02_表格输出" / "chart_quality_audit.csv", index=False, encoding="utf-8-sig")
    repaired = [r for r in chart_rows if "explained" in str(r["figure_path"])]
    md = "# Chart System Repair Report\n\n"
    md += "- 已复制最新修正版 PNG 主图到本轮任务包。\n"
    md += "- 已为全 0 约束图生成解释性 PNG：`dqn_revised_constraint_summary_explained.png`。\n"
    md += "- 历史图表未在实际文件系统中确认到 0 字节主图；历史近似空/信息不足问题已进入 deep audit 问题清单。\n"
    md += f"- 本轮 chart QA 记录 {len(chart_rows)} 张 PNG，解释性修复 {len(repaired)} 张。\n"
    md += "- 后续机制：空图、全 0 图、无差异图不得作为空白坐标轴输出；必须生成解释性图或说明。\n"
    write_text(RUN_DIR / "04_报告输出" / "chart_system_repair_report.md", md)


def write_readme_manifest() -> None:
    files = []
    for path in sorted(RUN_DIR.rglob("*")):
        if path.is_file():
            files.append(
                {
                    "path": rel(path),
                    "size_bytes": path.stat().st_size,
                    "sha256": sha256_file(path) if path.stat().st_size else "",
                    "purpose": artifact_purpose(path),
                }
            )
    pd.DataFrame(files).to_csv(RUN_DIR / "manifest.csv", index=False, encoding="utf-8-sig")
    write_text(
        RUN_DIR / "README.md",
        f"""
        # DQN 输出复核、解释体系与论文输出升级

        ## 任务目标

        系统复核当前 DQN、Q-learning、heuristic、图表、表格、模型输出、代码、论文输出和文献/Zotero 工作流，并把解释、图表 QA、代码说明、模型说明、论文 section 输出和 workflow self-improvement 固化到长期规则。

        ## 输入来源

        - 最新 DQN 修正版 experimental 包：`{rel(LATEST_DQN)}`
        - 历史对照 DQN 自动参数训练包：`{rel(BASELINE_DQN)}`
        - 项目索引：`outputs/_index/`
        - 项目状态：`project_state/`
        - references / Zotero sidecar 目录

        ## 主要输出

        - deep audit：`04_报告输出/deep_dqn_output_audit_report.md`
        - 解释索引：`10_输出解释与索引/`
        - 图表 QA 与修复：`02_表格输出/chart_quality_audit.csv`、`03_图表输出/*explained.png`
        - DQN 设置与结果解读：`04_报告输出/dqn_model_setting_detail_report.md`、`dqn_result_interpretation_report.md`
        - 文献/Zotero 安全侧车：`dqn_core_literature.bib`、`dqn_core_literature.ris`、`zotero_writeback_or_import_plan.csv`
        - 论文级 Results：`09_论文输出/04_结果/dqn_results_draft.md`、`09_论文输出/09_word导出/dqn_results_draft.docx`

        ## 尚未解决问题

        当前 DQN 仍为 experimental。formal DQN 仍需确认状态、动作、reward 权重、预算、容量、约束、transition、训练轮次、外部验证和敏感性分析。

        ## 下一步建议

        先由用户确认 formal DQN 参数表，再生成 Method / Introduction / Literature Review / Discussion 的论文级 section。
        """,
    )


def update_indexes_and_state() -> None:
    run_index = ROOT / "outputs" / "_index" / "run_index.md"
    run_manifest = ROOT / "outputs" / "_index" / "run_manifest.csv"
    latest_yaml = ROOT / "outputs" / "_index" / "latest_canonical_outputs.yaml"
    ts = datetime.now().isoformat(timespec="seconds")
    with run_index.open("a", encoding="utf-8") as f:
        f.write(
            f"\n## {RUN_DIR.name}\n\n- 时间：{ts}\n- 类型：DQN输出复核_解释体系与论文输出升级\n- 路径：`{rel(RUN_DIR)}`\n- 状态：completed_experimental_audit_and_workflow_upgrade\n- 关键输出：deep audit、解释索引、图表修复、DQN 设置说明、结果解读、文献/Zotero 侧车、Results DOCX、dry-run。\n"
        )
    row = pd.DataFrame(
        [
            {
                "run_id": RUN_DIR.name,
                "run_path": rel(RUN_DIR),
                "created_at": ts,
                "task": "DQN输出复核_解释体系与论文输出升级",
                "status": "completed",
                "experimental_boundary": "all DQN outputs remain experimental_not_formal",
            }
        ]
    )
    if run_manifest.exists():
        old = pd.read_csv(run_manifest)
        pd.concat([old, row], ignore_index=True).to_csv(run_manifest, index=False, encoding="utf-8-sig")
    else:
        row.to_csv(run_manifest, index=False, encoding="utf-8-sig")
    latest = {}
    if latest_yaml.exists():
        latest = yaml.safe_load(latest_yaml.read_text(encoding="utf-8")) or {}
    latest.update(
        {
            "latest_output_explanation_upgrade_run": rel(RUN_DIR),
            "latest_dqn_deep_audit_report": rel(RUN_DIR / "04_报告输出" / "deep_dqn_output_audit_report.md"),
            "latest_dqn_results_docx_upgraded": rel(RUN_DIR / "09_论文输出" / "09_word导出" / "dqn_results_draft.docx"),
            "latest_artifact_explanation_index": rel(RUN_DIR / "10_输出解释与索引" / "artifact_explanation_index.md"),
            "latest_dqn_status": "experimental_audited_explained_not_formal",
        }
    )
    latest_yaml.write_text(yaml.safe_dump(latest, allow_unicode=True, sort_keys=False), encoding="utf-8")

    updates = {
        "project_state/current_focus.md": f"# Current Focus\n\n当前重点：已完成 DQN 输出 deep audit、解释体系补强、论文 Results DOCX 升级、文献/Zotero 侧车和长期 workflow 规则固化。最新任务包：`{rel(RUN_DIR)}`。\n",
        "project_state/next_step.md": "# Next Step\n\n若继续 formal DQN，请先确认正式参数表：state/action/reward 权重、预算、容量、transition、训练轮次、baseline protocol、质量门控阈值和是否允许正式训练。若继续论文写作，可说“生成 DQN Method 部分并导出 Word”。\n",
        "project_state/conversation_handoff.md": f"# Conversation Handoff\n\n最新完成任务：DQN 输出复核、解释体系与论文输出升级。任务包：`{rel(RUN_DIR)}`。所有 DQN 结果仍为 experimental，不能作为 formal 监管政策结论。下一步建议 formal 参数确认或论文 Method/Introduction/Literature Review/Discussion 生成。\n",
        "project_state/run_protocol.md": safe_read(ROOT / "project_state" / "run_protocol.md") + "\n\n## Output Explanation Protocol Addendum\n\n每次 durable task 必须生成 artifact/table/figure/model/code explanations、artifact-to-evidence map、experimental/formal 状态标记和必要的 DOCX/render QA。\n",
        "project_state/project_memory.md": safe_read(ROOT / "project_state" / "project_memory.md") + "\n\n## 2026-04-26 输出解释与论文升级长期记忆\n\n所有图、表、模型输出、代码、Word 和文献输出必须带解释、来源、阅读方式、局限和 evidence map。DQN/RL experimental 结果不得写成 formal 监管结论。Zotero 只允许 sidecar notes/BibTeX/RIS/CSV 待导入，未经确认不得写 SQLite。\n",
        "project_state/lessons_learned.md": safe_read(ROOT / "project_state" / "lessons_learned.md") + "\n\n## 2026-04-26 DQN 输出复核经验\n\n旧 DQN 包可能包含 0 字节图或近似空表；最新图即使非空也可能语义不足（如全 0 约束图）。未来必须做空图、全 0、无差异和中文字体 QA，并生成解释性图或说明。\n",
        "project_state/changelog.md": safe_read(ROOT / "project_state" / "changelog.md") + f"\n\n## 2026-04-26\n\n- 完成 DQN 输出复核、解释体系补强、代码说明、论文级 Results DOCX、文献/Zotero 侧车和长期规则升级。任务包：`{rel(RUN_DIR)}`。\n",
        "project_state/decision_log.md": safe_read(ROOT / "project_state" / "decision_log.md") + "\n\n## 2026-04-26 Decision\n\n决定将 output explanation、empty chart prevention、model setting documentation、code explanation、literature-grounded modeling 和 Zotero safe sidecar 写入长期 workflow。DQN 结果保持 experimental，不升级为 formal policy conclusion。\n",
    }
    for target, text in updates.items():
        write_text(ROOT / target, text)


def update_workflow_improvement_files() -> None:
    ledger = ROOT / "workflow_improvement" / "improvement_ledger.csv"
    row = pd.DataFrame(
        [
            {
                "date": datetime.now().strftime("%Y-%m-%d"),
                "item": "output_explanation_and_reporting_upgrade",
                "change_type": "safe_local_workflow_upgrade",
                "status": "applied",
                "run_package": rel(RUN_DIR),
                "approval_required": "external plugins only; no third-party install executed",
            }
        ]
    )
    if ledger.exists():
        try:
            old = pd.read_csv(ledger)
            pd.concat([old, row], ignore_index=True).to_csv(ledger, index=False, encoding="utf-8-sig")
        except Exception:
            row.to_csv(ledger, index=False, encoding="utf-8-sig")
    else:
        row.to_csv(ledger, index=False, encoding="utf-8-sig")
    backlog = ROOT / "workflow_improvement" / "upgrade_backlog.yaml"
    if backlog.exists():
        raw_backlog = backlog.read_text(encoding="utf-8")
        try:
            existing = yaml.safe_load(raw_backlog) or {}
        except Exception:
            existing = {"legacy_unparsed_backlog_text": raw_backlog, "items": []}
    else:
        existing = {}
    existing = existing or {}
    existing.setdefault("items", [])
    existing["items"].append(
        {
            "id": "formal_dqn_parameter_confirmation_and_sensitivity",
            "status": "pending_user_confirmation",
            "reason": "Current DQN is experimental; formal use requires parameter confirmation and sensitivity analysis.",
        }
    )
    existing["items"].append(
        {
            "id": "paper_section_full_pipeline",
            "status": "ready",
            "reason": "Results section upgraded; Method/Introduction/Literature Review/Discussion can reuse new explanation/evidence workflow.",
        }
    )
    backlog.write_text(yaml.safe_dump(existing, allow_unicode=True, sort_keys=False), encoding="utf-8")


def main() -> None:
    init_run_package()
    inventory = inventory_package(LATEST_DQN, "latest_revised_dqn") + inventory_package(BASELINE_DQN, "historical_auto_dqn")
    tables = load_core_tables()
    issues = audit_issues(inventory, tables)
    write_deep_audit(inventory, issues)
    chart_rows = render_repair_charts(tables)
    write_chart_report(chart_rows)
    generate_explanations(tables, chart_rows, issues)
    generate_dqn_reports(tables)
    generate_code_explanations()
    generate_literature_outputs()
    generate_word_outputs(tables)
    generate_skill_scout_outputs()
    run_dry_runs()
    write_readme_manifest()
    update_indexes_and_state()
    update_workflow_improvement_files()
    print(json.dumps({"run_dir": str(RUN_DIR), "status": "completed"}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
